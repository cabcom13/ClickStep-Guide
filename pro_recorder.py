import sys
import os
import time
import json
import shutil
import math
import uuid
from datetime import datetime

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QLabel, QListWidget, QListWidgetItem,
                             QGraphicsView, QGraphicsScene, QGraphicsPixmapItem, QGraphicsItem,
                             QDockWidget, QTextEdit, QFileDialog, QMessageBox, QInputDialog,
                             QGraphicsEllipseItem, QGraphicsRectItem, QGraphicsTextItem,
                             QScrollArea, QColorDialog, QSlider, QSpinBox, QCheckBox,
                             QListWidgetItem, QAbstractItemView, QFontComboBox, QToolButton, QProgressBar)
from PyQt6.QtCore import Qt, QSize, QRectF, QPointF, pyqtSignal, QObject, QThread, QRect
from PyQt6.QtGui import (QColor, QFont, QIcon, QPixmap, QImage, QPainter, QPen, QBrush, 
                        QPainterPath, QTransform, QCursor, QFontDatabase)

from pynput import mouse, keyboard
from PIL import ImageGrab
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==================== DATA MODELS ====================

# ==================== DATA MODELS ====================

class RecordingOverlay(QWidget):
    """Transparent overlay to show recording status"""
    def __init__(self):
        super().__init__()
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint | Qt.WindowType.X11BypassWindowManagerHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents)
        
        layout = QHBoxLayout(self)
        self.lbl = QLabel("🔴 REC")
        self.lbl.setStyleSheet("color: red; font-weight: bold; font-size: 24px; background: rgba(0,0,0,180); padding: 5px 15px; border-radius: 15px; border: 1px solid red;")
        layout.addWidget(self.lbl)
        
        # Position top right
        screen = QApplication.primaryScreen().geometry()
        self.setGeometry(screen.width() - 170, 30, 150, 60)
        
        # Exclude from capture (Win 10 2004+)
        if os.name == 'nt':
            try:
                import ctypes
                # WDA_EXCLUDEFROMCAPTURE = 0x00000011
                ctypes.windll.user32.SetWindowDisplayAffinity(int(self.winId()), 0x00000011)
            except:
                pass

class Layer:
    def __init__(self, ltype, data, label, is_global=False, uid=None):
        self.type = ltype
        self.data = data
        self.label = label
        self.is_global = is_global
        self.visible = True
        self.uid = uid if uid else str(uuid.uuid4())

class LayerListWidget(QListWidget):
    """Custom ListWidget to handle layer drag and drop"""
    def __init__(self, parent=None, is_global_list=False, main_window=None):
        super().__init__(parent)
        self.is_global_list = is_global_list
        self.main_window = main_window
        self.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.setDragDropMode(QListWidget.DragDropMode.DragDrop)
        self.setDefaultDropAction(Qt.DropAction.MoveAction)
        
    def dropEvent(self, event):
        super().dropEvent(event)
        if self.main_window:
            self.main_window.handle_layer_drop(self.is_global_list)

# ==================== INTERACTIVE GRAPHICS ITEMS ====================

class ResizableRectItem(QGraphicsRectItem):
    """Photoshop-style resizable rectangle with corner handles"""
    def __init__(self, rect, color=QColor(255, 255, 255, 100), label=""):
        super().__init__(rect)
        self.item_type = 'rect' # Added item_type
        self.label = label
        self.handle_size = 10
        self.selected_handle = None
        self.base_color = color
        
        self.setPen(QPen(QColor(255, 255, 255), 2, Qt.PenStyle.DashLine))
        self.setBrush(QBrush(color))
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setAcceptHoverEvents(True)
        self.setCursor(Qt.CursorShape.SizeAllCursor)

    def get_handles(self):
        """Return positions of 8 resize handles"""
        r = self.rect()
        return {
            'tl': QPointF(r.left(), r.top()),
            'tr': QPointF(r.right(), r.top()),
            'bl': QPointF(r.left(), r.bottom()),
            'br': QPointF(r.right(), r.bottom()),
            't': QPointF(r.center().x(), r.top()),
            'b': QPointF(r.center().x(), r.bottom()),
            'l': QPointF(r.left(), r.center().y()),
            'r': QPointF(r.right(), r.center().y())
        }

    def get_handle_at(self, pos):
        """Check if position is over a handle"""
        handles = self.get_handles()
        for name, handle_pos in handles.items():
            if (pos - handle_pos).manhattanLength() < self.handle_size:
                return name
        return None

    def hoverMoveEvent(self, event):
        handle = self.get_handle_at(event.pos())
        if handle:
            cursors = {
                'tl': Qt.CursorShape.SizeFDiagCursor, 'br': Qt.CursorShape.SizeFDiagCursor,
                'tr': Qt.CursorShape.SizeBDiagCursor, 'bl': Qt.CursorShape.SizeBDiagCursor,
                't': Qt.CursorShape.SizeVerCursor, 'b': Qt.CursorShape.SizeVerCursor,
                'l': Qt.CursorShape.SizeHorCursor, 'r': Qt.CursorShape.SizeHorCursor
            }
            self.setCursor(cursors.get(handle, Qt.CursorShape.SizeAllCursor))
        else:
            self.setCursor(Qt.CursorShape.SizeAllCursor)
        super().hoverMoveEvent(event)

    def mousePressEvent(self, event):
        self.selected_handle = self.get_handle_at(event.pos())
        if not self.selected_handle:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.selected_handle:
            r = self.rect()
            pos = event.pos()
            
            if 'l' in self.selected_handle:
                r.setLeft(pos.x())
            if 'r' in self.selected_handle:
                r.setRight(pos.x())
            if 't' in self.selected_handle:
                r.setTop(pos.y())
            if 'b' in self.selected_handle:
                r.setBottom(pos.y())
            
            # Ensure minimum size
            if r.width() < 20: r.setWidth(20)
            if r.height() < 20: r.setHeight(20)
            
            self.setRect(r.normalized())
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self.selected_handle = None
        super().mouseReleaseEvent(event)

    def boundingRect(self):
        r = self.rect()
        margin = self.handle_size
        return r.adjusted(-margin, -margin, margin, margin)

    def paint(self, painter, option, widget):
        super().paint(painter, option, widget)
        
        # Draw handles when selected
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

class BlurItem(ResizableRectItem):
    """Blur annotation with Photoshop-style interaction"""
    def __init__(self, rect, is_global=False, bg_pixmap=None, uid=None):
        color = QColor(0, 255, 0, 80) if is_global else QColor(0, 0, 0, 150)
        super().__init__(rect, color, "Globaler Blur" if is_global else "Zensur")
        self.is_global = is_global
        self.item_type = 'blur'
        self.bg_pixmap = bg_pixmap
        self.blurred_cache = None
        self.last_pos = None
        self.last_rect = None
        self.uid = uid if uid else str(uuid.uuid4())
        # CacheMode disabled because we handle it manually
        self.setCacheMode(QGraphicsItem.CacheMode.NoCache)
    
    def itemChange(self, change, value):
        # Trigger repaint when position changes
        if change == QGraphicsItem.GraphicsItemChange.ItemPositionChange:
            self.prepareGeometryChange()
        return super().itemChange(change, value)
    
    def paint(self, painter, option, widget):
        # Render actual blurred content
        if self.bg_pixmap and self.rect().isValid():
            r = self.rect().toRect()
            pos = self.scenePos()
            
            # Check cache validity
            should_update = (
                self.blurred_cache is None or
                pos != self.last_pos or
                r != self.last_rect
            )
            
            if should_update:
                # Extract region from background
                source_rect = QRect(int(pos.x() + r.x()), int(pos.y() + r.y()), r.width(), r.height())
                
                # Ensure within bounds
                if source_rect.intersects(QRect(0, 0, self.bg_pixmap.width(), self.bg_pixmap.height())):
                    cropped = self.bg_pixmap.copy(source_rect)
                    
                    # Convert to QImage for blur
                    img = cropped.toImage()
                    if not img.isNull():
                        # Convert to numpy for OpenCV blur
                        width, height = img.width(), img.height()
                        ptr = img.bits()
                        ptr.setsize(height * width * 4)
                        arr = np.frombuffer(ptr, np.uint8).reshape((height, width, 4))

                        # Apply Gaussian Blur
                        blurred = cv2.GaussianBlur(arr, (75, 75), 40)

                        # Convert back to QImage
                        qimg = QImage(blurred.data, width, height, width * 4, QImage.Format.Format_RGBA8888)
                        self.blurred_cache = QPixmap.fromImage(qimg)
                        self.last_pos = pos
                        self.last_rect = r
                    else:
                        self.blurred_cache = None
                else:
                    self.blurred_cache = None

            if self.blurred_cache:
                painter.drawPixmap(r, self.blurred_cache)
        
        # Draw border
        color = QColor(0, 255, 0) if self.is_global else QColor(255, 255, 255)
        painter.setPen(QPen(color, 2 if self.isSelected() else 1))
        painter.setBrush(Qt.BrushStyle.NoBrush)
        painter.drawRect(self.rect())
        
        if self.is_global:
            painter.setFont(QFont("Segoe UI", 8))
            painter.setPen(QPen(color, 1))
            painter.drawText(self.rect().adjusted(5, 5, -5, -5), Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft, "GLOBAL")
        
        # Draw handles when selected
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

class ZoomItem(QGraphicsItem):
    """Interactive zoom box with target marker"""
    def __init__(self, rect, target_point, pixmap=None, is_global=False, uid=None):
        super().__init__()
        self.item_type = 'zoom'
        self.is_global = is_global
        self.uid = uid if uid else str(uuid.uuid4())
        self.box_rect = rect
        self.target = target_point
        self.pixmap = pixmap
        self.handle_size = 10
        self.selected_handle = None
        
        self.cached_zoom = None
        self.last_box_rect = None
        self.last_target = None

        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setAcceptHoverEvents(True)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        self.setCacheMode(QGraphicsItem.CacheMode.NoCache)
    
    def itemChange(self, change, value):
        # Update bounding rect when position changes
        if change == QGraphicsItem.GraphicsItemChange.ItemPositionChange:
            self.prepareGeometryChange()
        return super().itemChange(change, value)

    def boundingRect(self):
        # Include arrow in bounding box - MUST be in local coordinates
        # Target is global (scene) coords, self.pos() is global (scene) coords
        local_target = self.target - self.pos()
        box_center = self.box_rect.center()
        
        # Rect covering box and line to target
        arrow_rect = QRectF(local_target, box_center).normalized()
        
        # Merge with box rect and add buffer for stroke width and handles
        return self.box_rect.united(arrow_rect).adjusted(-20, -20, 20, 20)

    def get_handles(self):
        r = self.box_rect
        return {
            'tl': QPointF(r.left(), r.top()),
            'tr': QPointF(r.right(), r.top()),
            'bl': QPointF(r.left(), r.bottom()),
            'br': QPointF(r.right(), r.bottom())
        }

    def get_handle_at(self, pos):
        handles = self.get_handles()
        for name, handle_pos in handles.items():
            if (pos - handle_pos).manhattanLength() < self.handle_size:
                return name
        return None

    def hoverMoveEvent(self, event):
        handle = self.get_handle_at(event.pos())
        if handle:
            cursors = {
                'tl': Qt.CursorShape.SizeFDiagCursor, 'br': Qt.CursorShape.SizeFDiagCursor,
                'tr': Qt.CursorShape.SizeBDiagCursor, 'bl': Qt.CursorShape.SizeBDiagCursor
            }
            self.setCursor(cursors.get(handle, Qt.CursorShape.SizeAllCursor))
        else:
            self.setCursor(Qt.CursorShape.SizeAllCursor)

    def mousePressEvent(self, event):
        self.selected_handle = self.get_handle_at(event.pos())
        if not self.selected_handle:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.selected_handle:
            pos = event.pos()
            r = self.box_rect
            
            if 'l' in self.selected_handle:
                r.setLeft(pos.x())
            if 'r' in self.selected_handle:
                r.setRight(pos.x())
            if 't' in self.selected_handle:
                r.setTop(pos.y())
            if 'b' in self.selected_handle:
                r.setBottom(pos.y())
            
            if r.width() < 50: r.setWidth(50)
            if r.height() < 50: r.setHeight(50)
            
            self.box_rect = r.normalized()
            self.update()
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self.selected_handle = None
        super().mouseReleaseEvent(event)

    def paint(self, painter, option, widget):
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # 1. Draw Arrow (from box center to target, accounting for item position)
        center = self.box_rect.center()
        # Target is in scene coordinates, center is in item coordinates
        target_in_item = self.mapFromScene(self.target)
        painter.setPen(QPen(QColor(255, 255, 255, 200), 3, Qt.PenStyle.DashLine))
        painter.drawLine(center, target_in_item)
        
        # 2. Draw Zoom Box
        painter.setPen(QPen(QColor(255, 255, 255), 3))
        painter.setBrush(QBrush(QColor(20, 20, 20)))
        painter.drawRect(self.box_rect)
        
        # 3. Draw Magnified Content (if pixmap available)
        if self.pixmap:
            if (self.cached_zoom is None or
                self.box_rect != self.last_box_rect or
                self.target != self.last_target):

                # Crop area around target
                src_sz = self.box_rect.width() / 2
                crop_rect = QRectF(self.target.x()-src_sz/2, self.target.y()-src_sz/2, src_sz, src_sz)

                # Simple magnification logic
                cropped = self.pixmap.copy(crop_rect.toRect())
                self.cached_zoom = cropped.scaled(self.box_rect.size().toSize(),
                                      Qt.AspectRatioMode.KeepAspectRatio,
                                      Qt.TransformationMode.SmoothTransformation)
                self.last_box_rect = self.box_rect
                self.last_target = self.target

            if self.cached_zoom:
                painter.drawPixmap(self.box_rect.toRect(), self.cached_zoom)
            
            # Draw marker in zoom
            painter.setPen(QPen(QColor(255, 0, 0), 2))
            ic = self.box_rect.center()
            painter.drawEllipse(ic, 8, 8)

        # 4. Draw Handles
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

class EditableTextItem(QGraphicsTextItem):
    """Editable text with Photoshop-style interaction"""
    def __init__(self, text, color=QColor(255, 255, 255), is_global=False, uid=None):
        super().__init__(text)
        self.item_type = 'text'
        self.is_global = is_global
        self.uid = uid if uid else str(uuid.uuid4())
        self.setDefaultTextColor(color)
        self.setFont(QFont("Segoe UI", 18, QFont.Weight.Bold))
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemIsFocusable)
        self.setTextInteractionFlags(Qt.TextInteractionFlag.NoTextInteraction)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        self.handle_size = 12
        self.setAcceptHoverEvents(True)
        self.resizing = False
        self.resize_start_pos = None
        self.initial_font_size = 0

    def boundingRect(self):
        r = super().boundingRect()
        # Add space for handle at bottom right
        margin = self.handle_size
        return r.adjusted(0, 0, margin, margin)

    def paint(self, painter, option, widget):
        super().paint(painter, option, widget)
        if self.isSelected():
            # Draw resize handle (bottom right)
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            r = super().boundingRect() # Use original rect for handle position
            painter.drawEllipse(r.bottomRight(), self.handle_size/2, self.handle_size/2)

    def hoverMoveEvent(self, event):
        # Check if over handle
        r = self.boundingRect()
        handle_pos = r.bottomRight()
        if (event.pos() - handle_pos).manhattanLength() < self.handle_size:
            self.setCursor(Qt.CursorShape.SizeFDiagCursor)
        else:
            self.setCursor(Qt.CursorShape.SizeAllCursor)
        super().hoverMoveEvent(event)

    def mousePressEvent(self, event):
        r = self.boundingRect()
        handle_pos = r.bottomRight()
        if (event.pos() - handle_pos).manhattanLength() < self.handle_size:
            self.resizing = True
            self.resize_start_pos = event.scenePos()
            self.initial_font_size = self.font().pointSize()
            if self.initial_font_size <= 0: self.initial_font_size = 12
            event.accept()
        else:
            self.resizing = False
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.resizing:
            self.prepareGeometryChange()
            delta = event.scenePos() - self.resize_start_pos
            # Use Y delta to scale
            scale_factor = 1 + (delta.y() / 100.0)
            new_size = int(self.initial_font_size * scale_factor)
            new_size = max(8, min(new_size, 200)) # Clamp
            
            f = self.font()
            f.setPointSize(new_size)
            self.setFont(f)
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self.resizing = False
        super().mouseReleaseEvent(event)

    def update_font(self, family=None, size=None, bold=None, italic=None, underline=None):
        font = self.font()
        if family: font.setFamily(family)
        if size and size > 0: font.setPointSize(size)
        if bold is not None: font.setBold(bold)
        if italic is not None: font.setItalic(italic)
        if underline is not None: font.setUnderline(underline)
        self.setFont(font)

    def mouseDoubleClickEvent(self, event):
        # Open simple text dialog
        text, ok = QInputDialog.getText(None, "Text ändern", "Neuer Text:", text=self.toPlainText())
        if ok and text:
            self.setPlainText(text)

    def focusOutEvent(self, event):
        self.setTextInteractionFlags(Qt.TextInteractionFlag.NoTextInteraction)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        super().focusOutEvent(event)

class ClickMarkerItem(QGraphicsEllipseItem):
    """Professional click marker with number"""
    def __init__(self, x, y, number, color=QColor(0, 0, 255)):
        # Increased radius to 50 (100x100) to fit selection highlight (45) and pen width
        super().__init__(x-50, y-50, 100, 100)
        self.item_type = 'click'
        self.center_x = x
        self.center_y = y
        self.number = number
        self.marker_color = color
        
        self.setPen(QPen(QColor(255, 255, 255), 6))
        self.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        self.setCacheMode(QGraphicsItem.CacheMode.NoCache)
    
    def itemChange(self, change, value):
        # Update center coordinates when moved
        if change == QGraphicsItem.GraphicsItemChange.ItemPositionHasChanged:
            new_center = self.rect().center() + self.pos()
            self.center_x = int(new_center.x())
            self.center_y = int(new_center.y())
            
            # Update all zoom items that target this click marker
            if self.scene():
                for item in self.scene().items():
                    if hasattr(item, 'item_type') and item.item_type == 'zoom':
                        # Update zoom target to new click position
                        item.target = new_center
                        item.prepareGeometryChange()
                        item.update()
        return super().itemChange(change, value)

    def paint(self, painter, option, widget):
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        center = self.rect().center()
        
        # Outer white ring
        painter.setPen(QPen(QColor(255, 255, 255), 6))
        painter.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        painter.drawEllipse(center, 38, 38)
        
        # Inner colored ring
        painter.setPen(QPen(self.marker_color, 4))
        painter.drawEllipse(center, 32, 32)
        
        # Number
        painter.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        painter.setPen(QPen(QColor(0, 0, 0), 6))
        painter.drawText(self.rect(), Qt.AlignmentFlag.AlignCenter, self.number)
        painter.setPen(QPen(QColor(255, 255, 255), 3))
        painter.drawText(self.rect(), Qt.AlignmentFlag.AlignCenter, self.number)
        
        # Selection highlight
        if self.isSelected():
            painter.setPen(QPen(QColor(0, 175, 255), 3))
            painter.setBrush(QBrush(Qt.BrushStyle.NoBrush))
            painter.drawEllipse(center, 45, 45)

# ==================== EDITOR SCENE ====================

class EditorScene(QGraphicsScene):
    """Custom scene with drawing tools"""
    def __init__(self, editor):
        super().__init__()
        self.editor = editor
        self.draw_start = None
        self.preview_item = None

    def mousePressEvent(self, event):
        if self.editor.draw_mode and event.button() == Qt.MouseButton.LeftButton:
            self.draw_start = event.scenePos()
            
            if self.editor.draw_mode == 'text':
                text, ok = QInputDialog.getText(None, "Text", "Eingabe:")
                if ok and text:
                    color_dialog = QColorDialog()
                    color = color_dialog.getColor()
                    if color.isValid():
                        item = EditableTextItem(text, color)
                        item.setPos(self.draw_start)
                        self.addItem(item)
                self.editor.draw_mode = None
                self.editor.update_tool_buttons()
        else:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.draw_start and self.editor.draw_mode in ['blur', 'global_blur', 'zoom', 'crop']:
            # Live preview
            if self.preview_item:
                try:
                    self.removeItem(self.preview_item)
                except RuntimeError:
                    self.preview_item = None
            
            rect = QRectF(self.draw_start, event.scenePos()).normalized()
            colors = {
                'blur': QColor(0, 0, 0, 100),
                'global_blur': QColor(0, 255, 0, 100),
                'zoom': QColor(0, 175, 255, 100),
                'crop': QColor(255, 255, 0, 80)  # Yellow for crop
            }
            
            self.preview_item = QGraphicsRectItem(rect)
            self.preview_item.setPen(QPen(QColor(255, 255, 255), 2, Qt.PenStyle.DashLine))
            self.preview_item.setBrush(QBrush(colors.get(self.editor.draw_mode, QColor(255, 255, 255, 50))))
            self.addItem(self.preview_item)
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self.draw_start and self.editor.draw_mode:
            rect = QRectF(self.draw_start, event.scenePos()).normalized()
            
            if rect.width() > 20 and rect.height() > 20:
                if self.editor.draw_mode in ['blur', 'global_blur', 'zoom']:
                     self.editor.push_undo() # Undo before create
                
                if self.editor.draw_mode == 'blur':
                    item = BlurItem(rect, False, self.editor.current_pixmap)
                    self.addItem(item)
                elif self.editor.draw_mode == 'global_blur':
                    item = BlurItem(rect, True, self.editor.current_pixmap)
                    self.addItem(item)
                elif self.editor.draw_mode == 'zoom':
                    # Target is the center of the current step's click
                    target = QPointF(self.editor.steps[self.editor.current_idx].x, 
                                   self.editor.steps[self.editor.current_idx].y)
                    item = ZoomItem(rect, target, self.editor.current_pixmap)
                    self.addItem(item)
                elif self.editor.draw_mode == 'crop':
                    self.editor.push_undo() # Undo crop
                    
                    # Apply global crop and reload all steps
                    self.editor.global_crop = (int(rect.left()), int(rect.top()), 
                                             int(rect.right()), int(rect.bottom()))
                    
                    # Prevent further preview updates during reload
                    self.draw_start = None
                    if self.preview_item:
                        try:
                            self.removeItem(self.preview_item)
                        except RuntimeError:
                            pass
                        self.preview_item = None

                    # Reload current step with crop applied
                    self.editor.load_step(self.editor.current_idx)
                    
                    # Update all thumbnails to show crop
                    self.editor.update_thumbnails()
                    
                    QMessageBox.information(None, "Crop", 
                        f"Globaler Ausschnitt festgelegt: {rect.width():.0f}x{rect.height():.0f}px\n" +
                        "Wird auf alle Schritte angewendet!")
                    self.editor.draw_mode = None
                    self.editor.update_tool_buttons()
            
            if self.preview_item:
                try:
                    if self.preview_item.scene() == self:
                        self.removeItem(self.preview_item)
                except RuntimeError:
                    pass # Item already deleted (e.g. by scene.clear())
                self.preview_item = None
            
            self.draw_start = None
            self.editor.draw_mode = None
            self.editor.update_tool_buttons()
            self.editor.refresh_layer_list()
        else:
            super().mouseReleaseEvent(event)
            self.editor.refresh_layer_list()

# ==================== MAIN EDITOR ====================

class ZoomableGraphicsView(QGraphicsView):
    """Graphics view with mouse wheel zoom and middle button pan"""
    def __init__(self, scene):
        super().__init__(scene)
        self.zoom_factor = 1.0
        self.panning = False
        self.pan_start_pos = None
        
    def wheelEvent(self, event):
        # Zoom towards mouse cursor
        zoom_in_factor = 1.15
        zoom_out_factor = 1 / zoom_in_factor
        
        # Get the old position
        old_pos = self.mapToScene(event.position().toPoint())
        
        if event.angleDelta().y() > 0:
            factor = zoom_in_factor
            self.zoom_factor *= factor
        else:
            factor = zoom_out_factor
            self.zoom_factor *= factor
        
        # Limit zoom range
        if 0.1 < self.zoom_factor < 10:
            self.scale(factor, factor)
            
            # Get the new position
            new_pos = self.mapToScene(event.position().toPoint())
            
            # Move scene to keep mouse position
            delta = new_pos - old_pos
            self.translate(delta.x(), delta.y())
        else:
            self.zoom_factor /= factor  # Revert if out of range
    
    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.MiddleButton:
            # Start panning with middle mouse button
            self.panning = True
            self.pan_start_pos = event.pos()
            self.setCursor(Qt.CursorShape.ClosedHandCursor)
            event.accept()
        else:
            super().mousePressEvent(event)
    
    def mouseMoveEvent(self, event):
        if self.panning and self.pan_start_pos:
            # Pan the view
            delta = event.pos() - self.pan_start_pos
            self.pan_start_pos = event.pos()
            
            # Scroll the view
            self.horizontalScrollBar().setValue(self.horizontalScrollBar().value() - delta.x())
            self.verticalScrollBar().setValue(self.verticalScrollBar().value() - delta.y())
            event.accept()
        else:
            super().mouseMoveEvent(event)
    
    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MouseButton.MiddleButton:
            # Stop panning
            self.panning = False
            self.pan_start_pos = None
            self.setCursor(Qt.CursorShape.ArrowCursor)
            event.accept()
        else:
            super().mouseReleaseEvent(event)

class ProEditor(QMainWindow):
    def __init__(self, steps, globals, crop, save_cb, parent_window=None):
        super().__init__()
        self.recorder_window = parent_window
        self.setWindowTitle("ClickStep Guide Editor - Photoshop Style")
        self.resize(1600, 1000)
        self.setStyleSheet(self.get_stylesheet())
        
        self.steps = steps
        self.global_layers = globals
        self.global_crop = crop
        self.save_cb = save_cb
        self.current_idx = 0
        self.draw_mode = None
        self.current_project_name = None
        
        self.undo_stack = [] # List of snapshots
        
        self.scene = EditorScene(self)
        self.scene.selectionChanged.connect(self.update_properties)
        self.view = ZoomableGraphicsView(self.scene)
        self.view.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.view.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.view.setDragMode(QGraphicsView.DragMode.RubberBandDrag)
        self.view.setBackgroundBrush(QBrush(QColor("#0d0d0d")))
        # Use SmartViewportUpdate for better performance; artifacts handled by correct boundingRect
        self.view.setViewportUpdateMode(QGraphicsView.ViewportUpdateMode.SmartViewportUpdate)
        
        self.setCentralWidget(self.view)
        self.setup_ui()
        
        # Subtle Branding Footer
        self.statusBar().showMessage("ClickStep Guide Pro Engine | Professionelles Dokumentations-System Enabled")
        self.statusBar().setStyleSheet("color: #555; background: #1a1a1b; border-top: 1px solid #333;")
        
        if self.steps:
            self.load_step(0)

    def closeEvent(self, event):
        """Show recorder window when editor is closed"""
        if self.recorder_window:
            self.recorder_window.show()
        super().closeEvent(event)

    def get_stylesheet(self):
        return """
            QMainWindow, QWidget { 
                background-color: #1e1e1e; 
                color: #cccccc; 
                font-family: "Segoe UI", sans-serif;
                font-size: 13px;
            }
            
            /* Docks & Toolbars */
            QDockWidget {
                titlebar-close-icon: url(none);
                titlebar-normal-icon: url(none);
                border: 1px solid #333333;
            }
            QDockWidget::title {
                background: #252526;
                padding: 8px;
                font-weight: bold;
                color: #e0e0e0;
                border-bottom: 1px solid #333333;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }
            QToolBar {
                background: #252526;
                border-bottom: 1px solid #333333;
                spacing: 8px;
                padding: 6px;
            }
            QToolBar::separator {
                background: #3e3e42;
                width: 1px;
                margin: 4px 8px;
            }
            
            /* Lists */
            QListWidget {
                background-color: #252526;
                border: 1px solid #333333;
                border-radius: 4px;
                outline: none;
            }
            QListWidget::item {
                padding: 8px;
                border-bottom: 1px solid #2d2d2d;
                color: #cccccc;
            }
            QListWidget::item:selected {
                background-color: #094771; /* VS Code Blue Selection */
                color: white;
                border-left: 3px solid #007acc;
            }
            QListWidget::item:hover {
                background-color: #2a2d2e;
            }
            
            /* Buttons */
            QPushButton {
                background-color: #333333;
                color: #ffffff;
                border: 1px solid #3e3e42;
                padding: 6px 14px;
                border-radius: 4px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #3e3e42;
                border-color: #505050;
            }
            QPushButton:pressed {
                background-color: #1e1e1e;
                border-color: #007acc;
            }
            QPushButton:checked {
                background-color: #094771;
                border-color: #007acc;
                color: white;
            }
            QPushButton:disabled {
                background-color: #252526;
                color: #666666;
                border-color: #333333;
            }
            
            /* Special Buttons via ObjectName */
            QPushButton#AccentButton {
                background-color: #0078d4;
                border: 1px solid #0078d4;
                font-size: 14px;
            }
            QPushButton#AccentButton:hover {
                background-color: #106ebe;
            }
            QPushButton#DestructiveButton {
                background-color: transparent;
                border: 1px solid #d13438;
                color: #d13438;
            }
            QPushButton#DestructiveButton:hover {
                background-color: #d13438;
                color: white;
            }
            
            /* Scrollbars */
            QScrollBar:vertical {
                border: none;
                background: #1e1e1e;
                width: 12px;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #424242;
                min-height: 20px;
                border-radius: 6px;
                margin: 2px;
            }
            QScrollBar::handle:vertical:hover {
                background: #606060;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
        """

    def setup_ui(self):
        # Toolbar
        toolbar = self.addToolBar("Tools")
        toolbar.setMovable(False)
        toolbar.setIconSize(QSize(24, 24))
        
        self.btn_select = QPushButton("🏹 Select")
        self.btn_select.setCheckable(True)
        self.btn_select.setChecked(True)
        self.btn_select.clicked.connect(lambda: self.set_tool(None))
        
        self.btn_blur = QPushButton("🛡️ Blur")
        self.btn_blur.setCheckable(True)
        self.btn_blur.clicked.connect(lambda: self.set_tool('blur'))
        
        self.btn_zoom = QPushButton("🔍 Zoom")
        self.btn_zoom.setCheckable(True)
        self.btn_zoom.clicked.connect(lambda: self.set_tool('zoom'))
        
        self.btn_text = QPushButton("📝 Text")
        self.btn_text.setCheckable(True)
        self.btn_text.clicked.connect(lambda: self.set_tool('text'))
        
        self.btn_global_blur = QPushButton("🌍 Global Blur")
        self.btn_global_blur.setCheckable(True)
        self.btn_global_blur.clicked.connect(lambda: self.set_tool('global_blur'))
        
        self.btn_crop = QPushButton("✂️ Crop")
        self.btn_crop.setCheckable(True)
        self.btn_crop.clicked.connect(lambda: self.set_tool('crop'))
        
        toolbar.addWidget(self.btn_select)
        toolbar.addWidget(self.btn_blur)
        toolbar.addWidget(self.btn_zoom)
        toolbar.addWidget(self.btn_text)
        toolbar.addSeparator()
        toolbar.addWidget(self.btn_global_blur)
        toolbar.addWidget(self.btn_crop)
        toolbar.addSeparator()
        
        btn_delete = QPushButton("🗑️ Delete Layer")
        btn_delete.setObjectName("DestructiveButton")
        btn_delete.clicked.connect(self.delete_selected)
        toolbar.addWidget(btn_delete)
        
        btn_delete_step = QPushButton("❌ Step")
        btn_delete_step.setObjectName("DestructiveButton")
        btn_delete_step.setToolTip("Schritt löschen")
        btn_delete_step.clicked.connect(self.delete_step)
        toolbar.addWidget(btn_delete_step)
        
        toolbar.addSeparator()
        
        btn_save = QPushButton("� Speichern")
        btn_save.setToolTip("Strg+S")
        btn_save.setShortcut("Ctrl+S")
        btn_save.clicked.connect(lambda chk=False: self.save_project(save_as=False))
        toolbar.addWidget(btn_save)

        btn_save_as = QPushButton("📝 Unter...")
        btn_save_as.setToolTip("Speichern unter...")
        btn_save_as.clicked.connect(lambda chk=False: self.save_project(save_as=True))
        toolbar.addWidget(btn_save_as)
        
        toolbar.addSeparator()
        
        btn_undo = QPushButton("↩️ Zurück")
        btn_undo.setShortcut("Ctrl+Z")
        btn_undo.clicked.connect(self.undo)
        toolbar.addWidget(btn_undo)
        
        btn_export = QPushButton("💾 EXPORT")
        btn_export.setObjectName("AccentButton")
        btn_export.clicked.connect(self.on_export_clicked)
        toolbar.addWidget(btn_export)

        # Left Dock: Thumbnails
        self.dock_thumbs = QDockWidget("SCHRITTE", self)
        self.thumb_list = QListWidget()
        self.thumb_list.currentRowChanged.connect(self.on_step_changed)
        
        # Enable Drag and Drop reordering
        self.thumb_list.setDragDropMode(QAbstractItemView.DragDropMode.InternalMove)
        self.thumb_list.model().rowsMoved.connect(self.on_steps_reordered)
        
        self.dock_thumbs.setWidget(self.thumb_list)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.dock_thumbs)
        
        # Right Dock: Layers (Split into Step and Global)
        self.dock_layers = QDockWidget("EBENEN", self)
        self.layer_panel = QWidget()
        self.layer_layout = QVBoxLayout(self.layer_panel)
        
        # Step Layers
        step_label = QLabel("PROJEKT ENTWURF")
        step_label.setStyleSheet("font-weight: bold; color: #cccccc; padding: 5px; background: #252526;")
        self.layer_layout.addWidget(step_label)
        
        # Separator for clarity
        
        lbl_local = QLabel("BILD-EBENEN")
        lbl_local.setStyleSheet("font-size: 10px; font-weight: bold; color: #007acc; margin-top: 10px;")
        self.layer_layout.addWidget(lbl_local)
        
        self.step_layer_list = LayerListWidget(is_global_list=False, main_window=self)
        self.step_layer_list.itemClicked.connect(self.on_layer_clicked)
        self.layer_layout.addWidget(self.step_layer_list)
        
        # Buttons to move between categories
        btn_layout = QHBoxLayout()
        btn_to_global = QPushButton("↓ Zu Global")
        btn_to_global.clicked.connect(self.move_layer_to_global)
        btn_to_step = QPushButton("↑ Zu Lokal")
        btn_to_step.clicked.connect(self.move_layer_to_step)
        btn_layout.addWidget(btn_to_global)
        btn_layout.addWidget(btn_to_step)
        self.layer_layout.addLayout(btn_layout)
        
        lbl_global = QLabel("GLOBAL-EBENEN")
        lbl_global.setStyleSheet("font-size: 10px; font-weight: bold; color: #4ec9b0; margin-top: 10px;")
        self.layer_layout.addWidget(lbl_global)
        
        self.global_layer_list = LayerListWidget(is_global_list=True, main_window=self)
        self.global_layer_list.itemClicked.connect(self.on_layer_clicked)
        self.layer_layout.addWidget(self.global_layer_list)
        
        self.dock_layers.setWidget(self.layer_panel)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.dock_layers)

        # Right Dock: Properties
        self.dock_props = QDockWidget("EIGENSCHAFTEN", self)
        self.props_widget = QWidget()
        self.props_layout = QVBoxLayout(self.props_widget)
        self.props_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.dock_props.setWidget(self.props_widget)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.dock_props)
        self.tabifyDockWidget(self.dock_layers, self.dock_props)
        
        self.scene.selectionChanged.connect(self.update_properties)
        self.update_thumbnails()

    def on_steps_reordered(self, parent, start, end, destination, row):
        """Synchronize self.steps with QListWidget reordering"""
        # Note: QListWidget reorders internally, we need to rebuild our self.steps list 
        # based on the new order of data in the items (we store the original step index or object)
        new_steps = []
        for i in range(self.thumb_list.count()):
            item = self.thumb_list.item(i)
            # Find the step object associated with this item
            step_obj = item.data(Qt.ItemDataRole.UserRole)
            if step_obj:
                new_steps.append(step_obj)
        
        if new_steps:
            self.steps = new_steps
            # Refresh labels/thumbnails to update step numbers (#1, #2, etc)
            self.update_thumbnails()
            
    def update_thumbnails(self):
        """Update thumbnail list with step numbers"""
        # Block signals during rebuild to prevent on_step_changed firing
        self.thumb_list.blockSignals(True)
        self.thumb_list.clear()
        
        for i, step in enumerate(self.steps):
            # Create thumbnail with step number overlay
            img = step.raw_img
            h, w = img.shape[:2]
            
            # Resize for thumbnail
            thumb_h = 80
            thumb_w = int(w * thumb_h / h)
            thumb = cv2.resize(img, (thumb_w, thumb_h))
            
            # Add step number overlay
            cv2.putText(thumb, f"#{i+1}", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 
                       1.0, (0, 0, 0), 6, cv2.LINE_AA)
            cv2.putText(thumb, f"#{i+1}", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 
                       1.0, (255, 255, 255), 2, cv2.LINE_AA)
            
            # Convert to QPixmap
            rgb = cv2.cvtColor(thumb, cv2.COLOR_BGR2RGB)
            qimg = QImage(rgb.data, thumb_w, thumb_h, thumb_w*3, QImage.Format.Format_RGB888)
            pix = QPixmap.fromImage(qimg)
            
            item = QListWidgetItem(f"Schritt {i+1}")
            item.setIcon(QIcon(pix))
            item.setData(Qt.ItemDataRole.UserRole, step) # Store the step object for reordering sync
            self.thumb_list.addItem(item)
        
        self.thumb_list.blockSignals(False)
        self.thumb_list.setCurrentRow(self.current_idx)

    def on_step_changed(self, idx):
        """Handle step change from thumbnail list"""
        if idx < 0 or idx >= len(self.steps):
            return
        # Save current state before switching
        self.save_current_state()
        # Load new step
        self.load_step(idx)

    def delete_step(self):
        """Delete current step"""
        if len(self.steps) <= 1:
            QMessageBox.warning(self, "Fehler", "Mindestens ein Schritt muss vorhanden sein!")
            return
        
        reply = QMessageBox.question(self, 'Löschen', 
                                    f'Schritt {self.current_idx + 1} wirklich löschen?',
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            self.push_undo() # Save state before delete
            del self.steps[self.current_idx]
            
            # Adjust current index
            if self.current_idx >= len(self.steps):
                self.current_idx = len(self.steps) - 1
            
            self.update_thumbnails()
            self.load_step(self.current_idx)

    def on_export_clicked(self):
        self.save_current_state()
        self.save_cb(self.steps, self.global_layers, self.global_crop)

    def save_project(self, save_as=False):
        """Save project to disk"""
        self.save_current_state()
        
        name = self.current_project_name
        
        # Check if we need to ask for name
        if save_as is True or not name:
            prefill = name if name else f"Project_{datetime.now().strftime('%Y%m%d_%H%M')}"
            new_name, ok = QInputDialog.getText(self, "Projekt speichern", "Projektname:", text=prefill)
            if not ok or not new_name: return
            name = new_name
            self.current_project_name = name
            self.setWindowTitle(f"ClickStep Guide - {name}")
        
        base_path = os.path.join(os.getcwd(), "projects", name)
        img_path = os.path.join(base_path, "images")
        if os.path.exists(base_path) and (save_as is True):
             # Maybe warn? But standard Save As just overwrites/uses that name
             pass
             
        os.makedirs(img_path, exist_ok=True)
        
        data = {
            "global_crop": self.global_crop,
            "global_layers": [{"type": l.type, "data": l.data, "label": l.label, "uid": l.uid} for l in self.global_layers],
            "steps": []
        }
        
        for i, s in enumerate(self.steps):
            filename = f"step_{i}.png"
            cv2.imwrite(os.path.join(img_path, filename), s.raw_img)
            
            data["steps"].append({
                "image": filename,
                "description": s.description,
                "layers": [{"type": l.type, "data": l.data, "label": l.label} for l in s.layers]
            })
            
        with open(os.path.join(base_path, "project.json"), "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
        
        QMessageBox.information(self, "Erfolg", f"Projekt '{name}' gespeichert!")
        self.save_cb(None, None, None) # Trigger project list update in main window

    def set_tool(self, tool):
        self.draw_mode = tool
        self.update_tool_buttons()

    def update_tool_buttons(self):
        self.btn_select.setChecked(self.draw_mode is None)
        self.btn_blur.setChecked(self.draw_mode == 'blur')
        self.btn_zoom.setChecked(self.draw_mode == 'zoom')
        self.btn_text.setChecked(self.draw_mode == 'text')
        self.btn_global_blur.setChecked(self.draw_mode == 'global_blur')
        self.btn_crop.setChecked(self.draw_mode == 'crop')

    def delete_selected(self):
        """Delete selected layer items"""
        for item in self.scene.selectedItems():
            if hasattr(item, 'item_type') and item.item_type == 'click': 
                continue
            self.scene.removeItem(item)
        self.refresh_layer_list()

    def on_step_changed(self, idx):
        if idx < 0: return
        self.save_current_state()
        self.load_step(idx)

    def handle_layer_drop(self, target_is_global):
        """Called when an item is dropped into one of the layer lists"""
        # Find all items in the target list
        target_list = self.global_layer_list if target_is_global else self.step_layer_list
        
        for i in range(target_list.count()):
            item = target_list.item(i)
            uid = item.data(Qt.ItemDataRole.UserRole)
            if not uid: continue
            
            # Find in scene
            scene_item = next((x for x in self.scene.items() if getattr(x, 'uid', None) == uid), None)
            if scene_item:
                # Update status based on target list
                scene_item.is_global = target_is_global
                
                # Visual update
                if target_is_global and hasattr(scene_item, 'base_color'):
                     scene_item.base_color = QColor(0, 255, 0, 80) # Global color
                elif not target_is_global and hasattr(scene_item, 'base_color') and scene_item.item_type == 'blur':
                     scene_item.base_color = QColor(0, 0, 0, 150) # Local color
                
                scene_item.update()

        # Update lists to reflect changes cleanly
        self.refresh_layer_list()

    def on_layer_clicked(self, item):
        # Sync selection from list to scene
        uid = item.data(Qt.ItemDataRole.UserRole)
        if not uid: return
        
        # Find item by UUID
        scene_item = next((x for x in self.scene.items() if getattr(x, 'uid', None) == uid), None)
        
        if scene_item:
            self.scene.clearSelection()
            scene_item.setSelected(True)

    def move_layer_to_global(self):
        """Move selected step layer to global"""
        item = self.step_layer_list.currentItem()
        if not item:
            return
        
        idx = self.step_layer_list.row(item)
        scene_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and i.item_type != 'click' and not getattr(i, 'is_global', False)]
        scene_items.sort(key=lambda x: x.zValue())
        
        if idx < len(scene_items):
            scene_item = scene_items[idx]
            scene_item.is_global = True
            
            # Update visual indicator for blur items
            if hasattr(scene_item, 'base_color'):
                scene_item.base_color = QColor(0, 255, 0, 80)
                scene_item.update()
            
            self.refresh_layer_list()
            QMessageBox.information(self, "Erfolg", "Ebene wurde zu Global verschoben!")

    def move_layer_to_step(self):
        """Move selected global layer to step"""
        item = self.global_layer_list.currentItem()
        if not item:
            return
        
        idx = self.global_layer_list.row(item)
        scene_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and getattr(i, 'is_global', False)]
        scene_items.sort(key=lambda x: x.zValue())
        
        if idx < len(scene_items):
            scene_item = scene_items[idx]
            scene_item.is_global = False
            
            # Update visual indicator for blur items
            if hasattr(scene_item, 'base_color') and scene_item.item_type == 'blur':
                scene_item.base_color = QColor(0, 0, 0, 150)
                scene_item.update()
            
            self.refresh_layer_list()
            QMessageBox.information(self, "Erfolg", "Ebene wurde zu Bild verschoben!")

    def on_layer_moved_to_global(self):
        """Handle drag-drop between lists"""
        # This would need more complex implementation
        pass

    def refresh_layer_list(self):
        self.step_layer_list.clear()
        self.global_layer_list.clear()
        
        # Separate items by global status
        step_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and not getattr(i, 'is_global', False)]
        global_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and getattr(i, 'is_global', False)]
        
        step_items.sort(key=lambda x: x.zValue())
        global_items.sort(key=lambda x: x.zValue())
        
        # Populate step layers
        for i in step_items:
            name = f"{i.item_type.upper()}"
            if hasattr(i, 'toPlainText'):  # Text item
                name += f": {i.toPlainText()[:20]}"
            li = QListWidgetItem(name)
            li.setData(Qt.ItemDataRole.UserRole, getattr(i, 'uid', None))
            self.step_layer_list.addItem(li)
        
        # Populate global layers
        for i in global_items:
            name = f"🌍 {i.item_type.upper()}"
            if hasattr(i, 'toPlainText'):  # Text item
                name += f": {i.toPlainText()[:20]}"
            li = QListWidgetItem(name)
            li.setData(Qt.ItemDataRole.UserRole, getattr(i, 'uid', None))
            self.global_layer_list.addItem(li)

    def save_current_state(self):
        """Extract items from scene and sync back to data model"""
        if self.current_idx < 0 or not self.steps: return
        
        new_layers = []
        new_globals = []
        
        # Offset handling (add crop offset back to coordinates for storage)
        ox = getattr(self, 'current_offset_x', 0)
        oy = getattr(self, 'current_offset_y', 0)
        
        for item in self.scene.items():
            if not hasattr(item, 'item_type'): continue
            
            # Extract data based on type
            data = {}
            uid = getattr(item, 'uid', None)
            
            if item.item_type == 'blur':
                r = item.sceneBoundingRect()
                # Reverse offset (add crop start position)
                data = {'coords': (int(r.left()+ox), int(r.top()+oy), int(r.right()+ox), int(r.bottom()+oy))}
                layer = Layer('blur', data, item.label, getattr(item, 'is_global', False), uid)
                if getattr(item, 'is_global', False): new_globals.append(layer)
                else: new_layers.append(layer)
                
            elif item.item_type == 'zoom':
                r = item.box_rect.translated(item.pos())
                t = item.target
                data = {'x': int(r.left()+ox), 'y': int(r.top()+oy), 'size': int(r.width()), 
                        'target_x': int(t.x()+ox), 'target_y': int(t.y()+oy)}
                layer = Layer('zoom', data, "Zoom", getattr(item, 'is_global', False), uid)
                if getattr(item, 'is_global', False): new_globals.append(layer)
                else: new_layers.append(layer)
                
            elif item.item_type == 'text':
                p = item.pos()
                f = item.font()
                data = {
                    'text': item.toPlainText(), 
                    'x': int(p.x()+ox), 
                    'y': int(p.y()+oy), 
                    'color': (item.defaultTextColor().blue(), item.defaultTextColor().green(), item.defaultTextColor().red()),
                    'font': {
                        'family': f.family(),
                        'size': f.pointSize(),
                        'bold': f.bold(),
                        'italic': f.italic(),
                        'underline': f.underline()
                    }
                }
                layer = Layer('text', data, "Text", getattr(item, 'is_global', False), uid)
                if getattr(item, 'is_global', False): new_globals.append(layer)
                else: new_layers.append(layer)
                
            elif item.item_type == 'click':
                # Preserve click from current step but update position if moved
                p = item.rect().center() + item.pos()
                self.steps[self.current_idx].x = int(p.x()+ox)
                self.steps[self.current_idx].y = int(p.y()+oy)
                # Update the click layer separately or just keep it
        
        self.steps[self.current_idx].layers = [l for l in self.steps[self.current_idx].layers if l.type == 'click']
        for l in self.steps[self.current_idx].layers:
            if l.type == 'click':
                l.data['x'] = self.steps[self.current_idx].x
                l.data['y'] = self.steps[self.current_idx].y
        self.steps[self.current_idx].layers.extend(new_layers)
        
        if new_globals:
            # Smart update using UUIDs
            existing_map = {l.uid: l for l in self.global_layers}
            
            for ng in new_globals:
                if ng.uid in existing_map:
                    # Check if actually modified to avoid unnecessary updates? 
                    # For now just update data and label
                    existing_map[ng.uid].data = ng.data
                    existing_map[ng.uid].label = ng.label
                else:
                    self.global_layers.append(ng)
                    existing_map[ng.uid] = ng

    # ==================== UNDO SYSTEM ====================
    def push_undo(self):
        """Save current state to undo stack"""
        # First sync current scene to model to catch latest changes
        self.save_current_state()
        snapshot = self.capture_snapshot()
        self.undo_stack.append(snapshot)
        # Limit stack size
        if len(self.undo_stack) > 20:
            self.undo_stack.pop(0)
            
    def undo(self):
        """Restore last state"""
        if not self.undo_stack:
            return
            
        snapshot = self.undo_stack.pop()
        self.restore_snapshot(snapshot)
        QMessageBox.information(self, "Undo", "Schritt zurückgesetzt.")

    def capture_snapshot(self):
        """Create a deep copy of the project metadata"""
        # Copy global layers
        globals_copy = []
        for l in self.global_layers:
            import copy
            globals_copy.append(Layer(l.type, copy.deepcopy(l.data), l.label, l.is_global, l.uid))
            
        # Copy steps structure (images are referenced, not copied)
        steps_copy = []
        for s in self.steps:
            # We assume Step class structure. 
            # Step(raw_img, x, y, label)
            import copy
            # We need to recreate the Step object to detach layer list, but keep image ref
            # Find Step class in global scope or imported
            from pro_recorder import Step 
            
            new_step = Step(s.raw_img, s.x, s.y, getattr(s, 'label', ""))
            
            # Manually copy layers
            new_step.layers = []
            for l in s.layers:
                new_step.layers.append(Layer(l.type, copy.deepcopy(l.data), l.label, l.is_global, l.uid))
            
            steps_copy.append(new_step)
            
        return {
            "global_layers": globals_copy,
            "steps": steps_copy,
            "global_crop": self.global_crop,
            "current_idx": self.current_idx
        }

    def restore_snapshot(self, snapshot):
        """Restore project from snapshot"""
        self.global_layers = snapshot["global_layers"]
        self.steps = snapshot["steps"]
        self.global_crop = snapshot["global_crop"]
        idx = snapshot["current_idx"]
        
        if idx >= len(self.steps): idx = len(self.steps) - 1
        
        self.update_thumbnails()
        self.load_step(idx)

    def update_properties(self):
        try:
            # Check if scene is valid or deleted
            if not getattr(self, 'scene', None): return
            
            # Helper to clear layout recursively
            def clear_layout(layout):
                if layout is None: return
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget:
                        widget.deleteLater()
                    elif item.layout():
                        clear_layout(item.layout())
            
            clear_layout(self.props_layout)
            
            items = self.scene.selectedItems()
            if not items:
                # Add spacer to push content up if needed, or just label
                lbl = QLabel("Keine Auswahl")
                lbl.setStyleSheet("color: #666; font-style: italic;")
                self.props_layout.addWidget(lbl)
                self.props_layout.addStretch()
                return
            
            item = items[0]
            
            # Header
            header = QLabel(f"{type(item).__name__}")
            header.setStyleSheet("font-weight: bold; color: #0078d4; font-size: 14px; margin-bottom: 10px;")
            self.props_layout.addWidget(header)
            
            if isinstance(item, EditableTextItem):
                # TEXT FORMATTING CONTROLS
                
                # 1. Color Button
                btn_color = QPushButton("🎨 Farbe")
                btn_color.clicked.connect(lambda: self.change_text_color(item))
                self.props_layout.addWidget(btn_color)
                
                self.props_layout.addSpacing(10)
                self.props_layout.addWidget(QLabel("Schriftart:"))
                
                # 2. Font Family
                font_combo = QFontComboBox()
                font_combo.setCurrentFont(item.font())
                font_combo.currentFontChanged.connect(lambda f: item.update_font(family=f.family()))
                self.props_layout.addWidget(font_combo)
                
                # 3. Size & Styles Row
                row = QHBoxLayout()
                
                # Size
                spin_size = QSpinBox()
                spin_size.setRange(8, 200)
                sz = item.font().pointSize()
                spin_size.setValue(sz if sz > 0 else 18) # Default to 18 if invalid
                spin_size.valueChanged.connect(lambda s: item.update_font(size=s))
                row.addWidget(spin_size)
                
                # Bold
                btn_bold = QPushButton("B")
                btn_bold.setCheckable(True)
                btn_bold.setFixedWidth(30)
                btn_bold.setChecked(item.font().bold())
                btn_bold.setStyleSheet("font-weight: bold;")
                btn_bold.clicked.connect(lambda c: item.update_font(bold=c))
                row.addWidget(btn_bold)
                
                # Italic
                btn_italic = QPushButton("I")
                btn_italic.setCheckable(True)
                btn_italic.setFixedWidth(30)
                btn_italic.setChecked(item.font().italic())
                btn_italic.setStyleSheet("font-style: italic;")
                btn_italic.clicked.connect(lambda c: item.update_font(italic=c))
                row.addWidget(btn_italic)
                
                # Underline
                btn_underline = QPushButton("U")
                btn_underline.setCheckable(True)
                btn_underline.setFixedWidth(30)
                btn_underline.setChecked(item.font().underline())
                btn_underline.setStyleSheet("text-decoration: underline;")
                btn_underline.clicked.connect(lambda c: item.update_font(underline=c))
                row.addWidget(btn_underline)
                
                self.props_layout.addLayout(row)
                
            # Add general info (Optional)
            if hasattr(item, 'is_global'):
                status = "🌍 Global" if item.is_global else "📄 Lokal"
                self.props_layout.addWidget(QLabel(f"Status: {status}"))
                
        except RuntimeError:
            pass # Scene deleted during update

    def change_text_color(self, item):
        color = QColorDialog.getColor(item.defaultTextColor())
        if color.isValid():
            item.setDefaultTextColor(color)

    def update_thumbnails(self):
        self.thumb_list.clear()
        for i in range(len(self.steps)):
            self.thumb_list.addItem(f"Schritt {i+1}")

    def load_step(self, idx):
        if idx < 0 or idx >= len(self.steps): return
        self.current_idx = idx
        
        self.scene.clear()
        step = self.steps[idx]
        
        # 1. Background Image - Handle Global Crop
        img = step.raw_img
        
        # Calculate offsets based on crop
        offset_x, offset_y = 0, 0
        if self.global_crop:
            x1, y1, x2, y2 = self.global_crop
            h, w = img.shape[:2]
            # Clip crop rect to image bounds
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(w, x2), min(h, y2)
            
            if x2 > x1 and y2 > y1:
                img = img[y1:y2, x1:x2] # Crop the image logic
                offset_x, offset_y = x1, y1
                
        self.current_offset_x = offset_x
        self.current_offset_y = offset_y
        
        h, w = img.shape[:2]
        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        qimg = QImage(rgb.data, w, h, w*3, QImage.Format.Format_RGB888)
        self.current_pixmap = QPixmap.fromImage(qimg)
        
        bg_item = QGraphicsPixmapItem(self.current_pixmap)
        bg_item.setZValue(-100)
        self.scene.addItem(bg_item)
        
        # Helper helpers
        def shift_rect(c):
            # Transform global coordinates to cropped local coordinates
            return QRectF(c[0]-offset_x, c[1]-offset_y, c[2]-c[0], c[3]-c[1])
            
        def shift_point(x, y):
            return QPointF(x-offset_x, y-offset_y)
        
        # 2. Existing Layers
        
        # Global Layers (persist across steps)
        for gl in self.global_layers:
            # We need to render them relative to the crop
            # gl.data coordinates are always relative to the FULL ORIGINAL IMAGE
            
            if gl.type == 'blur':
                c = gl.data['coords']
                # Only add if visible in crop? QGraphicsView handles clipping anyway.
                item = BlurItem(shift_rect(c), True, self.current_pixmap, uid=gl.uid)
                self.scene.addItem(item)
            elif gl.type == 'zoom':
                d = gl.data
                rect = QRectF(d['x']-offset_x, d['y']-offset_y, d['size'], d['size'])
                target = shift_point(d['target_x'], d['target_y'])
                item = ZoomItem(rect, target, self.current_pixmap, is_global=True, uid=gl.uid)
                self.scene.addItem(item)
            elif gl.type == 'text':
                d = gl.data
                color = QColor(d['color'][2], d['color'][1], d['color'][0])
                item = EditableTextItem(d['text'], color, is_global=True, uid=gl.uid)
                item.setPos(d['x']-offset_x, d['y']-offset_y)
                if 'font' in d:
                    f = d['font']
                    item.update_font(f.get('family'), f.get('size'), f.get('bold'), f.get('italic'), f.get('underline'))
                self.scene.addItem(item)
                
        # Step Layers
        for l in step.layers:
            if l.type == 'click':
                item = ClickMarkerItem(l.data['x']-offset_x, l.data['y']-offset_y, str(idx+1))
                self.scene.addItem(item)
            elif l.type == 'blur':
                c = l.data['coords']
                item = BlurItem(shift_rect(c), False, self.current_pixmap)
                self.scene.addItem(item)
            elif l.type == 'zoom':
                d = l.data
                rect = QRectF(d['x']-offset_x, d['y']-offset_y, d['size'], d['size'])
                target = shift_point(d['target_x'], d['target_y'])
                item = ZoomItem(rect, target, self.current_pixmap, is_global=False)
                self.scene.addItem(item)
            elif l.type == 'text':
                d = l.data
                color = QColor(d['color'][2], d['color'][1], d['color'][0])
                item = EditableTextItem(d['text'], color, is_global=False)
                item.setPos(d['x']-offset_x, d['y']-offset_y)
                if 'font' in d:
                    f = d['font']
                    item.update_font(f.get('family'), f.get('size'), f.get('bold'), f.get('italic'), f.get('underline'))
                self.scene.addItem(item)
        
        # 3. Hardcoded Watermark (Non-Deletable) - Added once per step
        watermark = QGraphicsTextItem("Created with ClickStep Guide")
        watermark.setDefaultTextColor(QColor(150, 150, 150, 120))
        wm_font = QFont("Segoe UI", 10)
        watermark.setFont(wm_font)
        watermark.setZValue(9999) # Always on top
        # No flags = not selectable, not movable, not focusable
        watermark.setFlags(QGraphicsItem.GraphicsItemFlag(0))
        watermark.setAcceptedMouseButtons(Qt.MouseButton.NoButton) # Click through
        
        # Position bottom right of the image
        pw = self.current_pixmap.width()
        ph = self.current_pixmap.height()
        watermark.setPos(pw - 160, ph - 25)
        self.scene.addItem(watermark)
        
        self.scene.setSceneRect(0, 0, w, h)
        
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(10, lambda: self.view.fitInView(self.scene.sceneRect(), Qt.AspectRatioMode.KeepAspectRatio))
        
        self.refresh_layer_list()

# ==================== RECORDER (unchanged) ====================

class Step:
    def __init__(self, raw_img, x, y, label):
        self.raw_img = raw_img
        self.x, self.y = x, y
        self.description = label
        self.layers = [Layer('click', {'x': x, 'y': y}, label)]

class RecordingSignal(QObject):
    click_detected = pyqtSignal(int, int, str, object)

class RecordingThread(QThread):
    def __init__(self):
        super().__init__()
        self.signals = RecordingSignal()
        self.is_running = False

    def run(self):
        self.is_running = True
        self.mouse_listener = mouse.Listener(on_click=self.on_click)
        self.mouse_listener.start()
        while self.is_running:
            time.sleep(0.1)
        self.mouse_listener.stop()

    def on_click(self, x, y, button, pressed):
        if not pressed: return
        time.sleep(0.12)
        raw = ImageGrab.grab()
        img = cv2.cvtColor(np.array(raw), cv2.COLOR_RGB2BGR)
        self.signals.click_detected.emit(x, y, "Click", img)

class ProRecorder(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ClickStep Guide")
        self.resize(1280, 800)
        self.setStyleSheet("""
            QMainWindow { background-color: #121212; }
            QWidget { color: #e0e0e0; font-family: 'Segoe UI'; }
            QDockWidget::title { background-color: #1a1a1a; padding: 10px; font-weight: bold; }
            QListWidget { background-color: #1a1a1a; border: none; }
            QListWidget::item { padding: 10px; border-bottom: 1px solid #252525; }
            QListWidget::item:selected { background-color: #007acc; }
            QPushButton { background-color: #2a2a2a; border: none; padding: 10px 15px; border-radius: 4px; }
            QPushButton:hover { background-color: #3a3a3a; }
            QPushButton#RecordBtn { background-color: #007acc; font-size: 16px; padding: 20px 40px; }
            QPushButton#RecordBtn:hover { background-color: #008be2; }
        """)
        
        self.steps = []
        self.global_layers = []
        self.global_crop = None
        self.is_recording = False
        self.overlay = RecordingOverlay() # Create overlay
        
        self.recording_thread = RecordingThread()
        self.recording_thread.signals.click_detected.connect(self.handle_click)
        
        self.setup_ui()
        self.setup_hotkeys()
        self.update_project_list()

    def setup_hotkeys(self):
        """Setup global keyboard shortcuts"""
        h = {'<ctrl>+<alt>+s': lambda: self.btn_record.click(),
             '<ctrl>+<alt>+e': lambda: self.btn_record.click()}
        self.hotkey_thread = keyboard.GlobalHotKeys(h)
        self.hotkey_thread.start()

    def setup_ui(self):
        self.center_widget = QWidget()
        self.setCentralWidget(self.center_widget)
        
        # Main Layout: Horizontal (Left Sidebar | Right Content)
        main_layout = QHBoxLayout(self.center_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # --- LEFT SIDEBAR (Static Project List) ---
        sidebar = QWidget()
        sidebar.setFixedWidth(280)
        sidebar.setStyleSheet("background-color: #1a1a1a; border-right: 1px solid #333;")
        
        side_layout = QVBoxLayout(sidebar)
        side_layout.setContentsMargins(15, 20, 15, 20)
        side_layout.setSpacing(10)
        
        lbl_proj = QLabel("PROJEKTE")
        lbl_proj.setStyleSheet("color: #666; font-weight: bold; font-size: 12px; letter-spacing: 1px;")
        side_layout.addWidget(lbl_proj)
        
        self.proj_list = QListWidget()
        self.proj_list.setStyleSheet("""
            QListWidget { background: transparent; border: none; }
            QListWidget::item { padding: 10px; border-bottom: 1px solid #252525; }
            QListWidget::item:selected { background-color: #007acc; border-left: 3px solid white; }
            QListWidget::item:hover { background-color: #252526; }
        """)
        side_layout.addWidget(self.proj_list)
        
        # Sidebar Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(5)
        
        btn_open = QPushButton("📂 Öffnen")
        btn_open.setFixedHeight(40)
        btn_open.setStyleSheet("background-color: #333333; border: 1px solid #444;")
        btn_open.clicked.connect(self.load_project)
        
        btn_del = QPushButton("🗑️")
        btn_del.setFixedSize(40, 40)
        btn_del.setStyleSheet("background-color: #2b1111; color: #ff4444; border: 1px solid #500000;")
        btn_del.clicked.connect(self.delete_project)
        
        btn_layout.addWidget(btn_open)
        btn_layout.addWidget(btn_del)
        side_layout.addLayout(btn_layout)
        
        main_layout.addWidget(sidebar)
        
        # --- RIGHT CONTENT AREA ---
        content_area = QWidget()
        self.layout = QVBoxLayout(content_area)
        self.layout.setContentsMargins(50, 50, 50, 50)
        self.layout.setSpacing(30)
        
        if os.path.exists("logo.png"):
            logo_lbl = QLabel()
            pix = QPixmap("logo.png")
            if not pix.isNull():
                pix = pix.scaledToHeight(350, Qt.TransformationMode.SmoothTransformation) # Bigger Logo
                logo_lbl.setPixmap(pix)
                logo_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                self.layout.addWidget(logo_lbl)
        
        self.layout.addStretch() # Spacer
        
        self.btn_record = QPushButton("NEUE AUFNAHME STARTEN")
        self.btn_record.setObjectName("RecordBtn")
        self.btn_record.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_record.clicked.connect(self.toggle_recording)
        self.layout.addWidget(self.btn_record, 0, Qt.AlignmentFlag.AlignCenter)
        
        shortcuts = QLabel("Shortcuts: Ctrl+Alt+S (Start/Stop) | Ctrl+Alt+E (Editor)")
        shortcuts.setAlignment(Qt.AlignmentFlag.AlignCenter)
        shortcuts.setStyleSheet("color: #555; font-size: 13px;")
        self.layout.addWidget(shortcuts)
        
        self.layout.addStretch() # Bottom spacer
        
        # Attribution Footer
        footer = QLabel("© 2026 ClickStep Guide | High-Performance Documentation Engine")
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        footer.setStyleSheet("color: #333; font-size: 11px; margin-top: 20px; border-top: 1px solid #1a1a1a; padding-top: 10px;")
        self.layout.addWidget(footer)
        
        main_layout.addWidget(content_area)

    def toggle_recording(self):
        if not self.is_recording:
            self.is_recording = True
            self.steps = []
            self.btn_record.setText("AUFNAHME STOPPEN")
            self.showMinimized()
            
            # Ensure overlay is correctly positioned and visible
            screen = QApplication.primaryScreen().geometry()
            self.overlay.setGeometry(screen.width() - 170, 30, 150, 60)
            self.overlay.show() 
            self.recording_thread.start()
        else:
            self.stop_recording()

    def stop_recording(self):
        self.is_recording = False
        self.recording_thread.is_running = False
        self.overlay.hide() # Hide overlay
        self.btn_record.setText("NEUE AUFNAHME STARTEN")
        self.showNormal()
        if self.steps:
            self.open_editor()

    def handle_click(self, x, y, label, img):
        self.steps.append(Step(img, x, y, label))

    def open_editor(self):
        self.hide() # Hide main recorder
        self.editor = ProEditor(self.steps, self.global_layers, self.global_crop, self.final_export, parent_window=self)
        self.editor.show()

    def final_export(self, steps, global_layers, global_crop):
        if not steps: # This is a refresh signal for project list
            self.update_project_list()
            return
            
        path, _ = QFileDialog.getSaveFileName(self, "Export", "", "Word (*.docx)")
        if not path: return
        
        prog = QProgressBar()
        prog.setMaximum(len(steps))
        self.layout.addWidget(prog)
        
        try:
            doc = Document()
            doc.add_heading('ClickStep Guide - Anleitung', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, s in enumerate(steps):
                # 1. Prepare Base Image (Crop)
                raw = s.raw_img
                if global_crop:
                    x1, y1, x2, y2 = global_crop
                    raw = raw[y1:y2, x1:x2]
                
                canvas = raw.copy()
                ox, oy = (global_crop[0], global_crop[1]) if global_crop else (0, 0)
                
                # 2. Render all annotations (using the same logic as our editor's paint methods)
                # First Global Blur
                for gl in global_layers:
                    if gl.type == 'blur':
                        c = gl.data['coords']
                        self.render_blur_cv2(canvas, (c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy), True)
                
                # Then Step Layers
                for l in s.layers:
                    if l.type == 'blur':
                        c = l.data['coords']
                        self.render_blur_cv2(canvas, (c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy), False)
                    elif l.type == 'click':
                        self.render_click_cv2(canvas, l.data['x']-ox, l.data['y']-oy, i+1)
                    elif l.type == 'zoom':
                        self.render_zoom_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'text':
                        self.render_text_cv2(canvas, l.data, ox, oy)
                
                # 3. Render Permanent Watermark
                self.render_watermark_cv2(canvas)
                
                # 4. Add to Word
                tmp_file = f"export_tmp_{i}.png"
                cv2.imwrite(tmp_file, canvas)
                doc.add_heading(f"Schritt {i+1}", level=1)
                if s.description: doc.add_paragraph(s.description)
                doc.add_picture(tmp_file, width=Inches(6))
                os.remove(tmp_file)
                prog.setValue(i+1)

            # 4. Add Copyright & Info Footer at the end
            doc.add_page_break()
            doc.add_heading('Über dieses Dokument', level=1)
            p = doc.add_paragraph()
            p.add_run('Diese Anleitung wurde professionell erstellt mit ').italic = True
            p.add_run('ClickStep Guide').bold = True
            p.add_run('.')
            doc.add_paragraph(f"Erstellungsdatum: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
            doc.add_paragraph("© 2026 AutoGuide Automationsysteme")

            doc.save(path)
            QMessageBox.information(self, "Erfolg", f"Export nach {path} abgeschlossen!")
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Export fehlgeschlagen: {str(e)}")
        finally:
            self.layout.removeWidget(prog)
            prog.deleteLater()

    # OpenCV Rendering Helpers for Export
    def render_blur_cv2(self, img, coords, is_global):
        x1, y1, x2, y2 = coords
        h, w = img.shape[:2]
        x1, y1, x2, y2 = max(0, x1), max(0, y1), min(w, x2), min(h, y2)
        if x2 > x1 and y2 > y1:
            roi = img[y1:y2, x1:x2]
            img[y1:y2, x1:x2] = cv2.GaussianBlur(roi, (75, 75), 40)
            cv2.rectangle(img, (x1, y1), (x2, y2), (0, 255, 0) if is_global else (255, 255, 255), 2)

    def render_click_cv2(self, img, x, y, num):
        cv2.circle(img, (x, y), 38, (255, 255, 255), 6, cv2.LINE_AA)
        cv2.circle(img, (x, y), 32, (255, 0, 0), 4, cv2.LINE_AA)
        cv2.putText(img, str(num), (x-15, y+15), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (0,0,0), 6, cv2.LINE_AA)
        cv2.putText(img, str(num), (x-15, y+15), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (255,255,255), 3, cv2.LINE_AA)

    def render_zoom_cv2(self, img, data, ox, oy):
        zx, zy, sz = data['x']-ox, data['y']-oy, data['size']
        tx, ty = data['target_x']-ox, data['target_y']-oy
        src_sz = sz // 2
        
        # Crop area
        x1, y1 = max(0, tx - src_sz//2), max(0, ty - src_sz//2)
        x2, y2 = min(img.shape[1], tx + src_sz//2), min(img.shape[0], ty + src_sz//2)
        if x2 > x1 and y2 > y1:
            roi = img[y1:y2, x1:x2]
            zoomed = cv2.resize(roi, (sz, sz))
            img[zy:zy+sz, zx:zx+sz] = zoomed
            cv2.rectangle(img, (zx, zy), (zx+sz, zy+sz), (255, 255, 255), 3)
            cv2.arrowedLine(img, (zx+sz//2, zy+sz//2), (tx, ty), (255, 255, 255), 3)

    def render_text_cv2(self, img, data, ox, oy):
        x, y = data['x']-ox, data['y']-oy
        color = data.get('color', (255, 255, 255)) # BGR
        cv2.putText(img, data['text'], (x, y), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (0,0,0), 6, cv2.LINE_AA)
        cv2.putText(img, data['text'], (x, y), cv2.FONT_HERSHEY_SIMPLEX, 1.2, color, 2, cv2.LINE_AA)

    def render_watermark_cv2(self, img):
        """Draw a permanent watermark in the bottom right corner"""
        h, w = img.shape[:2]
        text = "Created with ClickStep Guide"
        font = cv2.FONT_HERSHEY_SIMPLEX
        scale = 0.6
        thickness = 1
        color = (180, 180, 180) # Light grey
        
        size = cv2.getTextSize(text, font, scale, thickness)[0]
        tx, ty = w - size[0] - 15, h - 15
        
        # Shadow for visibility on bright backgrounds
        cv2.putText(img, text, (tx+1, ty+1), font, scale, (20, 20, 20), thickness, cv2.LINE_AA)
        cv2.putText(img, text, (tx, ty), font, scale, color, thickness, cv2.LINE_AA)

    def update_project_list(self):
        """Update project list in sidebar"""
        self.proj_list.clear()
        path = os.path.join(os.getcwd(), "projects")
        if not os.path.exists(path): 
            os.makedirs(path)
        for d in os.listdir(path):
            if os.path.isdir(os.path.join(path, d)):
                item = QListWidgetItem(d)
                self.proj_list.addItem(item)

    def load_project(self):
        """Load selected project"""
        item = self.proj_list.currentItem()
        if not item: return
        
        project_path = os.path.join(os.getcwd(), "projects", item.text(), "project.json")
        if not os.path.exists(project_path):
            QMessageBox.warning(self, "Fehler", "Projekt-Datei nicht gefunden!")
            return
        
        try:
            with open(project_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Load steps
            self.steps = []
            img_path = os.path.join(os.getcwd(), "projects", item.text(), "images")
            for step_data in data.get("steps", []):
                img_file = os.path.join(img_path, step_data["image"])
                if os.path.exists(img_file):
                    img = cv2.imread(img_file)
                    step = Step(img, 0, 0, step_data.get("description", ""))
                    step.layers = [] # Reset default click
                    for l_data in step_data.get("layers", []):
                        step.layers.append(Layer(l_data['type'], l_data['data'], l_data.get('label', 'Layer')))
                    
                    # Update step x, y from the first click layer found
                    click_l = next((l for l in step.layers if l.type == 'click'), None)
                    if click_l:
                        step.x, step.y = click_l.data['x'], click_l.data['y']
                        
                    self.steps.append(step)
            
            self.global_crop = data.get("global_crop")
            self.global_layers = []
            for gl_data in data.get("global_layers", []):
                self.global_layers.append(Layer(gl_data['type'], gl_data['data'], gl_data.get('label', 'Global Layer'), True))
            
            if self.steps:
                self.open_editor()
            else:
                QMessageBox.warning(self, "Warnung", "Keine Schritte im Projekt gefunden!")
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Projekt konnte nicht geladen werden: {str(e)}")

    def delete_project(self):
        """Delete selected project"""
        item = self.proj_list.currentItem()
        if not item: return
        
        reply = QMessageBox.question(self, 'Löschen', 
                                    f"Projekt '{item.text()}' wirklich löschen?", 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                shutil.rmtree(os.path.join(os.getcwd(), "projects", item.text()))
                self.update_project_list()
                QMessageBox.information(self, "Erfolg", "Projekt gelöscht!")
            except Exception as e:
                QMessageBox.critical(self, "Fehler", f"Löschen fehlgeschlagen: {str(e)}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ProRecorder()
    window.show()
    sys.exit(app.exec())
