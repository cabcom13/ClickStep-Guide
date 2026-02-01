import sys
import os
import time
import json
import shutil
import math
import uuid
import queue
import ctypes
from datetime import datetime

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QLabel, QFileDialog, QListWidget, QListWidgetItem, 
                             QDockWidget, QToolBar, QSlider, QInputDialog, QMessageBox, 
                             QGraphicsScene, QGraphicsView, QGraphicsItem, QGraphicsRectItem,
                             QGraphicsPixmapItem, QGraphicsTextItem, QGraphicsLineItem, QMenu,
                             QSpinBox, QColorDialog, QFontComboBox, QComboBox, QDialog, QLineEdit, 
                             QDialogButtonBox, QAbstractItemView, QCheckBox, QTextEdit, QFrame,
                             QFormLayout, QGroupBox, QRadioButton, QButtonGroup, QProgressBar)
from PyQt6.QtCore import Qt, QTimer, QPointF, QRectF, QRect, QSize, pyqtSignal, QObject, QLineF, QThread
from PyQt6.QtGui import (QPixmap, QPainter, QPen, QColor, QFont, QAction, QIcon, 
                         QBrush, QImage, QPainterPath)

from pynput import mouse, keyboard
from PIL import ImageGrab
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==================== DATA MODELS ====================

# ==================== DATA MODELS ====================

class ModernDialog(QDialog):
    """Custom unified dialog for inputs and confirmations to replace native popups"""
    def __init__(self, title, message, mode="input", default_text="", parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Dialog)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        
        layout = QVBoxLayout()
        self.setLayout(layout)
        
        # Stylized Background
        self.bg = QWidget()
        self.bg.setStyleSheet("""
            QWidget {
                background-color: #2b2b2b;
                border: 1px solid #454545;
                border-radius: 8px;
            }
            QLabel { color: #e0e0e0; font-family: "Segoe UI"; font-size: 14px; border: none; }
            QLineEdit, QTextEdit { 
                background: #333; color: white; border: 1px solid #555; 
                padding: 6px; border-radius: 4px; font-family: "Segoe UI";
            }
            QPushButton {
                background: #0078d4; color: white; border: none; padding: 6px 15px; border-radius: 4px;
                font-family: "Segoe UI"; font-weight: bold;
            }
            QPushButton:hover { background: #106ebe; }
            QPushButton#CancelBtn { background: #d13438; }
            QPushButton#CancelBtn:hover { background: #a80000; }
        """)
        bg_layout = QVBoxLayout(self.bg)
        layout.addWidget(self.bg)
        
        # Title
        lbl_title = QLabel(title)
        lbl_title.setStyleSheet("font-weight: bold; font-size: 16px; margin-bottom: 5px;")
        bg_layout.addWidget(lbl_title)
        
        # Message
        lbl_msg = QLabel(message)
        lbl_msg.setWordWrap(True)
        bg_layout.addWidget(lbl_msg)
        
        self.input_field = None
        
        if mode == "input" or mode == "multiline":
            if mode == "multiline":
                from PyQt6.QtWidgets import QTextEdit
                self.input_field = QTextEdit()
                self.input_field.setPlainText(default_text)
                self.input_field.setMinimumHeight(100)
            else:
                self.input_field = QLineEdit()
                self.input_field.setText(default_text)
            bg_layout.addWidget(self.input_field)
            
        elif mode == "list":
             self.input_field = QComboBox()
             self.input_field.addItems(default_text) # default_text is list here
             bg_layout.addWidget(self.input_field)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        btn_ok = QPushButton("OK" if mode != "confirm" else "Ja")
        btn_ok.clicked.connect(self.accept)
        
        btn_cancel = QPushButton("Abbrechen" if mode != "confirm" else "Nein")
        btn_cancel.setObjectName("CancelBtn")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_cancel)
        btn_layout.addWidget(btn_ok)
        bg_layout.addLayout(btn_layout)

    def get_text(self):
        if hasattr(self.input_field, 'toPlainText'):
            return self.input_field.toPlainText()
        if hasattr(self.input_field, 'currentText'):
            return self.input_field.currentText()
        return self.input_field.text() if self.input_field else ""

class RecordingOverlay(QWidget):
    """Modern pill-shaped recording status overlay"""
    def __init__(self):
        super().__init__()
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint | Qt.WindowType.Tool)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents)
        
        # Main Layout
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Container Frame (The Pill)
        self.container = QFrame()
        self.container.setObjectName("Pill")
        self.container.setStyleSheet("""
            QFrame#Pill {
                background-color: rgba(25, 25, 25, 240);
                border: 1px solid #444;
                border-radius: 20px;
            }
        """)
        
        # Content Layout
        inner_layout = QHBoxLayout(self.container)
        inner_layout.setContentsMargins(15, 8, 15, 8)
        inner_layout.setSpacing(12)
        
        # Red Dot (Pulsing ideally, but static for now)
        self.dot_lbl = QLabel("●")
        self.dot_lbl.setStyleSheet("color: #ff3333; font-size: 14px; font-family: 'Segoe UI Symbol';")
        
        # REC Text
        self.rec_lbl = QLabel("REC")
        self.rec_lbl.setStyleSheet("color: white; font-weight: 700; font-family: 'Segoe UI'; font-size: 13px; letter-spacing: 0.5px;")
        
        # Separator
        sep = QFrame()
        sep.setFixedWidth(1)
        sep.setFixedHeight(14)
        sep.setStyleSheet("background-color: #555;")
        
        # Steps
        self.step_lbl = QLabel("Steps: 0")
        self.step_lbl.setStyleSheet("color: #dddddd; font-family: 'Segoe UI'; font-size: 13px; font-weight: 500;")
        
        inner_layout.addWidget(self.dot_lbl)
        inner_layout.addWidget(self.rec_lbl)
        inner_layout.addWidget(sep)
        inner_layout.addWidget(self.step_lbl)
        
        layout.addWidget(self.container)
        
        # Initial Position (Top Center-Right)
        screen = QApplication.primaryScreen().geometry()
        pill_w = 190
        pill_h = 56 # Height implicitly handled by layout, but widget height needed
        self.setGeometry(screen.width() - pill_w - 40, 40, pill_w, pill_h)
        
        # Exclude from capture (Win 10 2004+)
        if os.name == 'nt':
            try:
                import ctypes
                user32 = ctypes.windll.user32
                # WDA_EXCLUDEFROMCAPTURE = 0x00000011
                user32.SetWindowDisplayAffinity(int(self.winId()), 0x00000011) 
            except:
                pass

    def update_steps(self, count):
        self.step_lbl.setText(f"Steps: {count}")

class Layer:
    def __init__(self, ltype, data, label, is_global=False, uid=None):
        self.type = ltype
        self.data = data
        self.label = label
        self.is_global = is_global
        self.visible = True
        self.uid = uid if uid else str(uuid.uuid4())

# Global Click Marker Settings
class ClickMarkerSettings:
    """Singleton class for global click marker appearance settings"""
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        if self._initialized:
            return
        self._initialized = True
        
        # Default settings
        self.color = QColor(0, 168, 255)  # Neon Blue
        self.text_color = QColor(255, 255, 255) # White text default
        self.size = 40  # Radius in pixels
        self.border_width = 3
        self.show_glow = True
        self.number_size = 16
        
        # Load from file if exists
        self.load()
    
    def save(self):
        """Save settings to file"""
        settings = {
            'color': [self.color.red(), self.color.green(), self.color.blue(), self.color.alpha()],
            'text_color': [self.text_color.red(), self.text_color.green(), self.text_color.blue()],
            'size': self.size,
            'border_width': self.border_width,
            'show_glow': self.show_glow,
            'number_size': self.number_size
        }
        try:
            os.makedirs('.settings', exist_ok=True)
            with open('.settings/marker_settings.json', 'w') as f:
                json.dump(settings, f)
        except Exception as e:
            print(f"Failed to save settings: {e}")
            
    def load(self):
        """Load settings from file"""
        try:
            if os.path.exists('.settings/marker_settings.json'):
                with open('.settings/marker_settings.json', 'r') as f:
                    data = json.load(f)
                    
                    if 'color' in data:
                        c = data['color']
                        # Handle old format (RGB) vs new (RGBA)
                        if len(c) == 4:
                            self.color = QColor(c[0], c[1], c[2], c[3])
                        else:
                            self.color = QColor(c[0], c[1], c[2])
                    
                    if 'text_color' in data:
                        c = data['text_color']
                        self.text_color = QColor(c[0], c[1], c[2])
                        
                    self.size = data.get('size', 40)
                    self.border_width = data.get('border_width', 3)
                    self.show_glow = data.get('show_glow', True)
                    self.number_size = data.get('number_size', 16)
        except Exception as e:
            print(f"Failed to load settings: {e}")
            try:
                with open('.settings/marker_settings.json', 'w') as f:
                    json.dump({}, f)
            except: pass

class AppSettings:
    """Handles global application settings like shortcuts and theme"""
    def __init__(self):
        self.path = os.path.join(os.environ.get('LOCALAPPDATA', os.path.expanduser("~")), "ClickStepGuide", "app_settings.json")
        self.theme = "dark" # "dark" or "light"
        self.shortcut_record = "<ctrl>+<alt>+s"
        self.shortcut_editor = "<ctrl>+<alt>+e"
        self.load()

    def load(self):
        if os.path.exists(self.path):
            try:
                with open(self.path, 'r') as f:
                    data = json.load(f)
                    self.theme = data.get("theme", "dark")
                    self.shortcut_record = data.get("shortcut_record", "<ctrl>+<alt>+s")
                    self.shortcut_editor = data.get("shortcut_editor", "<ctrl>+<alt>+e")
            except: pass

    def save(self):
        os.makedirs(os.path.dirname(self.path), exist_ok=True)
        try:
            with open(self.path, 'w') as f:
                json.dump({
                    "theme": self.theme, 
                    "shortcut_record": self.shortcut_record,
                    "shortcut_editor": self.shortcut_editor
                }, f)
        except: pass

class SettingsDialog(QDialog):
    """Modern Settings Dialog to configure shortcuts and appearance"""
    def __init__(self, settings, parent=None):
        super().__init__(parent)
        self.settings = settings
        self.setWindowTitle("App Einstellungen")
        self.setFixedWidth(450)
        self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowStaysOnTopHint)
        
        # Style based on current theme if parent provided
        if parent:
            self.setStyleSheet(parent.styleSheet())
        
        layout = QVBoxLayout(self)
        
        # --- SHORTCUTS GROUP ---
        shortcut_group = QGroupBox("Tastenkombinationen (Hotkeys)")
        shortcut_layout = QFormLayout(shortcut_group)
        shortcut_layout.setContentsMargins(15, 20, 15, 15)
        shortcut_layout.setSpacing(15)
        
        self.edit_record = QLineEdit(self.settings.shortcut_record)
        self.edit_record.setPlaceholderText("<ctrl>+<alt>+s")
        
        self.edit_editor = QLineEdit(self.settings.shortcut_editor)
        self.edit_editor.setPlaceholderText("<ctrl>+<alt>+e")
        
        shortcut_layout.addRow("Aufnahme Start/Stop:", self.edit_record)
        shortcut_layout.addRow("Editor öffnen:", self.edit_editor)
        
        info_lbl = QLabel("<small>Format: &lt;ctrl&gt;+&lt;alt&gt;+Taste (z.B. &lt;ctrl&gt;+&lt;alt&gt;+s)</small>")
        info_lbl.setStyleSheet("color: #888; margin-top: -5px;")
        shortcut_layout.addRow("", info_lbl)
        
        layout.addWidget(shortcut_group)
        
        # --- THEME GROUP ---
        theme_group = QGroupBox("Erscheinungsbild")
        theme_layout = QVBoxLayout(theme_group)
        theme_layout.setContentsMargins(15, 15, 15, 15)
        
        self.radio_dark = QRadioButton("Dark Mode (Empfohlen)")
        self.radio_light = QRadioButton("Light Mode")
        
        if self.settings.theme == "dark": self.radio_dark.setChecked(True)
        else: self.radio_light.setChecked(True)
        
        theme_layout.addWidget(self.radio_dark)
        theme_layout.addWidget(self.radio_light)
        layout.addWidget(theme_group)
        
        layout.addStretch()
        
        # --- BUTTONS ---
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Einstellungen Speichern")
        btn_save.setStyleSheet("""
            QPushButton { 
                background: #007acc; color: white; padding: 10px; 
                border-radius: 4px; font-weight: bold; 
            }
            QPushButton:hover { background: #008be2; }
        """)
        btn_save.clicked.connect(self.accept)
        
        btn_cancel = QPushButton("Abbrechen")
        btn_cancel.setStyleSheet("""
            QPushButton { 
                background: transparent; border: 1px solid #444; 
                color: #ccc; padding: 10px; border-radius: 4px; 
            }
            QPushButton:hover { background: #333; }
        """)
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_cancel)
        btn_layout.addWidget(btn_save)
        layout.addLayout(btn_layout)

    def get_settings(self):
        return {
            "theme": "dark" if self.radio_dark.isChecked() else "light",
            "shortcut_record": self.edit_record.text().lower(),
            "shortcut_editor": self.edit_editor.text().lower()
        }

class LayerListWidget(QListWidget):
    """Custom ListWidget to handle layer drag and drop"""
    def __init__(self, parent=None, is_global_list=False, main_window=None):
        super().__init__(parent)
        self.is_global_list = is_global_list
        self.main_window = main_window
        self.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.setDragDropMode(QListWidget.DragDropMode.DragDrop)
        self.setDefaultDropAction(Qt.DropAction.MoveAction)
        
    def dropEvent(self, event):
        super().dropEvent(event)
        if self.main_window:
            self.main_window.handle_layer_drop(self.is_global_list)

# ==================== INTERACTIVE GRAPHICS ITEMS ====================

class ResizableRectItem(QGraphicsRectItem):
    """Photoshop-style resizable rectangle with corner handles"""
    def __init__(self, rect, color=QColor(255, 255, 255, 100), label=""):
        super().__init__(rect)
        self.item_type = 'rect' # Added item_type
        self.label = label
        self.handle_size = 10
        self.selected_handle = None
        self.base_color = color
        
        self.setPen(QPen(QColor(255, 255, 255), 2, Qt.PenStyle.DashLine))
        self.setBrush(QBrush(color))
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setAcceptHoverEvents(True)
        self.setCursor(Qt.CursorShape.SizeAllCursor)

    def get_handles(self):
        """Return positions of 8 resize handles"""
        r = self.rect()
        return {
            'tl': QPointF(r.left(), r.top()),
            'tr': QPointF(r.right(), r.top()),
            'bl': QPointF(r.left(), r.bottom()),
            'br': QPointF(r.right(), r.bottom()),
            't': QPointF(r.center().x(), r.top()),
            'b': QPointF(r.center().x(), r.bottom()),
            'l': QPointF(r.left(), r.center().y()),
            'r': QPointF(r.right(), r.center().y())
        }

    def get_handle_at(self, pos):
        """Check if position is over a handle"""
        handles = self.get_handles()
        for name, handle_pos in handles.items():
            if (pos - handle_pos).manhattanLength() < self.handle_size:
                return name
        return None

    def hoverMoveEvent(self, event):
        handle = self.get_handle_at(event.pos())
        if handle:
            cursors = {
                'tl': Qt.CursorShape.SizeFDiagCursor, 'br': Qt.CursorShape.SizeFDiagCursor,
                'tr': Qt.CursorShape.SizeBDiagCursor, 'bl': Qt.CursorShape.SizeBDiagCursor,
                't': Qt.CursorShape.SizeVerCursor, 'b': Qt.CursorShape.SizeVerCursor,
                'l': Qt.CursorShape.SizeHorCursor, 'r': Qt.CursorShape.SizeHorCursor
            }
            self.setCursor(cursors.get(handle, Qt.CursorShape.SizeAllCursor))
        else:
            self.setCursor(Qt.CursorShape.SizeAllCursor)
        super().hoverMoveEvent(event)

    def mousePressEvent(self, event):
        self.selected_handle = self.get_handle_at(event.pos())
        if not self.selected_handle:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.selected_handle:
            r = self.rect()
            pos = event.pos()
            
            if 'l' in self.selected_handle:
                r.setLeft(pos.x())
            if 'r' in self.selected_handle:
                r.setRight(pos.x())
            if 't' in self.selected_handle:
                r.setTop(pos.y())
            if 'b' in self.selected_handle:
                r.setBottom(pos.y())
            
            # Ensure minimum size
            if r.width() < 20: r.setWidth(20)
            if r.height() < 20: r.setHeight(20)
            
            self.setRect(r.normalized())
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self.selected_handle = None
        super().mouseReleaseEvent(event)

    def paint(self, painter, option, widget):
        super().paint(painter, option, widget)
        
        # Draw handles when selected
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

class BlurItem(ResizableRectItem):
    """Blur annotation with Photoshop-style interaction"""
    def __init__(self, rect, is_global=False, bg_pixmap=None, uid=None):
        # No base color (fully transparent), only border logic remains in paint
        super().__init__(rect, QColor(0, 0, 0, 0), "Globaler Blur" if is_global else "Zensur")
        self.is_global = is_global
        self.item_type = 'blur'
        self.bg_pixmap = bg_pixmap
        self.blurred_cache = None
        self.blur_strength = 40 # Default blur radius
        self.uid = uid if uid else str(uuid.uuid4())
        # Disable caching to refresh blur dynamically
        self.setCacheMode(QGraphicsItem.CacheMode.NoCache)
    
    def itemChange(self, change, value):
        # Trigger repaint when position changes
        if change == QGraphicsItem.GraphicsItemChange.ItemPositionChange:
            self.prepareGeometryChange()
        return super().itemChange(change, value)
    
    def paint(self, painter, option, widget):
        # Render actual blurred content
        if self.bg_pixmap and self.rect().isValid():
            r = self.rect().toRect()
            pos = self.scenePos()
            
            # Extract region from background
            source_rect = QRect(int(pos.x() + r.x()), int(pos.y() + r.y()), r.width(), r.height())
            
            # Ensure within bounds
            if source_rect.intersects(QRect(0, 0, self.bg_pixmap.width(), self.bg_pixmap.height())):
                cropped = self.bg_pixmap.copy(source_rect)
                
                # Convert to QImage for blur
                img = cropped.toImage()
                if not img.isNull():
                    # Convert to numpy for OpenCV blur
                    width, height = img.width(), img.height()
                    ptr = img.bits()
                    ptr.setsize(height * width * 4)
                    arr = np.frombuffer(ptr, np.uint8).reshape((height, width, 4))
                    
                    
                    # Apply Gaussian Blur with dynamic strength (must be odd)
                    k = self.blur_strength | 1 # Ensure odd
                    blurred = cv2.GaussianBlur(arr, (k, k), 0)

                    
                    # Convert back to QImage
                    qimg = QImage(blurred.data, width, height, width * 4, QImage.Format.Format_RGBA8888)
                    blurred_pixmap = QPixmap.fromImage(qimg)
                    
                    # Draw blurred content
                    painter.drawPixmap(r, blurred_pixmap)
        
        # Draw border
        if self.isSelected():
            painter.setPen(QPen(QColor(0, 175, 255), 2, Qt.PenStyle.DashLine))
        else:
             # Very subtle border when not selected, or invisible?
             # Let's make it invisible unless hovered/selected for cleaner look
             painter.setPen(Qt.PenStyle.NoPen)
        
        painter.setBrush(Qt.BrushStyle.NoBrush)
        painter.drawRect(self.rect())
        
        if self.is_global:
            color = QColor(0, 255, 0)
            painter.setFont(QFont("Segoe UI", 8))
            painter.setPen(QPen(color, 1))
            painter.drawText(self.rect().adjusted(5, 5, -5, -5), Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft, "GLOBAL")
        
        # Draw handles when selected
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

class ZoomItem(QGraphicsItem):
    """Interactive zoom box with target marker"""
    def __init__(self, rect, target_point, pixmap=None, is_global=False, uid=None):
        super().__init__()
        self.item_type = 'zoom'
        self.is_global = is_global
        self.uid = uid if uid else str(uuid.uuid4())
        self.box_rect = rect
        self.target = target_point
        self.pixmap = pixmap
        self.handle_size = 10
        self.selected_handle = None
        self.border_color = QColor(255, 255, 255) # Default White
        
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setAcceptHoverEvents(True)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        self.setCacheMode(QGraphicsItem.CacheMode.NoCache)
    
    def itemChange(self, change, value):
        # Update bounding rect when position changes
        if change == QGraphicsItem.GraphicsItemChange.ItemPositionChange:
            self.prepareGeometryChange()
        return super().itemChange(change, value)

    def boundingRect(self):
        # Include arrow in bounding box - MUST be in local coordinates
        # Target is global (scene) coords, self.pos() is global (scene) coords
        local_target = self.target - self.pos()
        box_center = self.box_rect.center()
        
        # Rect covering box and line to target
        arrow_rect = QRectF(local_target, box_center).normalized()
        
        # Merge with box rect and add buffer for stroke width and handles
        return self.box_rect.united(arrow_rect).adjusted(-20, -20, 20, 20)

    def get_handles(self):
        r = self.box_rect
        return {
            'tl': QPointF(r.left(), r.top()),
            'tr': QPointF(r.right(), r.top()),
            'bl': QPointF(r.left(), r.bottom()),
            'br': QPointF(r.right(), r.bottom())
        }

    def get_handle_at(self, pos):
        handles = self.get_handles()
        for name, handle_pos in handles.items():
            if (pos - handle_pos).manhattanLength() < self.handle_size:
                return name
        return None

    def hoverMoveEvent(self, event):
        handle = self.get_handle_at(event.pos())
        if handle:
            cursors = {
                'tl': Qt.CursorShape.SizeFDiagCursor, 'br': Qt.CursorShape.SizeFDiagCursor,
                'tr': Qt.CursorShape.SizeBDiagCursor, 'bl': Qt.CursorShape.SizeBDiagCursor
            }
            self.setCursor(cursors.get(handle, Qt.CursorShape.SizeAllCursor))
        else:
            self.setCursor(Qt.CursorShape.SizeAllCursor)

    def mousePressEvent(self, event):
        self.selected_handle = self.get_handle_at(event.pos())
        if not self.selected_handle:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.selected_handle:
            pos = event.pos()
            r = self.box_rect
            
            if 'l' in self.selected_handle:
                r.setLeft(pos.x())
            if 'r' in self.selected_handle:
                r.setRight(pos.x())
            if 't' in self.selected_handle:
                r.setTop(pos.y())
            if 'b' in self.selected_handle:
                r.setBottom(pos.y())
            
            if r.width() < 50: r.setWidth(50)
            if r.height() < 50: r.setHeight(50)
            
            self.box_rect = r.normalized()
            self.update()
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self.selected_handle = None
        super().mouseReleaseEvent(event)

    def paint(self, painter, option, widget):
        if not self.scene(): return
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # 1. Draw Arrow (from box center to target, accounting for item position)
        center = self.box_rect.center()
        # Target is in scene coordinates, center is in item coordinates
        target_in_item = self.mapFromScene(self.target)
        
        # Calculate line start/end to avoid drawing over the click marker text
        # Vector from target to box center
        line_vec = center - target_in_item
        length = math.sqrt(line_vec.x()**2 + line_vec.y()**2)
        
        if length > 30: # Only draw if far enough apart
            # Normalize and scale to offset start by 30px (radius of marker)
            offset_x = (line_vec.x() / length) * 35 
            offset_y = (line_vec.y() / length) * 35
            
            start_point = QPointF(target_in_item.x() + offset_x, target_in_item.y() + offset_y)
            
            
            # Use custom border color with transparency for line
            line_col = QColor(self.border_color)
            line_col.setAlpha(200)
            painter.setPen(QPen(line_col, 3, Qt.PenStyle.DashLine))
            painter.drawLine(center, start_point)
        
        # 2. Draw Zoom Box
        painter.setPen(QPen(self.border_color, 3))
        painter.setBrush(QBrush(QColor(20, 20, 20)))
        painter.drawRect(self.box_rect)
        
        # 3. Draw Magnified Content - Get current pixmap from scene
        current_pixmap = self.pixmap
        
        # Try to get the background image from the scene
        if self.scene():
            for item in self.scene().items():
                if isinstance(item, QGraphicsPixmapItem) and not hasattr(item, 'item_type'):
                    current_pixmap = item.pixmap()
                    break
        
        if current_pixmap:
            # Crop area around target (2x magnification)
            src_sz = self.box_rect.width() / 2
            crop_rect = QRectF(self.target.x()-src_sz/2, self.target.y()-src_sz/2, src_sz, src_sz)
            
            # Crop and scale with aspect ratio preserved
            cropped = current_pixmap.copy(crop_rect.toRect())
            scaled = cropped.scaled(self.box_rect.size().toSize(), 
                                  Qt.AspectRatioMode.KeepAspectRatio, 
                                  Qt.TransformationMode.SmoothTransformation)
            
            # Calculate where the image is actually drawn (centered in box)
            img_rect = QRectF(scaled.rect())
            x_offset = (self.box_rect.width() - img_rect.width()) / 2
            y_offset = (self.box_rect.height() - img_rect.height()) / 2
            draw_rect = QRectF(self.box_rect.x() + x_offset, 
                              self.box_rect.y() + y_offset,
                              img_rect.width(), 
                              img_rect.height())
            
            painter.drawPixmap(draw_rect.toRect(), scaled)
            
            # Draw marker at the center of the actual image (not box center)
            painter.setPen(QPen(QColor(255, 0, 0), 2))
            marker_center = draw_rect.center()
            painter.drawEllipse(marker_center, 8, 8)

        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

class SpotlightItem(QGraphicsRectItem):
    """Focus Spotlight: Dims everything OUTSIDE the rect"""
    def __init__(self, x, y, w, h, is_global=False, uid=None):
        super().__init__(x, y, w, h)
        self.item_type = 'spotlight'
        self.is_global = is_global
        self.uid = uid if uid else str(uuid.uuid4())
        
        # Settings
        self.dim_opacity = 0.6 # Renamed from opacity to avoid shadowing QGraphicsItem.opacity()
        self.color = QColor(0, 0, 0)
        self.spotlight_shape = 'rect' # or 'ellipse'
        self.handle_size = 10
        self.selected_handle = None
        
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                      QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                      QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setAcceptHoverEvents(True)
        
    def boundingRect(self):
        # Must return the whole scene rect to draw the global dimmer
        if not self.scene(): return self.rect().adjusted(-10, -10, 10, 10)
        return self.mapRectFromScene(self.scene().sceneRect())

    def shape(self):
        # Interaction area is only the box + handles
        path = QPainterPath()
        path.addRect(self.rect().adjusted(-10, -10, 10, 10))
        return path
        
    def get_handles(self):
        r = self.rect()
        return {
            'tl': r.topLeft(),
            'tr': r.topRight(),
            'bl': r.bottomLeft(),
            'br': r.bottomRight()
        }

    def paint(self, painter, option, widget):
        if not self.scene(): return
        
        sr = self.scene().sceneRect()
        r = self.rect()
        scene_bound = self.mapRectFromScene(sr)
        
        # Draw Dimmer
        painter.setPen(Qt.PenStyle.NoPen)
        dim_color = QColor(self.color)
        dim_color.setAlphaF(self.dim_opacity)
        painter.setBrush(dim_color)
        
        path = QPainterPath()
        path.addRect(scene_bound)
        
        hole = QPainterPath()
        if self.spotlight_shape == 'ellipse':
            hole.addEllipse(r)
        else:
            hole.addRect(r)
            
        final_path = path.subtracted(hole)
        painter.drawPath(final_path)
        
        # Outline & Handles
        if self.isSelected():
            painter.setBrush(Qt.BrushStyle.NoBrush)
            painter.setPen(QPen(QColor(255, 255, 255), 2, Qt.PenStyle.DashLine))
            if self.spotlight_shape == 'ellipse':
                painter.drawEllipse(r)
            else:
                painter.drawRect(r)
                
            # Handles
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, 4, 4)

    def hoverMoveEvent(self, event):
        pos = event.pos()
        cursor = Qt.CursorShape.SizeAllCursor
        for name, hpos in self.get_handles().items():
            if (pos - hpos).manhattanLength() < 10:
                if name in ['tl', 'br']: cursor = Qt.CursorShape.SizeFDiagCursor
                elif name in ['tr', 'bl']: cursor = Qt.CursorShape.SizeBDiagCursor
                break
        self.setCursor(cursor)
        super().hoverMoveEvent(event)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            pos = event.pos()
            for name, hpos in self.get_handles().items():
                if (pos - hpos).manhattanLength() < 10:
                    self.selected_handle = name
                    return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.selected_handle:
            pos = event.pos()
            r = self.rect()
            
            if self.selected_handle == 'tl': r.setTopLeft(pos)
            elif self.selected_handle == 'tr': r.setTopRight(pos)
            elif self.selected_handle == 'bl': r.setBottomLeft(pos)
            elif self.selected_handle == 'br': r.setBottomRight(pos)
            
            self.setRect(r.normalized())
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self.selected_handle = None
        super().mouseReleaseEvent(event)
        
    def itemChange(self, change, value):
        if change == QGraphicsItem.GraphicsItemChange.ItemPositionChange:
             self.prepareGeometryChange()
        return super().itemChange(change, value)

class ArrowItem(QGraphicsItem):
    """Standalone arrow with start/end handles"""
    def __init__(self, start_point, end_point, color=QColor(255, 0, 0), width=4, is_global=False, uid=None):
        super().__init__()
        self.item_type = 'arrow'
        self.start = start_point
        self.end = end_point
        self.color = color
        self.width = width
        self.is_global = is_global
        self.uid = uid if uid else str(uuid.uuid4())
        self.handle_size = 10
        self.selected_handle = None
        
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                      QGraphicsItem.GraphicsItemFlag.ItemIsSelectable | 
                      QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setAcceptHoverEvents(True)

    def boundingRect(self):
        return QRectF(self.start, self.end).normalized().adjusted(-20, -20, 20, 20)
        
    def paint(self, painter, option, widget):
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # Draw Arrow
        line_vec = self.end - self.start
        length = math.sqrt(line_vec.x()**2 + line_vec.y()**2)
        angle = math.atan2(line_vec.y(), line_vec.x())
        arrow_size = 15 + self.width

        # Shorten line endpoint so the round cap doesn't stick out of the tip
        # The arrowhead covers the end, so we stop the line a bit early
        shorten = 0
        if length > 10:
             # Shorten by half arrow size roughly, ensuring line ends inside head
             shorten = min(length, arrow_size * 0.6)
        
        line_end = self.end - QPointF(math.cos(angle) * shorten, math.sin(angle) * shorten)
        
        painter.setPen(QPen(self.color, self.width, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
        painter.drawLine(self.start, line_end)
        
        # Draw Arrowhead
        if length > 10:
            p1 = self.end - QPointF(math.cos(angle - math.pi/6) * arrow_size, 
                                  math.sin(angle - math.pi/6) * arrow_size)
            p2 = self.end - QPointF(math.cos(angle + math.pi/6) * arrow_size, 
                                  math.sin(angle + math.pi/6) * arrow_size)
            
            path = QPainterPath()
            path.moveTo(self.end)
            path.lineTo(p1)
            path.lineTo(p2)
            path.closeSubpath()
            
            painter.setBrush(QBrush(self.color))
            painter.setPen(Qt.PenStyle.NoPen)
            painter.drawPath(path)

        # Draw Handles
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            painter.drawEllipse(self.start, 5, 5)
            painter.drawEllipse(self.end, 5, 5)

    def mousePressEvent(self, event):
        pos = event.pos()
        if (pos - self.start).manhattanLength() < 15:
            self.selected_handle = 'start'
        elif (pos - self.end).manhattanLength() < 15:
            self.selected_handle = 'end'
        else:
            self.selected_handle = None
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.selected_handle == 'start':
            self.start = event.pos()
            self.prepareGeometryChange()
            self.update()
        elif self.selected_handle == 'end':
            self.end = event.pos()
            self.prepareGeometryChange()
            self.update()
        else:
            super().mouseMoveEvent(event)

class IconItem(ResizableRectItem):
    """Scalable icon stamp (Monochrome/Tintable symbols)"""
    ICONS = {
        'check': '✔', 
        'cross': '✖', 
        'warn': '⚠', 
        'info': 'ℹ', 
        'star': '★', 
        'idea': '💡',
        'arrow_up': '⬆',
        'arrow_down': '⬇',
        'heart': '❤'
    }
    
    def __init__(self, pos, icon_type='check', size=60, color=QColor(255, 0, 0), is_global=False, uid=None):
        rect = QRectF(0, 0, size, size)
        super().__init__(rect, color=QColor(0,0,0,0)) # Transparent background
        self.item_type = 'icon'
        self.setPos(pos)
        self.icon_type = icon_type
        self.icon_color = color
        self.is_global = is_global
        self.uid = uid if uid else str(uuid.uuid4())
        
    def paint(self, painter, option, widget):
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # 1. Selection border and handles (from ResizableRectItem logic)
        if self.isSelected():
            painter.setPen(QPen(QColor(0, 175, 255), 2, Qt.PenStyle.DashLine))
            painter.drawRect(self.rect())
        
        # 2. Draw Icon
        text = self.ICONS.get(self.icon_type, '?')
        # Adjust font size to fit rect
        font_size = int(min(self.rect().width(), self.rect().height()) * 0.8)
        font = QFont("Segoe UI Symbol", font_size) # Segoe UI Symbol is better for monochrome
        if font_size <= 0: font_size = 1
        font.setPointSize(font_size)
        painter.setFont(font)
        
        painter.setPen(QPen(self.icon_color))
        painter.drawText(self.rect(), Qt.AlignmentFlag.AlignCenter, text)
        
        # 3. Draw Handles
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            for handle_pos in self.get_handles().values():
                painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

class InfoBoxItem(ZoomItem):
    """Text box with moveable arrow pointing to target from border"""
    def __init__(self, rect, target_point, text="Info", is_global=False, uid=None):
        super().__init__(rect, target_point, None, is_global, uid)
        self.item_type = 'infobox'
        self.text = text
        self.bg_color = QColor(40, 40, 40, 220)
        self.text_color = QColor(255, 255, 255)
        self.font_obj = QFont("Segoe UI", 12)
        self.border_width = 2
        self.corner_radius = 5
        self.h_align = 'left' # left, center, right
        self.v_align = 'top'  # top, center, bottom
        # Inherits border_color from ZoomItem (default white)

    def font(self):
        return self.font_obj
    
    def setFont(self, font):
        self.font_obj = font
        self.update()

    def update_font(self, family=None, size=None, bold=None, italic=None, underline=None):
        font = self.font()
        if family: font.setFamily(family)
        if size is not None:
            # Validate size before setting
            if size > 0:
                font.setPointSize(size)
            else:
                font.setPointSize(12) # Fallback to default size
        if bold is not None: font.setBold(bold)
        if italic is not None: font.setItalic(italic)
        if underline is not None: font.setUnderline(underline)
        self.setFont(font)

    def get_handles(self):
        handles = super().get_handles()
        # Add target handle (mapped to local coords)
        if self.scene():
            handles['target'] = self.mapFromScene(self.target)
        return handles

    def mouseMoveEvent(self, event):
        if self.selected_handle == 'target':
            self.target = event.scenePos()
            self.prepareGeometryChange()
            self.update()
        else:
            super().mouseMoveEvent(event)

    def get_intersection_point(self, line_vec):
        """Calculate where the line from center to target leaves the box"""
        # Liang-Barsky simplified for drawing from center
        rect = self.box_rect
        hw = rect.width() / 2
        hh = rect.height() / 2
        
        dx = line_vec.x()
        dy = line_vec.y()
        
        if dx == 0 and dy == 0: return rect.center()
        
        # Calculate slope factors to edges
        tx = hw / abs(dx) if dx != 0 else float('inf')
        ty = hh / abs(dy) if dy != 0 else float('inf')
        
        scale = min(tx, ty)
        
        return rect.center() + QPointF(dx * scale, dy * scale)

    def paint(self, painter, option, widget):
        if not self.scene(): return
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        center = self.box_rect.center()
        target_in_item = self.mapFromScene(self.target)
        line_vec = target_in_item - center
        
        # 1. Calc Start Point (Intersection with Box Border)
        start_point = self.get_intersection_point(line_vec)
        
        dist = (target_in_item - start_point).manhattanLength()
        
        # 2. Draw Arrow
        if dist > 10:
            # Use custom border color for arrow
            arrow_col = QColor(self.border_color)
            arrow_col.setAlpha(200)
            painter.setPen(QPen(arrow_col, 2, Qt.PenStyle.DashLine))
            painter.drawLine(start_point, target_in_item)
            
            # Draw dot at target
            painter.setBrush(QBrush(self.border_color))
            painter.setPen(Qt.PenStyle.NoPen)
            painter.drawEllipse(target_in_item, 4, 4)

        # 3. Draw Box Background
        painter.setPen(QPen(self.border_color, self.border_width))
        painter.setBrush(QBrush(self.bg_color))
        painter.drawRoundedRect(self.box_rect, self.corner_radius, self.corner_radius)
        
        # 4. Draw Text
        painter.setFont(self.font_obj)
        painter.setPen(QPen(self.text_color))
        text_rect = self.box_rect.adjusted(10, 10, -10, -10)
        
        # Determine alignment flags
        h_flags = {
            'left': Qt.AlignmentFlag.AlignLeft,
            'center': Qt.AlignmentFlag.AlignHCenter,
            'right': Qt.AlignmentFlag.AlignRight
        }
        v_flags = {
            'top': Qt.AlignmentFlag.AlignTop,
            'center': Qt.AlignmentFlag.AlignVCenter,
            'bottom': Qt.AlignmentFlag.AlignBottom
        }
        
        align_flags = h_flags.get(self.h_align, Qt.AlignmentFlag.AlignLeft) | \
                      v_flags.get(self.v_align, Qt.AlignmentFlag.AlignTop) | \
                      Qt.TextFlag.TextWordWrap
        
        painter.drawText(text_rect, align_flags, self.text)

        # 5. Draw Handles
        if self.isSelected():
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            
            handles = self.get_handles()
            for name, handle_pos in handles.items():
                # Highlight target handle differently
                if name == 'target':
                     painter.setBrush(QBrush(QColor(255, 50, 50))) # Red for target
                     painter.drawEllipse(handle_pos, 6, 6) # Slightly larger
                     painter.setBrush(QBrush(QColor(0, 175, 255))) # Reset for others
                else:
                    painter.drawEllipse(handle_pos, self.handle_size/2, self.handle_size/2)

    def mouseDoubleClickEvent(self, event):
        text, ok = QInputDialog.getMultiLineText(None, "Text ändern", "Neuer Text:", self.text)
        if ok and text:
            self.text = text
            self.update()

class EditableTextItem(QGraphicsTextItem):
    """Editable text with Photoshop-style interaction"""
    def __init__(self, text, color=QColor(255, 255, 255), is_global=False, uid=None):
        super().__init__(text)
        self.item_type = 'text'
        self.is_global = is_global
        self.uid = uid if uid else str(uuid.uuid4())
        self.setDefaultTextColor(color)
        self.setFont(QFont("Segoe UI", 18, QFont.Weight.Bold))
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemIsFocusable)
        self.setTextInteractionFlags(Qt.TextInteractionFlag.NoTextInteraction)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        self.handle_size = 12
        self.setAcceptHoverEvents(True)
        self.resizing = False
        self.resize_start_pos = None
        self.initial_font_size = 0

    def paint(self, painter, option, widget):
        super().paint(painter, option, widget)
        if self.isSelected():
            # Draw resize handle (bottom right)
            painter.setBrush(QBrush(QColor(0, 175, 255)))
            painter.setPen(QPen(Qt.GlobalColor.white, 1))
            r = self.boundingRect()
            painter.drawEllipse(r.bottomRight(), self.handle_size/2, self.handle_size/2)

    def hoverMoveEvent(self, event):
        # Check if over handle
        r = self.boundingRect()
        handle_pos = r.bottomRight()
        if (event.pos() - handle_pos).manhattanLength() < self.handle_size:
            self.setCursor(Qt.CursorShape.SizeFDiagCursor)
        else:
            self.setCursor(Qt.CursorShape.SizeAllCursor)
        super().hoverMoveEvent(event)

    def mousePressEvent(self, event):
        r = self.boundingRect()
        handle_pos = r.bottomRight()
        if (event.pos() - handle_pos).manhattanLength() < self.handle_size:
            self.resizing = True
            self.resize_start_pos = event.scenePos()
            self.initial_font_size = self.font().pointSize()
            if self.initial_font_size <= 0: self.initial_font_size = 12
            event.accept()
        else:
            self.resizing = False
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.resizing:
            self.prepareGeometryChange()
            delta = event.scenePos() - self.resize_start_pos
            # Use Y delta to scale
            scale_factor = 1 + (delta.y() / 100.0)
            new_size = int(self.initial_font_size * scale_factor)
            new_size = max(8, min(new_size, 200)) # Clamp
            
            f = self.font()
            f.setPointSize(new_size)
            self.setFont(f)
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self.resizing = False
        super().mouseReleaseEvent(event)

    def update_font(self, family=None, size=None, bold=None, italic=None, underline=None):
        font = self.font()
        if family: font.setFamily(family)
        if size is not None:
            # Validate size before setting
            if size > 0:
                font.setPointSize(size)
            else:
                font.setPointSize(18) # Fallback to default size
        if bold is not None: font.setBold(bold)
        if italic is not None: font.setItalic(italic)
        if underline is not None: font.setUnderline(underline)
        self.setFont(font)

    def mouseDoubleClickEvent(self, event):
        # Open simple text dialog
        text, ok = QInputDialog.getText(None, "Text ändern", "Neuer Text:", text=self.toPlainText())
        if ok and text:
            self.setPlainText(text)

    def focusOutEvent(self, event):
        self.setTextInteractionFlags(Qt.TextInteractionFlag.NoTextInteraction)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        super().focusOutEvent(event)

class ClickMarkerItem(QGraphicsItem):
    """Professional click marker with modern design (Glow + Target)"""
    def __init__(self, x, y, number, color=None):
        super().__init__()
        self.item_type = 'click'
        self.center_x = x
        self.center_y = y
        self.number = number
        
        # Get global settings (singleton)
        self.settings = ClickMarkerSettings()
        self.marker_color = color if color else self.settings.color
        
        self.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable | 
                     QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
                     QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setCursor(Qt.CursorShape.SizeAllCursor)
        self.setCacheMode(QGraphicsItem.CacheMode.NoCache)
    
    def boundingRect(self):
        # Read from settings dynamically so changes are immediately visible
        rect_size = self.settings.size * 2.5
        half = rect_size / 2
        return QRectF(self.center_x - half, self.center_y - half, rect_size, rect_size)
    
    def itemChange(self, change, value):
        if change == QGraphicsItem.GraphicsItemChange.ItemPositionHasChanged:
            # Update center logical coordinates
            # boundingRect is relative to pos(), so center is roughly pos() + (0,0) offset
            scene_pos = self.scenePos()
            # We defined center_x/y as initial positions. The item moves relative to that.
            # But simpler: the dot is at (0,0) in local coords if we center rect there.
            # Let's keep logic simple: Updates zoom targets.
            
            if self.scene():
                center_point = self.sceneBoundingRect().center()
                for item in self.scene().items():
                    if hasattr(item, 'item_type') and item.item_type == 'zoom':
                        item.target = center_point
                        item.prepareGeometryChange()
                        item.update()
        return super().itemChange(change, value)

    def paint(self, painter, option, widget):
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        c = QPointF(self.center_x, self.center_y)
        
        # Use settings for all dimensions
        inner_radius = self.settings.size * 0.6
        glow_radius = self.settings.size * 0.9
        
        # USE LIVE COLOR FROM SETTINGS
        current_color = self.settings.color
        
        # Check if transparent
        is_transparent = current_color.alpha() == 0
        
        # 1. Drop Shadow (skip if transparent)
        if not is_transparent:
            painter.setPen(Qt.PenStyle.NoPen)
            painter.setBrush(QColor(0, 0, 0, 100))
            painter.drawEllipse(c + QPointF(2, 2), inner_radius * 0.85, inner_radius * 0.85)
        
        # 2. Outer Glow (optional, skip if transparent)
        if self.settings.show_glow and not is_transparent:
            glow_color = QColor(current_color)
            glow_color.setAlpha(60)
            painter.setBrush(glow_color)
            painter.drawEllipse(c, glow_radius, glow_radius)
        
        # 3. Main Ring
        painter.setPen(QPen(QColor(255, 255, 255), self.settings.border_width))
        
        if is_transparent:
            painter.setBrush(Qt.BrushStyle.NoBrush)
        else:
            painter.setBrush(current_color)
            
        painter.drawEllipse(c, inner_radius, inner_radius)
        
        # 4. Number (Use settings text_color)
        painter.setFont(QFont("Segoe UI", self.settings.number_size, QFont.Weight.Bold))
        
        # Use saved text_color or default to white
        text_col = getattr(self.settings, 'text_color', QColor(255, 255, 255))
        painter.setPen(QPen(text_col, 1))
        
        font_metrics = painter.fontMetrics()
        text_w = font_metrics.horizontalAdvance(str(self.number))
        text_h = font_metrics.capHeight()
        
        pos_point = QPointF(c.x() - text_w/2, c.y() + text_h/2)
        painter.drawText(pos_point, str(self.number))
        
        # 5. Selection Ring
        if self.isSelected():
            painter.setPen(QPen(QColor(255, 255, 255), 2, Qt.PenStyle.DashLine))
            painter.setBrush(Qt.BrushStyle.NoBrush)
            painter.drawEllipse(c, self.settings.size * 1.05, self.settings.size * 1.05)

# ==================== EDITOR SCENE ====================

class EditorScene(QGraphicsScene):
    """Custom scene with drawing tools"""
    def __init__(self, editor):
        super().__init__()
        self.editor = editor
        self.draw_start = None
        self.preview_item = None

    def mousePressEvent(self, event):
        if self.editor.draw_mode and event.button() == Qt.MouseButton.LeftButton:
            self.draw_start = event.scenePos()
            
            if self.editor.draw_mode == 'text':
                dlg = ModernDialog("Text", "Text für das Label:", mode="input", parent=self.editor)
                if dlg.exec():
                    text = dlg.get_text()
                    if text:
                        color_dialog = QColorDialog()
                        color = color_dialog.getColor()
                        if color.isValid():
                            item = EditableTextItem(text, color)
                            item.setPos(self.draw_start)
                            self.addItem(item)
                self.editor.draw_mode = None
                self.editor.update_tool_buttons()
        else:
            super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.draw_start and self.editor.draw_mode in ['blur', 'global_blur', 'zoom', 'crop', 'infobox', 'arrow']:
            # Live preview logic
            if self.preview_item:
                try:
                    if self.preview_item.scene() == self:
                        self.removeItem(self.preview_item)
                except RuntimeError:
                    pass
            
            rect = QRectF(self.draw_start, event.scenePos()).normalized()
            
            if self.editor.draw_mode == 'arrow':
                 self.preview_item = QGraphicsLineItem(QLineF(self.draw_start, event.scenePos()))
                 self.preview_item.setPen(QPen(QColor(255, 0, 0), 4))
                 self.addItem(self.preview_item)
                 return

            # Rect preview for others
            colors = {
                'blur': QColor(0, 0, 0, 100),
                'global_blur': QColor(0, 255, 0, 100),
                'zoom': QColor(0, 175, 255, 100),
                'infobox': QColor(255, 165, 0, 100),
                'crop': QColor(255, 255, 0, 80),
                'spotlight': QColor(0, 0, 0, 150)
            }
            
            self.preview_item = QGraphicsRectItem(rect)
            self.preview_item.setPen(QPen(QColor(255, 255, 255), 2, Qt.PenStyle.DashLine))
            self.preview_item.setBrush(QBrush(colors.get(self.editor.draw_mode, QColor(255, 255, 255, 50))))
            self.addItem(self.preview_item)
        else:
            super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self.draw_start and self.editor.draw_mode:
            rect = QRectF(self.draw_start, event.scenePos()).normalized()
            
            
            if self.editor.draw_mode == 'arrow':
                item = ArrowItem(self.draw_start, event.scenePos())
                self.addItem(item)
                self.draw_start = None
                if self.preview_item: self.removeItem(self.preview_item)
                self.preview_item = None
                self.editor.draw_mode = None
                self.editor.update_tool_buttons()
                return

            if rect.width() > 20 and rect.height() > 20:
                if self.editor.draw_mode in ['blur', 'global_blur', 'zoom', 'infobox', 'spotlight']:
                     self.editor.push_undo() # Undo before create
                
                if self.editor.draw_mode == 'blur':
                    item = BlurItem(rect, False, self.editor.current_pixmap)
                    self.addItem(item)
                elif self.editor.draw_mode == 'spotlight':
                    item = SpotlightItem(rect.x(), rect.y(), rect.width(), rect.height())
                    self.addItem(item)
                elif self.editor.draw_mode == 'global_blur':
                    item = BlurItem(rect, True, self.editor.current_pixmap)
                    self.addItem(item)
                elif self.editor.draw_mode == 'zoom':
                    # Target is the center of the current step's click
                    target = QPointF(self.editor.steps[self.editor.current_idx].x, 
                                   self.editor.steps[self.editor.current_idx].y)
                    item = ZoomItem(rect, target, self.editor.current_pixmap)
                    self.addItem(item)
                elif self.editor.draw_mode == 'infobox':
                    dlg = ModernDialog("Info Box", "Text eingeben:", mode="multiline", parent=self.editor)
                    if dlg.exec():
                        text = dlg.get_text()
                        if text:
                            target = QPointF(self.editor.steps[self.editor.current_idx].x, 
                                           self.editor.steps[self.editor.current_idx].y)
                            item = InfoBoxItem(rect, target, text)
                            self.addItem(item)
                elif self.editor.draw_mode == 'crop':
                    self.editor.push_undo() # Undo crop
                    
                    # Apply global crop and reload all steps
                    self.editor.global_crop = (int(rect.left()), int(rect.top()), 
                                             int(rect.right()), int(rect.bottom()))
                    
                    # Reload current step with crop applied
                    self.editor.load_step(self.editor.current_idx)
                    
                    # Update all thumbnails to show crop
                    self.editor.update_thumbnails()
                    
                    QMessageBox.information(None, "Crop", 
                        f"Globaler Ausschnitt festgelegt: {rect.width():.0f}x{rect.height():.0f}px\n" +
                        "Wird auf alle Schritte angewendet!")
                    self.editor.draw_mode = None
                    self.editor.update_tool_buttons()
            
            if self.preview_item:
                try:
                    if self.preview_item.scene() == self:
                        self.removeItem(self.preview_item)
                except RuntimeError:
                    pass # Item already deleted (e.g. by scene.clear())
                self.preview_item = None
            
            self.draw_start = None
            self.editor.draw_mode = None
            self.editor.update_tool_buttons()
            self.editor.refresh_layer_list()
        else:
            super().mouseReleaseEvent(event)
            self.editor.refresh_layer_list()

# ==================== MAIN EDITOR ====================

class ZoomableGraphicsView(QGraphicsView):
    """Graphics view with mouse wheel zoom and middle button pan"""
    def __init__(self, scene):
        super().__init__(scene)
        self.zoom_factor = 1.0
        self.panning = False
        self.pan_start_pos = None
        
    def wheelEvent(self, event):
        # Zoom towards mouse cursor
        zoom_in_factor = 1.15
        zoom_out_factor = 1 / zoom_in_factor
        
        # Get the old position
        old_pos = self.mapToScene(event.position().toPoint())
        
        if event.angleDelta().y() > 0:
            factor = zoom_in_factor
            self.zoom_factor *= factor
        else:
            factor = zoom_out_factor
            self.zoom_factor *= factor
        
        # Limit zoom range
        if 0.1 < self.zoom_factor < 10:
            self.scale(factor, factor)
            
            # Get the new position
            new_pos = self.mapToScene(event.position().toPoint())
            
            # Move scene to keep mouse position
            delta = new_pos - old_pos
            self.translate(delta.x(), delta.y())
        else:
            self.zoom_factor /= factor  # Revert if out of range
    
    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.MiddleButton:
            # Start panning with middle mouse button
            self.panning = True
            self.pan_start_pos = event.pos()
            self.setCursor(Qt.CursorShape.ClosedHandCursor)
            event.accept()
        else:
            super().mousePressEvent(event)
    
    def mouseMoveEvent(self, event):
        if self.panning and self.pan_start_pos:
            # Pan the view
            delta = event.pos() - self.pan_start_pos
            self.pan_start_pos = event.pos()
            
            # Scroll the view
            self.horizontalScrollBar().setValue(self.horizontalScrollBar().value() - delta.x())
            self.verticalScrollBar().setValue(self.verticalScrollBar().value() - delta.y())
            event.accept()
        else:
            super().mouseMoveEvent(event)
    
    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MouseButton.MiddleButton:
            # Stop panning
            self.panning = False
            self.pan_start_pos = None
            self.setCursor(Qt.CursorShape.ArrowCursor)
            event.accept()
        else:
            super().mouseReleaseEvent(event)

class ProEditor(QMainWindow):
    def __init__(self, steps, globals, crop, save_cb, project_name=None, parent_window=None, settings=None):
        super().__init__()
        self.settings = settings if settings else AppSettings()
        self.recorder_window = parent_window
        self.setWindowTitle("ClickStep Guide Editor - Photoshop Style")
        self.resize(1600, 1000)
        self.setWindowIcon(QIcon(resource_path("assets/icon.ico")))
        self.setStyleSheet(self.get_stylesheet())
        
        self.steps = steps
        self.global_layers = globals
        self.global_crop = crop
        self.save_cb = save_cb
        self.current_idx = 0
        self.draw_mode = None
        self.current_project_name = project_name
        if project_name:
            self.setWindowTitle(f"ClickStep Guide Editor - {project_name}")
        
        self.undo_stack = [] # List of snapshots
        
        self.scene = EditorScene(self)
        self.view = ZoomableGraphicsView(self.scene)
        self.view.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.view.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.view.setDragMode(QGraphicsView.DragMode.RubberBandDrag)
        self.view.setBackgroundBrush(QBrush(QColor("#f8fafc")))
        # Use FullViewportUpdate to prevent smearing/artifacts during zoom/panning
        self.view.setViewportUpdateMode(QGraphicsView.ViewportUpdateMode.FullViewportUpdate)
        
        self.setCentralWidget(self.view)
        self.setup_ui()
        
        # Initialize project directory
        base = os.environ.get('LOCALAPPDATA', os.path.expanduser("~"))
        self.project_base_path = os.path.join(base, "ClickStepGuide", "projects")
        os.makedirs(self.project_base_path, exist_ok=True)
        
        # CRITICAL: Connect selection changes to properties panel updates
        self.scene.selectionChanged.connect(self.update_properties)
        
        # Subtle Branding Footer
        self.statusBar().showMessage("ClickStep Guide Pro Engine | Professionelles Dokumentations-System Enabled")
        self.statusBar().setStyleSheet("color: #64748b; background: #f8fafc; border-top: 1px solid #e2e8f0;")
        
        if self.steps:
            self.load_step(0)

    def get_project_dir(self):
        """Returns the project directory in Local AppData for Store compliance"""
        return self.project_base_path
        
    def closeEvent(self, event):
        """Show recorder window when editor is closed"""
        if self.recorder_window:
            self.recorder_window.show()
        super().closeEvent(event)

    def get_stylesheet(self):
        is_dark = self.settings.theme == "dark"
        if is_dark:
            return """
                QMainWindow, QWidget { background-color: #1e1e1e; color: #cccccc; font-family: "Segoe UI", sans-serif; font-size: 13px; }
                QDockWidget { titlebar-close-icon: url(none); titlebar-normal-icon: url(none); border: 1px solid #333333; }
                QDockWidget::title { background: #252526; padding: 8px; font-weight: bold; color: #e0e0e0; border-bottom: 1px solid #333333; text-transform: uppercase; letter-spacing: 0.5px; }
                QToolBar { background: #252526; border-bottom: 1px solid #333333; spacing: 8px; padding: 6px; }
                QListWidget { background-color: #252526; border: 1px solid #333333; border-radius: 4px; outline: none; }
                QListWidget::item { padding: 8px; border-bottom: 1px solid #2d2d2d; color: #cccccc; }
                QListWidget::item:selected { background-color: #094771; color: white; border-left: 3px solid #007acc; }
                QPushButton { background-color: #333333; color: #ffffff; border: 1px solid #3e3e42; padding: 6px 14px; border-radius: 4px; font-weight: 600; }
                QPushButton:hover { background-color: #3e3e42; border-color: #505050; }
                QPushButton:pressed { background-color: #1e1e1e; border-color: #007acc; }
                QPushButton:checked { background-color: #094771; border-color: #007acc; color: white; }
                QScrollBar:vertical { border: none; background: #1e1e1e; width: 12px; margin: 0; }
                QScrollBar::handle:vertical { background: #424242; min-height: 20px; border-radius: 6px; margin: 2px; }
                QScrollBar::handle:vertical:hover { background: #606060; }
                QSlider::groove:horizontal { background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2d2d2d, stop:1 #1e1e1e); height: 6px; border-radius: 3px; }
                QSlider::handle:horizontal { background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3b82f6, stop:1 #2563eb); width: 18px; height: 18px; margin: -6px 0; border-radius: 9px; border: 2px solid #1e40af; }
                QSlider::handle:horizontal:hover { background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #60a5fa, stop:1 #3b82f6); border-color: #3b82f6; }
                QSlider::sub-page:horizontal { background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #3b82f6, stop:1 #60a5fa); border-radius: 3px; }
                QSpinBox { background: #2d2d2d; color: #e0e0e0; border: 1px solid #3e3e42; border-radius: 4px; padding: 4px 8px; }
                QSpinBox:hover { border-color: #3b82f6; }
                QComboBox { background: #2d2d2d; color: #e0e0e0; border: 1px solid #3e3e42; border-radius: 4px; padding: 4px 8px; }
                QComboBox:hover { border-color: #3b82f6; }
            """
        else:
            return """
                QMainWindow, QWidget { background-color: #ffffff; color: #202124; font-family: "Segoe UI", sans-serif; font-size: 13px; }
                QDockWidget { titlebar-close-icon: url(none); titlebar-normal-icon: url(none); border: 1px solid #e0e0e0; }
                QDockWidget::title { background: #f8f9fa; padding: 8px; font-weight: bold; color: #5f6368; border-bottom: 1px solid #e0e0e0; text-transform: uppercase; letter-spacing: 0.5px; }
                QToolBar { background: #f8f9fa; border-bottom: 1px solid #dadce0; spacing: 8px; padding: 6px; }
                QListWidget { background-color: #ffffff; border: 1px solid #dadce0; border-radius: 4px; outline: none; }
                QListWidget::item { padding: 8px; border-bottom: 1px solid #f1f3f4; color: #3c4043; }
                QListWidget::item:selected { background-color: #e8f0fe; color: #1967d2; border-left: 3px solid #1a73e8; }
                QPushButton { background-color: #ffffff; color: #3c4043; border: 1px solid #dadce0; padding: 6px 14px; border-radius: 4px; font-weight: 600; }
                QPushButton:hover { background-color: #f8f9fa; border-color: #bdc1c6; }
                QPushButton:pressed { background-color: #f1f3f4; border-color: #1a73e8; }
                QPushButton:checked { background-color: #e8f0fe; border-color: #1a73e8; color: #1967d2; }
                QScrollBar:vertical { border: none; background: #ffffff; width: 12px; margin: 0; }
                QScrollBar::handle:vertical { background: #dadce0; min-height: 20px; border-radius: 6px; margin: 2px; }
                QScrollBar::handle:vertical:hover { background: #bdc1c6; }
                QSlider::groove:horizontal { background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #e2e8f0, stop:1 #cbd5e1); height: 6px; border-radius: 3px; }
                QSlider::handle:horizontal { background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3b82f6, stop:1 #2563eb); width: 18px; height: 18px; margin: -6px 0; border-radius: 9px; border: 2px solid #1e40af; }
                QSlider::handle:horizontal:hover { background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #60a5fa, stop:1 #3b82f6); border-color: #3b82f6; }
                QSlider::sub-page:horizontal { background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #3b82f6, stop:1 #60a5fa); border-radius: 3px; }
                QSpinBox { background: white; color: #334155; border: 1px solid #cbd5e1; border-radius: 6px; padding: 4px 8px; }
                QSpinBox:hover { border-color: #3b82f6; background: #f8fafc; }
                QComboBox { background: white; color: #334155; border: 1px solid #cbd5e1; border-radius: 6px; padding: 4px 8px; }
                QComboBox:hover { border-color: #3b82f6; background: #f8fafc; }
            """

    def setup_ui(self):
        # Office-style Toolbar with grouped sections
        toolbar = self.addToolBar("Tools")
        toolbar.setMovable(False)
        toolbar.setIconSize(QSize(20, 20))
        toolbar.setStyleSheet("""
            QToolBar {
                background: white;
                border-bottom: 1px solid #d1d5db;
                spacing: 0px;
                padding: 4px 8px;
            }
            QToolBar::separator {
                background: #d1d5db;
                width: 1px;
                margin: 4px 12px;
            }
            QPushButton {
                background: transparent;
                color: #374151;
                border: 1px solid transparent;
                padding: 4px 8px;
                border-radius: 3px;
                font-size: 11px;
                min-width: 50px;
            }
            QPushButton:hover {
                background: #f3f4f6;
                border-color: #e5e7eb;
            }
            QPushButton:pressed {
                background: #e5e7eb;
            }
            QPushButton:checked {
                background: #dbeafe;
                border-color: #93c5fd;
                color: #1e40af;
            }
            QPushButton[objectName="DestructiveButton"] {
                color: #dc2626;
            }
            QPushButton[objectName="DestructiveButton"]:hover {
                background: #fee2e2;
                border-color: #fca5a5;
            }
            QPushButton[objectName="AccentButton"] {
                background: #10b981;
                color: white;
                font-weight: 600;
            }
            QPushButton[objectName="AccentButton"]:hover {
                background: #059669;
            }
            QLabel {
                color: #6b7280;
                font-size: 9px;
                padding: 0px 4px;
            }
        """)
        
        # Helper function to create a toolbar group
        def create_group(label_text, buttons):
            group_widget = QWidget()
            group_layout = QVBoxLayout(group_widget)
            group_layout.setContentsMargins(0, 0, 0, 0)
            group_layout.setSpacing(2)
            
            # Buttons row
            btn_row = QWidget()
            btn_layout = QHBoxLayout(btn_row)
            btn_layout.setContentsMargins(0, 0, 0, 0)
            btn_layout.setSpacing(2)
            
            for btn in buttons:
                btn_layout.addWidget(btn)
            
            # Label
            label = QLabel(label_text)
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            group_layout.addWidget(btn_row)
            group_layout.addWidget(label)
            
            return group_widget
        
        # Create buttons
        self.btn_select = QPushButton("🏹 Select")
        self.btn_select.setCheckable(True)
        self.btn_select.setChecked(True)
        self.btn_select.clicked.connect(lambda: self.set_tool(None))
        
        self.btn_arrow = QPushButton("↗️ Pfeil")
        self.btn_arrow.setCheckable(True)
        self.btn_arrow.clicked.connect(lambda: self.set_tool('arrow'))
        
        self.btn_text = QPushButton("📝 Text")
        self.btn_text.setCheckable(True)
        self.btn_text.clicked.connect(lambda: self.set_tool('text'))
        
        self.btn_infobox = QPushButton("💬 Info")
        self.btn_infobox.setCheckable(True)
        self.btn_infobox.clicked.connect(lambda: self.set_tool('infobox'))
        
        self.btn_icon = QPushButton("⭐ Icon")
        self.btn_icon.clicked.connect(self.add_icon_dialog)
        
        self.btn_zoom = QPushButton("🔍 Zoom")
        self.btn_zoom.setCheckable(True)
        self.btn_zoom.clicked.connect(lambda: self.set_tool('zoom'))
        
        self.btn_blur = QPushButton("🛡️ Blur")
        self.btn_blur.setCheckable(True)
        self.btn_blur.clicked.connect(lambda: self.set_tool('blur'))
        
        self.btn_spotlight = QPushButton("🔦 Fokus")
        self.btn_spotlight.setCheckable(True)
        self.btn_spotlight.setToolTip("Spotlight Effekt hinzufügen")
        self.btn_spotlight.clicked.connect(lambda: self.set_tool('spotlight'))
        
        self.btn_global_blur = QPushButton("🌍 Global")
        self.btn_global_blur.setCheckable(True)
        self.btn_global_blur.clicked.connect(lambda: self.set_tool('global_blur'))
        
        self.btn_crop = QPushButton("✂️ Crop")
        self.btn_crop.setCheckable(True)
        self.btn_crop.clicked.connect(lambda: self.set_tool('crop'))
        
        btn_delete = QPushButton("🗑️ Delete")
        btn_delete.setObjectName("DestructiveButton")
        btn_delete.clicked.connect(self.delete_selected)
        
        btn_delete_step = QPushButton("❌ Step")
        btn_delete_step.setObjectName("DestructiveButton")
        btn_delete_step.setToolTip("Schritt löschen")
        btn_delete_step.clicked.connect(self.delete_step)
        
        btn_save = QPushButton("💾 Save")
        btn_save.setToolTip("Strg+S")
        btn_save.setShortcut("Ctrl+S")
        btn_save.clicked.connect(lambda chk=False: self.save_project(save_as=False))

        btn_save_as = QPushButton("📝 Save As")
        btn_save_as.setToolTip("Speichern unter...")
        btn_save_as.clicked.connect(lambda chk=False: self.save_project(save_as=True))
        
        btn_undo = QPushButton("↩️ Undo")
        btn_undo.setShortcut("Ctrl+Z")
        btn_undo.clicked.connect(self.undo)
        
        btn_export = QPushButton("💾 EXPORT")
        btn_export.setObjectName("AccentButton")
        btn_export.clicked.connect(self.on_export_clicked)
        
        # Add grouped sections to toolbar
        toolbar.addWidget(create_group("Werkzeuge", [self.btn_select, self.btn_arrow, self.btn_text]))
        toolbar.addSeparator()
        
        toolbar.addWidget(create_group("Anmerkungen", [self.btn_infobox, self.btn_icon, self.btn_zoom]))
        toolbar.addSeparator()
        
        toolbar.addWidget(create_group("Effekte", [self.btn_blur, self.btn_spotlight, self.btn_global_blur, self.btn_crop]))
        toolbar.addSeparator()
        
        toolbar.addWidget(create_group("Bearbeiten", [btn_delete, btn_delete_step, btn_undo]))
        toolbar.addSeparator()
        
        toolbar.addWidget(create_group("Datei", [btn_save, btn_save_as]))
        toolbar.addSeparator()
        
        toolbar.addWidget(btn_export)
        

        # Left Dock: Thumbnails
        self.dock_thumbs = QDockWidget("SCHRITTE", self)
        self.thumb_list = QListWidget()
        self.thumb_list.currentRowChanged.connect(self.on_step_changed)
        
        # Enable Drag and Drop reordering
        self.thumb_list.setDragDropMode(QAbstractItemView.DragDropMode.InternalMove)
        self.thumb_list.model().rowsMoved.connect(self.on_steps_reordered)
        
        self.dock_thumbs.setWidget(self.thumb_list)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.dock_thumbs)
        
        # Right Dock: Layers (Modern 3-Row Grid Layout)
        self.dock_layers = QDockWidget("EBENEN", self)
        self.layer_panel = QWidget()
        self.layer_layout = QVBoxLayout(self.layer_panel)
        self.layer_layout.setSpacing(8)
        self.layer_layout.setContentsMargins(8, 8, 8, 8)
        
        # Row 1: Project Header
        project_header = QLabel("📋 PROJEKT ENTWURF")
        project_header.setStyleSheet("""
            font-weight: 700; 
            color: #1e293b; 
            padding: 10px; 
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #e0e7ff, stop:1 #dbeafe);
            border-radius: 6px;
            font-size: 12px;
            letter-spacing: 0.5px;
        """)
        self.layer_layout.addWidget(project_header)
        
        # Row 2: Step Layers Section
        step_container = QWidget()
        step_container.setStyleSheet("""
            QWidget {
                background: rgba(255, 255, 255, 0.6);
                border: 1px solid #e2e8f0;
                border-radius: 8px;
            }
        """)
        step_container_layout = QVBoxLayout(step_container)
        step_container_layout.setContentsMargins(8, 8, 8, 8)
        step_container_layout.setSpacing(6)
        
        lbl_local = QLabel("🎨 BILD-EBENEN")
        lbl_local.setStyleSheet("""
            font-size: 11px; 
            font-weight: 700; 
            color: #3b82f6; 
            padding: 4px 0;
            letter-spacing: 0.5px;
        """)
        step_container_layout.addWidget(lbl_local)
        
        self.step_layer_list = LayerListWidget(is_global_list=False, main_window=self)
        self.step_layer_list.itemClicked.connect(self.on_layer_clicked)
        self.step_layer_list.setMaximumHeight(180)
        step_container_layout.addWidget(self.step_layer_list)
        
        self.layer_layout.addWidget(step_container)
        
        # Row 3: Transfer Buttons (Compact)
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(6)
        
        btn_to_global = QPushButton("↓ Global")
        btn_to_global.setStyleSheet("""
            QPushButton {
                background: rgba(59, 130, 246, 0.1);
                color: #3b82f6;
                border: 1px solid rgba(59, 130, 246, 0.3);
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover {
                background: rgba(59, 130, 246, 0.2);
                border-color: #3b82f6;
            }
        """)
        btn_to_global.clicked.connect(self.move_layer_to_global)
        
        btn_to_step = QPushButton("↑ Lokal")
        btn_to_step.setStyleSheet("""
            QPushButton {
                background: rgba(16, 185, 129, 0.1);
                color: #10b981;
                border: 1px solid rgba(16, 185, 129, 0.3);
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover {
                background: rgba(16, 185, 129, 0.2);
                border-color: #10b981;
            }
        """)
        btn_to_step.clicked.connect(self.move_layer_to_step)
        
        btn_layout.addWidget(btn_to_global)
        btn_layout.addWidget(btn_to_step)
        self.layer_layout.addLayout(btn_layout)
        
        # Row 4: Global Layers Section
        global_container = QWidget()
        global_container.setStyleSheet("""
            QWidget {
                background: rgba(255, 255, 255, 0.6);
                border: 1px solid #e2e8f0;
                border-radius: 8px;
            }
        """)
        global_container_layout = QVBoxLayout(global_container)
        global_container_layout.setContentsMargins(8, 8, 8, 8)
        global_container_layout.setSpacing(6)
        
        lbl_global = QLabel("🌍 GLOBAL-EBENEN")
        lbl_global.setStyleSheet("""
            font-size: 11px; 
            font-weight: 700; 
            color: #10b981; 
            padding: 4px 0;
            letter-spacing: 0.5px;
        """)
        global_container_layout.addWidget(lbl_global)
        
        self.global_layer_list = LayerListWidget(is_global_list=True, main_window=self)
        self.global_layer_list.itemClicked.connect(self.on_layer_clicked)
        self.global_layer_list.setMaximumHeight(180)
        global_container_layout.addWidget(self.global_layer_list)
        
        self.layer_layout.addWidget(global_container)
        
        # Add stretch to push everything to top
        self.layer_layout.addStretch()
        
        self.dock_layers.setWidget(self.layer_panel)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.dock_layers)

        # Right Dock: Properties
        self.dock_props = QDockWidget("EIGENSCHAFTEN", self)
        self.props_widget = QWidget()
        self.props_layout = QVBoxLayout(self.props_widget)
        self.props_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.dock_props.setWidget(self.props_widget)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.dock_props)
        self.tabifyDockWidget(self.dock_layers, self.dock_props)
        
        # Make EBENEN the default active tab
        self.dock_layers.raise_()
        
        # Bottom Dock: Documentation
        self.dock_docs = QDockWidget("DOKUMENTATION / ANLEITUNG", self)
        self.dock_docs.setContentsMargins(0, 5, 0, 0)
        # Allow docking but default to bottom
        self.doc_widget = QWidget()
        doc_layout = QVBoxLayout(self.doc_widget)
        doc_layout.setContentsMargins(10, 10, 10, 10)
        
        lbl_instr = QLabel("<b>Schritt-Beschreibung:</b>")
        lbl_instr.setStyleSheet("color: #007acc;")
        doc_layout.addWidget(lbl_instr)
        
        self.txt_description = QTextEdit()
        self.txt_description.setPlaceholderText("Beschreibe hier was der Benutzer tun soll (z.B. 'Klicke auf Speichern')...")
        self.txt_description.setStyleSheet("background: #252526; color: #fff; border: 1px solid #3e3e42; font-size: 14px; font-family: 'Segoe UI'; border-radius: 4px; padding: 8px;")
        self.txt_description.setFixedHeight(80) 
        self.txt_description.textChanged.connect(self.save_description)
        doc_layout.addWidget(self.txt_description)
        
        self.dock_docs.setWidget(self.doc_widget)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.dock_docs)

        # Signal connection already done in __init__, removed duplicate
        self.update_thumbnails()

    def on_steps_reordered(self, parent, start, end, destination, row):
        """Synchronize self.steps with QListWidget reordering"""
        # Note: QListWidget reorders internally, we need to rebuild our self.steps list 
        # based on the new order of data in the items (we store the original step index or object)
        new_steps = []
        for i in range(self.thumb_list.count()):
            item = self.thumb_list.item(i)
            # Find the step object associated with this item
            step_obj = item.data(Qt.ItemDataRole.UserRole)
            if step_obj:
                new_steps.append(step_obj)
        
        if new_steps:
            self.steps = new_steps
            # Refresh labels/thumbnails to update step numbers (#1, #2, etc)
            self.update_thumbnails()
            
    def update_thumbnails(self):
        """Update thumbnail list with step numbers"""
        # Block signals during rebuild to prevent on_step_changed firing
        self.thumb_list.blockSignals(True)
        self.thumb_list.clear()
        
        for i, step in enumerate(self.steps):
            # Create thumbnail with step number overlay
            img = step.raw_img
            h, w = img.shape[:2]
            
            # Resize for thumbnail
            thumb_h = 80
            thumb_w = int(w * thumb_h / h)
            thumb = cv2.resize(img, (thumb_w, thumb_h))
            
            # Add step number overlay
            cv2.putText(thumb, f"#{i+1}", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 
                       1.0, (0, 0, 0), 6, cv2.LINE_AA)
            cv2.putText(thumb, f"#{i+1}", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 
                       1.0, (255, 255, 255), 2, cv2.LINE_AA)
            
            # Convert to QPixmap
            rgb = cv2.cvtColor(thumb, cv2.COLOR_BGR2RGB)
            qimg = QImage(rgb.data, thumb_w, thumb_h, thumb_w*3, QImage.Format.Format_RGB888)
            pix = QPixmap.fromImage(qimg)
            
            item = QListWidgetItem(f"Schritt {i+1}")
            item.setIcon(QIcon(pix))
            item.setData(Qt.ItemDataRole.UserRole, step) # Store the step object for reordering sync
            self.thumb_list.addItem(item)
        
        self.thumb_list.blockSignals(False)
        self.thumb_list.setCurrentRow(self.current_idx)

    def on_step_changed(self, idx):
        """Handle step change from thumbnail list"""
        if idx < 0 or idx >= len(self.steps):
            return
        # Save current state before switching
        self.save_current_state()
        # Load new step
        self.load_step(idx)

    def delete_step(self):
        """Delete current step"""
        if len(self.steps) <= 1:
            QMessageBox.warning(self, "Fehler", "Mindestens ein Schritt muss vorhanden sein!")
            return
        
        dlg = ModernDialog('Löschen', f'Schritt {self.current_idx + 1} wirklich löschen?', mode="confirm", parent=self)
        
        if dlg.exec():
            self.push_undo() # Save state before delete
            del self.steps[self.current_idx]
            
            # Adjust current index
            if self.current_idx >= len(self.steps):
                self.current_idx = len(self.steps) - 1
            
            self.update_thumbnails()
            self.load_step(self.current_idx)

    def on_export_clicked(self):
        self.save_current_state()
        self.save_cb(self.steps, self.global_layers, self.global_crop)

    def save_project(self, save_as=False):
        """Save project to disk"""
        self.save_current_state()
        
        name = self.current_project_name
        
        # Check if we need to ask for name
        if save_as is True or not name:
            prefill = name if name else f"Project_{datetime.now().strftime('%Y%m%d_%H%M')}"
            dlg = ModernDialog("Projekt speichern", "Projektname:", mode="input", default_text=prefill, parent=self)
            
            if not dlg.exec(): return
            new_name = dlg.get_text()
            if not new_name: return
            
            name = new_name
            self.current_project_name = name
            self.setWindowTitle(f"ClickStep Guide - {name}")
        
        base_path = os.path.join(self.get_project_dir(), name)
        img_path = os.path.join(base_path, "images")
        if os.path.exists(base_path) and (save_as is True):
             # Maybe warn? But standard Save As just overwrites/uses that name
             pass
             
        os.makedirs(img_path, exist_ok=True)
        
        data = {
            "global_crop": self.global_crop,
            "global_layers": [{"type": l.type, "data": l.data, "label": l.label, "uid": l.uid} for l in self.global_layers],
            "steps": []
        }
        
        for i, s in enumerate(self.steps):
            filename = f"step_{i}.png"
            cv2.imwrite(os.path.join(img_path, filename), s.raw_img)
            
            data["steps"].append({
                "image": filename,
                "description": s.description,
                "layers": [{"type": l.type, "data": l.data, "label": l.label} for l in s.layers]
            })
            
        with open(os.path.join(base_path, "project.json"), "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
        
        # Use simple label overlay or status bar instead of annoying popup?
        # For now, just status bar, or a non-blocking modern toast
        self.statusBar().showMessage(f"Projekt '{name}' erfolgreich gespeichert!", 3000)
        # QMessageBox.information(self, "Erfolg", f"Projekt '{name}' gespeichert!")
        self.save_cb(None, None, None) # Trigger project list update in main window

    def set_tool(self, tool):
        self.draw_mode = tool
        self.update_tool_buttons()

    def update_tool_buttons(self):
        self.btn_select.setChecked(self.draw_mode is None)
        self.btn_arrow.setChecked(self.draw_mode == 'arrow')
        self.btn_infobox.setChecked(self.draw_mode == 'infobox')
        self.btn_blur.setChecked(self.draw_mode == 'blur')
        self.btn_zoom.setChecked(self.draw_mode == 'zoom')
        self.btn_text.setChecked(self.draw_mode == 'text')
        self.btn_global_blur.setChecked(self.draw_mode == 'global_blur')
        self.btn_global_blur.setChecked(self.draw_mode == 'global_blur')
        self.btn_crop.setChecked(self.draw_mode == 'crop')
        self.btn_spotlight.setChecked(self.draw_mode == 'spotlight')

    def delete_selected(self):
        """Delete selected layer items"""
        items = self.scene.selectedItems()
        if not items: return
        
        print(f"DEBUG: Deleting {len(items)} items...")
        
        self.push_undo() # Save state before delete
        
        # Block signals to prevent UI updates during deletion
        self.scene.blockSignals(True)
        try:
            # Clear selection first
            self.scene.clearSelection()
            
            # Remove items
            for item in items:
                # Protect ClickMarkers from deletion
                if hasattr(item, 'item_type') and item.item_type == 'click': 
                    continue
                try:
                    if item.scene() == self.scene:
                        self.scene.removeItem(item)
                except RuntimeError as e:
                    print(f"Error removing item: {e}")
                except Exception as e:
                    print(f"Unexpected error removing item: {e}")
        finally:
            self.scene.blockSignals(False)
        
        # Manually trigger update once after everything is done
        self.update_properties()
        
        # NOW save the state (items are already gone from scene)
        try:
            self.save_current_state()
            self.refresh_layer_list()
        except Exception as e:
            print(f"Error saving state after delete: {e}")

    def keyPressEvent(self, event):
        """Handle Delete key to remove items"""
        # Check if we are currently editing text
        focus_item = self.scene.focusItem()
        if isinstance(focus_item, QGraphicsTextItem) and (focus_item.textInteractionFlags() & Qt.TextInteractionFlag.TextEditable):
             # Let the text item handle the delete key (delete character)
             super().keyPressEvent(event)
             return

        if event.key() == Qt.Key.Key_Delete:
            self.delete_selected()
            event.accept()
        else:
            super().keyPressEvent(event)

    def open_marker_settings(self):
        """Open click marker settings dialog"""
        from marker_settings_dialog import ClickMarkerSettingsDialog
        dialog = ClickMarkerSettingsDialog(self)
        if dialog.exec():
            # Refresh all markers - force complete redraw
            for item in self.scene.items():
                if hasattr(item, 'item_type') and item.item_type == 'click':
                    item.prepareGeometryChange()  # Notify scene of size change
                    item.update()  # Trigger repaint
            
            # Force scene update
            self.scene.update()
            self.view.viewport().update()
            
            self.statusBar().showMessage("Marker-Einstellungen gespeichert!", 3000)

    def on_step_changed(self, idx):
        if idx < 0: return
        self.save_current_state()
        self.load_step(idx)

    def handle_layer_drop(self, target_is_global):
        """Called when an item is dropped into one of the layer lists"""
        # Find all items in the target list
        target_list = self.global_layer_list if target_is_global else self.step_layer_list
        
        for i in range(target_list.count()):
            item = target_list.item(i)
            uid = item.data(Qt.ItemDataRole.UserRole)
            if not uid: continue
            
            # Find in scene
            scene_item = next((x for x in self.scene.items() if getattr(x, 'uid', None) == uid), None)
            if scene_item:
                # Update status based on target list
                scene_item.is_global = target_is_global
                
                # Visual update
                if target_is_global and hasattr(scene_item, 'base_color'):
                     scene_item.base_color = QColor(0, 255, 0, 80) # Global color
                elif not target_is_global and hasattr(scene_item, 'base_color') and scene_item.item_type == 'blur':
                     scene_item.base_color = QColor(0, 0, 0, 150) # Local color
                
                scene_item.update()

        # Update lists to reflect changes cleanly
        self.refresh_layer_list()

    def on_layer_clicked(self, item):
        # Sync selection from list to scene
        uid = item.data(Qt.ItemDataRole.UserRole)
        if not uid: return
        
        # Find item by UUID
        scene_item = next((x for x in self.scene.items() if getattr(x, 'uid', None) == uid), None)
        
        if scene_item:
            self.scene.clearSelection()
            scene_item.setSelected(True)

    def move_layer_to_global(self):
        """Move selected step layer to global"""
        item = self.step_layer_list.currentItem()
        if not item:
            return
        
        idx = self.step_layer_list.row(item)
        scene_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and i.item_type != 'click' and not getattr(i, 'is_global', False)]
        scene_items.sort(key=lambda x: x.zValue())
        
        if idx < len(scene_items):
            scene_item = scene_items[idx]
            scene_item.is_global = True
            
            # Update visual indicator for blur items
            if hasattr(scene_item, 'base_color'):
                scene_item.base_color = QColor(0, 255, 0, 80)
                scene_item.update()
            
            self.refresh_layer_list()
            QMessageBox.information(self, "Erfolg", "Ebene wurde zu Global verschoben!")

    def move_layer_to_step(self):
        """Move selected global layer to step"""
        item = self.global_layer_list.currentItem()
        if not item:
            return
        
        idx = self.global_layer_list.row(item)
        scene_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and getattr(i, 'is_global', False)]
        scene_items.sort(key=lambda x: x.zValue())
        
        if idx < len(scene_items):
            scene_item = scene_items[idx]
            scene_item.is_global = False
            
            # Update visual indicator for blur items
            if hasattr(scene_item, 'base_color') and scene_item.item_type == 'blur':
                scene_item.base_color = QColor(0, 0, 0, 150)
                scene_item.update()
            
            self.refresh_layer_list()
            QMessageBox.information(self, "Erfolg", "Ebene wurde zu Bild verschoben!")

    def on_layer_moved_to_global(self):
        """Handle drag-drop between lists"""
        # This would need more complex implementation
        pass

    def refresh_layer_list(self):
        self.step_layer_list.clear()
        self.global_layer_list.clear()
        
        # Separate items by global status
        step_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and not getattr(i, 'is_global', False)]
        global_items = [i for i in self.scene.items() if hasattr(i, 'item_type') and getattr(i, 'is_global', False)]
        
        step_items.sort(key=lambda x: x.zValue())
        global_items.sort(key=lambda x: x.zValue())
        
        # Populate step layers
        for i in step_items:
            name = f"{i.item_type.upper()}"
            if hasattr(i, 'toPlainText'):  # Text item
                name += f": {i.toPlainText()[:20]}"
            li = QListWidgetItem(name)
            li.setData(Qt.ItemDataRole.UserRole, getattr(i, 'uid', None))
            self.step_layer_list.addItem(li)
        
        # Populate global layers
        for i in global_items:
            name = f"🌍 {i.item_type.upper()}"
            if hasattr(i, 'toPlainText'):  # Text item
                name += f": {i.toPlainText()[:20]}"
            li = QListWidgetItem(name)
            li.setData(Qt.ItemDataRole.UserRole, getattr(i, 'uid', None))
            self.global_layer_list.addItem(li)

    def save_current_state(self):
        """Extract items from scene and sync back to data model"""
        if self.current_idx < 0 or not self.steps: return
        
        new_layers = []
        new_globals = []
        
        # Offset handling (add crop offset back to coordinates for storage)
        ox = getattr(self, 'current_offset_x', 0)
        oy = getattr(self, 'current_offset_y', 0)
        
        for item in self.scene.items():
            if not hasattr(item, 'item_type'): continue
            
            # Extract data based on type
            data = {}
            uid = getattr(item, 'uid', None)
            
    
            if item.item_type == 'blur':
                r = item.sceneBoundingRect()
                # Reverse offset (add crop start position)
                data = {
                    'coords': (int(r.left()+ox), int(r.top()+oy), int(r.right()+ox), int(r.bottom()+oy)),
                    'strength': getattr(item, 'blur_strength', 40)
                }
                layer = Layer('blur', data, item.label, getattr(item, 'is_global', False), uid)
                if getattr(item, 'is_global', False): new_globals.append(layer)
                else: new_layers.append(layer)
                
            elif item.item_type == 'zoom':
                r = item.box_rect.translated(item.pos())
                t = item.target
                col = item.border_color
                data = {'x': int(r.left()+ox), 'y': int(r.top()+oy), 'size': int(r.width()), 
                        'target_x': int(t.x()+ox), 'target_y': int(t.y()+oy),
                        'color': (col.red(), col.green(), col.blue())}
                layer = Layer('zoom', data, "Zoom", getattr(item, 'is_global', False), uid)
                if getattr(item, 'is_global', False): new_globals.append(layer)
                else: new_layers.append(layer)

            elif item.item_type == 'arrow':
                s, e = item.start + item.pos(), item.end + item.pos() # Add item pos if moved
                col = item.color
                data = {'sx': int(s.x()+ox), 'sy': int(s.y()+oy), 'ex': int(e.x()+ox), 'ey': int(e.y()+oy),
                        'color': (col.red(), col.green(), col.blue()), 'width': item.width}
                layer = Layer('arrow', data, "Arrow", getattr(item, 'is_global', False), uid)
                if getattr(item, 'is_global', False): new_globals.append(layer)
                else: new_layers.append(layer)

            elif item.item_type == 'icon':
                p = item.pos()
                r = item.rect()
                col = item.icon_color
                data = {'x': int(p.x()+ox), 'y': int(p.y()+oy), 'w': int(r.width()), 'h': int(r.height()), 
                        'type': item.icon_type, 'color': (col.red(), col.green(), col.blue())}
                layer = Layer('icon', data, item.icon_type.title(), getattr(item, 'is_global', False), uid)
                if getattr(item, 'is_global', False): new_globals.append(layer)
                else: new_layers.append(layer)

            elif item.item_type == 'infobox':
                try:
                    r = item.box_rect.translated(item.pos())
                    t = item.target
                    col = item.border_color
                    f = item.font()
                    bg = item.bg_color
                    data = {'x': int(r.left()+ox), 'y': int(r.top()+oy), 'w': int(r.width()), 'h': int(r.height()),
                            'target_x': int(t.x()+ox), 'target_y': int(t.y()+oy), 'text': item.text,
                            'color': (col.red(), col.green(), col.blue()),
                            'bg_color': (bg.red(), bg.green(), bg.blue(), bg.alpha()),
                            'text_color': (item.text_color.red(), item.text_color.green(), item.text_color.blue()),
                            'border_width': item.border_width,
                            'corner_radius': item.corner_radius,
                            'h_align': item.h_align,
                            'v_align': item.v_align,
                            'font': {
                                'family': f.family(),
                                'size': f.pointSize() if f.pointSize() > 0 else 12,
                                'bold': f.bold(),
                                'italic': f.italic(),
                                'underline': f.underline()
                            }}
                    layer = Layer('infobox', data, "InfoBox", getattr(item, 'is_global', False), uid)
                    if getattr(item, 'is_global', False): new_globals.append(layer)
                    else: new_layers.append(layer)
                except (RuntimeError, AttributeError):
                    # Item was deleted or is being deleted, skip it
                    pass
                
            elif item.item_type == 'spotlight':
                 # Use rect() which is local and position
                 r = item.rect().translated(item.pos())
                 col = item.color
                 data = {
                     'x': int(r.left()+ox), 'y': int(r.top()+oy),
                     'w': int(r.width()), 'h': int(r.height()),
                     'opacity': getattr(item, 'dim_opacity', 0.6),
                     'shape': getattr(item, 'spotlight_shape', 'rect'),
                     'color': (col.red(), col.green(), col.blue())
                 }
                 layer = Layer('spotlight', data, "Spotlight", getattr(item, 'is_global', False), uid)
                 if getattr(item, 'is_global', False): new_globals.append(layer)
                 else: new_layers.append(layer)

            elif item.item_type == 'text':
                try:
                    p = item.pos()
                    f = item.font()
                    data = {
                        'text': item.toPlainText(), 
                        'x': int(p.x()+ox), 
                        'y': int(p.y()+oy), 
                        'color': (item.defaultTextColor().blue(), item.defaultTextColor().green(), item.defaultTextColor().red()),
                        'font': {
                            'family': f.family(),
                            'size': f.pointSize() if f.pointSize() > 0 else 18,
                            'bold': f.bold(),
                            'italic': f.italic(),
                            'underline': f.underline()
                        }
                    }
                    layer = Layer('text', data, "Text", getattr(item, 'is_global', False), uid)
                    if getattr(item, 'is_global', False): new_globals.append(layer)
                    else: new_layers.append(layer)
                except (RuntimeError, AttributeError):
                    # Item was deleted or is being deleted, skip it
                    pass
                
            elif item.item_type == 'click':
                # Preserve click from current step but update position if moved
                p = item.boundingRect().center() + item.pos()
                self.steps[self.current_idx].x = int(p.x()+ox)
                self.steps[self.current_idx].y = int(p.y()+oy)
                # Update the click layer separately or just keep it
        
        self.steps[self.current_idx].layers = [l for l in self.steps[self.current_idx].layers if l.type == 'click']
        for l in self.steps[self.current_idx].layers:
            if l.type == 'click':
                l.data['x'] = self.steps[self.current_idx].x
                l.data['y'] = self.steps[self.current_idx].y
        self.steps[self.current_idx].layers.extend(new_layers)
        
        if new_globals:
            # Smart update using UUIDs
            existing_map = {l.uid: l for l in self.global_layers}
            
            for ng in new_globals:
                if ng.uid in existing_map:
                    # Check if actually modified to avoid unnecessary updates? 
                    # For now just update data and label
                    existing_map[ng.uid].data = ng.data
                    existing_map[ng.uid].label = ng.label
                else:
                    self.global_layers.append(ng)
                    existing_map[ng.uid] = ng

    # ==================== UNDO SYSTEM ====================
    def push_undo(self):
        """Save current state to undo stack"""
        # First sync current scene to model to catch latest changes
        self.save_current_state()
        snapshot = self.capture_snapshot()
        self.undo_stack.append(snapshot)
        # Limit stack size
        if len(self.undo_stack) > 20:
            self.undo_stack.pop(0)
            
    def undo(self):
        """Restore last state"""
        if not self.undo_stack:
            return
            
        snapshot = self.undo_stack.pop()
        self.restore_snapshot(snapshot)
        QMessageBox.information(self, "Undo", "Schritt zurückgesetzt.")

    def capture_snapshot(self):
        """Create a deep copy of the project metadata"""
        # Copy global layers
        globals_copy = []
        for l in self.global_layers:
            import copy
            globals_copy.append(Layer(l.type, copy.deepcopy(l.data), l.label, l.is_global, l.uid))
            
        # Copy steps structure (images are referenced, not copied)
        steps_copy = []
        for s in self.steps:
            # We assume Step class structure. 
            # Step(raw_img, x, y, label)
            import copy
            # We need to recreate the Step object to detach layer list, but keep image ref
            new_step = Step(s.raw_img, s.x, s.y, getattr(s, 'label', ""))
            
            # Manually copy layers
            new_step.layers = []
            for l in s.layers:
                new_step.layers.append(Layer(l.type, copy.deepcopy(l.data), l.label, l.is_global, l.uid))
            
            steps_copy.append(new_step)
            
        return {
            "global_layers": globals_copy,
            "steps": steps_copy,
            "global_crop": self.global_crop,
            "current_idx": self.current_idx
        }

    def restore_snapshot(self, snapshot):
        """Restore project from snapshot"""
        self.global_layers = snapshot["global_layers"]
        self.steps = snapshot["steps"]
        self.global_crop = snapshot["global_crop"]
        idx = snapshot["current_idx"]
        
        if idx >= len(self.steps): idx = len(self.steps) - 1
        
        self.update_thumbnails()
        self.load_step(idx)

    def add_icon_dialog(self):
        # Quick dialog to pick icon type
        items = ["check", "cross", "warn", "info", "star", "idea"]
        dlg = ModernDialog("Icon wählen", "Typ:", mode="list", default_text=items, parent=self)
        
        if dlg.exec():
             item_text = dlg.get_text()
             # Add to center of view
             center = self.scene.views()[0].mapToScene(self.scene.views()[0].rect().center())
             icon = IconItem(center, item_text)
             self.scene.addItem(icon)

    def set_border_color_arrow(self, item, color):
        item.color = color
        item.update()
        
    def set_icon_type(self, item, name):
         item.icon_type = name
         item.update()

    def update_properties(self):
        try:
            # Check if scene is valid or deleted
            if not getattr(self, 'scene', None): return
            
            # Helper to clear layout recursively
            def clear_layout(layout):
                if layout is None: return
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget:
                        widget.deleteLater()
                    elif item.layout():
                        clear_layout(item.layout())
            
            clear_layout(self.props_layout)
            
            items = self.scene.selectedItems()
            if not items:
                # Add spacer to push content up if needed, or just label
                lbl = QLabel("Keine Auswahl")
                lbl.setStyleSheet("color: #666; font-style: italic;")
                self.props_layout.addWidget(lbl)
                self.props_layout.addStretch()
                return
            
            # Additional safety: check if item is still in scene
            if not items[0].scene():
                lbl = QLabel("Keine Auswahl")
                lbl.setStyleSheet("color: #666; font-style: italic;")
                self.props_layout.addWidget(lbl)
                self.props_layout.addStretch()
                return
            
            item = items[0]
            
            # Header
            header = QLabel(f"{type(item).__name__}")
            header.setStyleSheet("font-weight: bold; color: #0078d4; font-size: 14px; margin-bottom: 10px;")
            self.props_layout.addWidget(header)
            
            # Debug: Print item type
            print(f"DEBUG: Selected item type: {type(item).__name__}, has item_type: {hasattr(item, 'item_type')}, item_type value: {getattr(item, 'item_type', 'N/A')}")
            
            if isinstance(item, BlurItem):
                # BLUR CONTROLS
                self.props_layout.addWidget(QLabel("Unschärfe-Stärke:"))
                
                slider = QSlider(Qt.Orientation.Horizontal)
                slider.setRange(1, 150)
                slider.setValue(item.blur_strength)
                slider.valueChanged.connect(lambda v: self.update_blur_strength(item, v))
                self.props_layout.addWidget(slider)

            if isinstance(item, ArrowItem):
                 self.props_layout.addWidget(QLabel("Pfeilfarbe:"))
                 self.create_color_palette(item, lambda i, c: self.set_border_color_arrow(i, c), item.color)
                 
                 self.props_layout.addWidget(QLabel("Dicke:"))
                 spin = QSpinBox()
                 spin.setRange(1, 20)
                 spin.setValue(item.width)
                 spin.valueChanged.connect(lambda v: setattr(item, 'width', v) or item.update())
                 self.props_layout.addWidget(spin)

            if isinstance(item, IconItem):
                self.props_layout.addWidget(QLabel("Icon Farbe:"))
                self.create_color_palette(item, self.set_icon_color, item.icon_color)
                
                self.props_layout.addWidget(QLabel("Icon Typ:"))
                combo = QComboBox()
                combo.addItems(IconItem.ICONS.keys())
                combo.setCurrentText(item.icon_type)
                combo.currentTextChanged.connect(lambda t: self.set_icon_type(item, t))
                self.props_layout.addWidget(combo)
                
                self.props_layout.addWidget(QLabel("Größe:"))
                spin = QSpinBox()
                spin.setRange(20, 500)
                r = item.rect()
                spin.setValue(int(r.width()))
                spin.valueChanged.connect(lambda v: item.setRect(QRectF(0, 0, v, v)) or item.update())
                self.props_layout.addWidget(spin)

            if isinstance(item, SpotlightItem):
                self.props_layout.addWidget(QLabel("🎥 Spotlight Fokus"))
                
                self.props_layout.addWidget(QLabel("Deckkraft:"))
                slider = QSlider(Qt.Orientation.Horizontal)
                slider.setRange(0, 100)
                slider.setValue(int(item.dim_opacity * 100))
                slider.valueChanged.connect(lambda v: setattr(item, 'dim_opacity', v/100.0) or item.update())
                self.props_layout.addWidget(slider)
                
                self.props_layout.addWidget(QLabel("Form:"))
                combo = QComboBox()
                combo.addItems(['rect', 'ellipse'])
                combo.setCurrentText(item.spotlight_shape)
                combo.currentTextChanged.connect(lambda t: setattr(item, 'spotlight_shape', t) or item.update())
                self.props_layout.addWidget(combo)
                
                self.props_layout.addWidget(QLabel("Größe (BxH):"))
                row = QHBoxLayout()
                spin_w = QSpinBox()
                spin_w.setRange(10, 5000)
                spin_w.setValue(int(item.rect().width()))
                spin_w.valueChanged.connect(lambda v: item.setRect(0, 0, v, item.rect().height()) or item.update())
                row.addWidget(spin_w)
                
                spin_h = QSpinBox()
                spin_h.setRange(10, 5000)
                spin_h.setValue(int(item.rect().height()))
                spin_h.valueChanged.connect(lambda v: item.setRect(0, 0, item.rect().width(), v) or item.update())
                row.addWidget(spin_h)
                self.props_layout.addLayout(row)

            if isinstance(item, ClickMarkerItem):
                # GLOBAL MARKER SETTINGS (affects ALL markers)
                settings = ClickMarkerSettings()
                
                self.props_layout.addWidget(QLabel("<b>🌍 Globale Marker-Einstellungen</b>"))
                self.props_layout.addWidget(QLabel("<small>(Gilt für alle Klickmarker)</small>"))
                self.props_layout.addSpacing(10)
                
                # Color
                self.props_layout.addWidget(QLabel("Farbe:"))
                
                # Color preset buttons
                presets = [QColor(0, 168, 255), QColor(255, 80, 80), QColor(80, 255, 80), 
                          QColor(255, 255, 255), QColor(255, 165, 0), QColor(255, 0, 255)]
                
                preset_layout = QHBoxLayout()
                preset_layout.setSpacing(5)
                
                for col in presets:
                    btn = QPushButton()
                    btn.setFixedSize(24, 24)
                    btn.setStyleSheet(f"background-color: {col.name()}; border: 1px solid #555; border-radius: 12px;")
                    btn.clicked.connect(lambda checked, c=col: self.set_marker_color_global(item, c))
                    preset_layout.addWidget(btn)
                
                preset_layout.addStretch()
                self.props_layout.addLayout(preset_layout)
                
                # Custom color and Transparent buttons
                btn_layout = QHBoxLayout()
                btn_custom = QPushButton("🎨 Custom...")
                btn_custom.clicked.connect(lambda: self.choose_custom_marker_color())
                btn_layout.addWidget(btn_custom)
                
                btn_transparent = QPushButton("⭕ Transparent")
                btn_transparent.clicked.connect(lambda: self.set_marker_transparent())
                btn_layout.addWidget(btn_transparent)
                self.props_layout.addLayout(btn_layout)
                
                # Text Color
                self.props_layout.addSpacing(10)
                self.props_layout.addWidget(QLabel("Schriftfarbe:"))
                text_presets = [QColor(255, 255, 255), QColor(0, 0, 0)]
                text_layout = QHBoxLayout()
                for col in text_presets:
                    btn = QPushButton("Weiß" if col.name() == "#ffffff" else "Schwarz")
                    btn.setStyleSheet(f"background-color: {col.name()}; color: {'black' if col.name() == '#ffffff' else 'white'}; border: 1px solid #555;")
                    btn.clicked.connect(lambda checked, c=col: self.set_marker_text_color_global(c))
                    text_layout.addWidget(btn)
                self.props_layout.addLayout(text_layout)
                
                # Size
                self.props_layout.addWidget(QLabel(f"Größe: {settings.size}px"))
                size_slider = QSlider(Qt.Orientation.Horizontal)
                size_slider.setRange(20, 80)
                size_slider.setValue(settings.size)
                size_slider.valueChanged.connect(lambda v: self.update_marker_size_global(v))
                self.props_layout.addWidget(size_slider)
                
                # Border Width
                self.props_layout.addWidget(QLabel(f"Randstärke: {settings.border_width}px"))
                border_slider = QSlider(Qt.Orientation.Horizontal)
                border_slider.setRange(1, 10)
                border_slider.setValue(settings.border_width)
                border_slider.valueChanged.connect(lambda v: self.update_marker_border_global(v))
                self.props_layout.addWidget(border_slider)
                
                # Number Size
                self.props_layout.addWidget(QLabel(f"Schriftgröße: {settings.number_size}pt"))
                number_slider = QSlider(Qt.Orientation.Horizontal)
                number_slider.setRange(8, 32)
                number_slider.setValue(settings.number_size)
                number_slider.valueChanged.connect(lambda v: self.update_marker_number_size_global(v))
                self.props_layout.addWidget(number_slider)
                
                # Glow Effect
                glow_check = QCheckBox("Leuchteffekt anzeigen")
                glow_check.setChecked(settings.show_glow)
                glow_check.stateChanged.connect(lambda s: self.update_marker_glow_global(s == Qt.CheckState.Checked.value))
                self.props_layout.addWidget(glow_check)

            if isinstance(item, (ZoomItem, InfoBoxItem)):
                # BOX COLOR CONTROLS
                self.props_layout.addWidget(QLabel("Randfarbe:"))
                if not hasattr(item, 'border_color'): item.border_color = QColor(255, 255, 255)
                self.create_color_palette(item, self.set_border_color, item.border_color)

            if isinstance(item, InfoBoxItem):
                self.props_layout.addSpacing(10)
                self.props_layout.addWidget(QLabel("Hintergrundfarbe:"))
                self.create_color_palette(item, self.set_bg_color, item.bg_color, allow_transparent=True)
                
                self.props_layout.addSpacing(10)
                self.props_layout.addWidget(QLabel("Randstärke / Abrundung:"))
                row = QHBoxLayout()
                spin_bw = QSpinBox()
                spin_bw.setRange(0, 20)
                spin_bw.setValue(item.border_width)
                spin_bw.setToolTip("Randstärke")
                spin_bw.valueChanged.connect(lambda v: setattr(item, 'border_width', v) or item.update())
                row.addWidget(spin_bw)
                
                spin_cr = QSpinBox()
                spin_cr.setRange(0, 100)
                spin_cr.setValue(item.corner_radius)
                spin_cr.setToolTip("Abrundung")
                spin_cr.valueChanged.connect(lambda v: setattr(item, 'corner_radius', v) or item.update())
                row.addWidget(spin_cr)
                self.props_layout.addLayout(row)

                # Text Alignment
                self.props_layout.addSpacing(10)
                self.props_layout.addWidget(QLabel("Textausrichtung:"))
                
                h_layout = QHBoxLayout()
                for align in ['left', 'center', 'right']:
                    icon = "⬅️" if align == 'left' else "⏺️" if align == 'center' else "➡️"
                    btn = QPushButton(icon)
                    btn.setCheckable(True)
                    btn.setChecked(item.h_align == align)
                    btn.setToolTip(f"Horizontal: {align}")
                    btn.clicked.connect(lambda checked, a=align: setattr(item, 'h_align', a) or item.update() or self.update_properties())
                    h_layout.addWidget(btn)
                self.props_layout.addLayout(h_layout)

                v_layout = QHBoxLayout()
                for align in ['top', 'center', 'bottom']:
                    icon = "🔝" if align == 'top' else "↔️" if align == 'center' else "⬇️"
                    btn = QPushButton(icon)
                    btn.setCheckable(True)
                    btn.setChecked(item.v_align == align)
                    btn.setToolTip(f"Vertikal: {align}")
                    btn.clicked.connect(lambda checked, a=align: setattr(item, 'v_align', a) or item.update() or self.update_properties())
                    v_layout.addWidget(btn)
                self.props_layout.addLayout(v_layout)

                self.props_layout.addSpacing(10)
                self.props_layout.addWidget(QLabel("Schriftfarbe:"))
                self.create_color_palette(item, self.set_text_color, item.text_color)
                self.create_font_controls(item)

            if isinstance(item, EditableTextItem):
                # TEXT FORMATTING CONTROLS
                self.props_layout.addSpacing(5)
                self.props_layout.addWidget(QLabel("Schriftfarbe:"))
                self.create_color_palette(item, self.set_text_color, item.defaultTextColor() if hasattr(item, 'defaultTextColor') else item.text_color)
                self.create_font_controls(item)
                
            # Add general info (Optional)
            if hasattr(item, 'is_global'):
                status = "🌍 Global" if item.is_global else "📄 Lokal"
                self.props_layout.addWidget(QLabel(f"Status: {status}"))
            
            # Add spacer to push content to top
            self.props_layout.addStretch()
                
            
            # Add spacer to push content to top
            self.props_layout.addStretch()
        except RuntimeError:
            pass # Scene deleted during update
            
    def update_blur_strength(self, item, val):
        item.blur_strength = val
        item.update() # Trigger repaint
    
    def refresh_all_markers(self):
        """Helper to refresh all click markers in the scene"""
        for item in self.scene.items():
            if hasattr(item, 'item_type') and item.item_type == 'click':
                item.prepareGeometryChange()
                item.update()
        self.scene.update()
        self.view.viewport().update()
    
    def set_marker_color_global(self, item, color):
        """Update marker color globally"""
        settings = ClickMarkerSettings()
        settings.color = color
        settings.save()
        self.refresh_all_markers()
        self.update_properties()  # Refresh panel
    
    def choose_custom_marker_color(self):
        """Open color picker for custom marker color"""
        settings = ClickMarkerSettings()
        color = QColorDialog.getColor(settings.color, self, "Marker-Farbe wählen")
        if color.isValid():
            settings.color = color
            settings.save()
            self.refresh_all_markers()
            self.update_properties()
    
    def set_marker_transparent(self):
        """Set marker fill to transparent (only border visible)"""
        settings = ClickMarkerSettings()
        # Use transparent color (alpha = 0)
        settings.color = QColor(0, 0, 0, 0)
        settings.save()
        self.refresh_all_markers()
        self.update_properties()
    
    def set_marker_text_color_global(self, color):
        """Update marker text color globally"""
        settings = ClickMarkerSettings()
        settings.text_color = color
        settings.save()
        self.refresh_all_markers()
        self.update_properties()
    
    def update_marker_size_global(self, value):
        """Update marker size globally"""
        settings = ClickMarkerSettings()
        settings.size = value
        settings.save()
        self.refresh_all_markers()
        self.update_properties()  # Refresh label
    
    def update_marker_border_global(self, value):
        """Update marker border width globally"""
        settings = ClickMarkerSettings()
        settings.border_width = value
        settings.save()
        self.refresh_all_markers()
        self.update_properties()  # Refresh label
    
    def update_marker_number_size_global(self, value):
        """Update marker number size globally"""
        settings = ClickMarkerSettings()
        settings.number_size = value
        settings.save()
        self.refresh_all_markers()
        self.update_properties()  # Refresh label
    
    def update_marker_glow_global(self, enabled):
        """Update marker glow effect globally"""
        settings = ClickMarkerSettings()
        settings.show_glow = enabled
        settings.save()
        self.refresh_all_markers()
        
    def create_color_palette(self, item, setter_func, current_color, allow_transparent=False):
         presets = [QColor(255, 255, 255), QColor(255, 80, 80), QColor(80, 255, 80), QColor(80, 150, 255),
                    QColor(255, 165, 0), QColor(255, 255, 0), QColor(255, 0, 255)]
         
         preset_layout = QHBoxLayout()
         preset_layout.setSpacing(8)
         
         for col in presets:
             btn = QPushButton()
             btn.setFixedSize(32, 32)
             btn.setStyleSheet(f"""
                 QPushButton {{
                     background-color: {col.name()}; 
                     border: 2px solid #cbd5e1; 
                     border-radius: 16px;
                 }}
                 QPushButton:hover {{
                     border: 3px solid #3b82f6;
                     transform: scale(1.1);
                 }}
             """)
             btn.clicked.connect(lambda checked, c=col: setter_func(item, c))
             preset_layout.addWidget(btn)
             
         if allow_transparent:
             btn_trans = QPushButton("❌")
             btn_trans.setFixedSize(32, 32)
             btn_trans.setToolTip("Transparent")
             btn_trans.setStyleSheet("""
                 QPushButton {
                     background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                         stop:0 #fee2e2, stop:1 #fca5a5);
                     border: 2px solid #cbd5e1; 
                     border-radius: 16px; 
                     color: #dc2626;
                     font-weight: 700;
                 }
                 QPushButton:hover {
                     border: 3px solid #dc2626;
                 }
             """)
             btn_trans.clicked.connect(lambda: setter_func(item, QColor(0,0,0,0)))
             preset_layout.addWidget(btn_trans)

         preset_layout.addStretch()
         self.props_layout.addLayout(preset_layout)
         
         btn_custom = QPushButton("Custom...")
         btn_custom.setStyleSheet("""
             QPushButton {
                 background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                     stop:0 #f8fafc, stop:1 #e2e8f0);
                 color: #334155;
                 border: 1px solid #cbd5e1;
                 padding: 8px 16px;
                 border-radius: 6px;
                 font-weight: 600;
             }
             QPushButton:hover {
                 background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                     stop:0 #e2e8f0, stop:1 #cbd5e1);
                 border-color: #94a3b8;
             }
         """)
         btn_custom.clicked.connect(lambda: self.change_generic_color(item, setter_func, current_color))
         self.props_layout.addWidget(btn_custom)
         
    def change_generic_color(self, item, setter, current):
        color = QColorDialog.getColor(current)
        if color.isValid():
             setter(item, color)

    def set_border_color(self, item, color):
        item.border_color = color
        item.update()

    def set_bg_color(self, item, color):
        item.bg_color = color
        item.update()

    def set_text_color(self, item, color):
        if hasattr(item, 'setDefaultTextColor'):
            item.setDefaultTextColor(color)
        else:
            item.text_color = color
            item.update()

    def set_icon_color(self, item, color):
        item.icon_color = color
        item.update()

    def create_font_controls(self, item):
        try:
            self.props_layout.addSpacing(10)
            self.props_layout.addWidget(QLabel("Schriftart:"))
            
            # 1. Font Family
            font_combo = QFontComboBox()
            font_combo.setCurrentFont(item.font())
            font_combo.currentFontChanged.connect(lambda f: item.update_font(family=f.family()))
            self.props_layout.addWidget(font_combo)
            
            # 2. Size & Styles Row
            row = QHBoxLayout()
            
            # Size
            spin_size = QSpinBox()
            spin_size.setRange(8, 200)
            try:
                sz = item.font().pointSize()
                if sz <= 0: sz = item.font().pixelSize()
                spin_size.setValue(sz if sz > 0 else 18) # Default to 18 if invalid
            except (RuntimeError, AttributeError):
                spin_size.setValue(18)
            spin_size.valueChanged.connect(lambda s: item.update_font(size=s))
            row.addWidget(spin_size)
            
            # Bold
            btn_bold = QPushButton("B")
            btn_bold.setCheckable(True)
            btn_bold.setFixedWidth(30)
            # Handle cases where font might be a property or method
            try:
                f = item.font()
                btn_bold.setChecked(f.bold())
            except (RuntimeError, AttributeError):
                btn_bold.setChecked(False)
            btn_bold.setStyleSheet("font-weight: bold;")
            btn_bold.clicked.connect(lambda c: item.update_font(bold=c))
            row.addWidget(btn_bold)
            
            # Italic
            btn_italic = QPushButton("I")
            btn_italic.setCheckable(True)
            btn_italic.setFixedWidth(30)
            try:
                btn_italic.setChecked(f.italic())
            except:
                btn_italic.setChecked(False)
            btn_italic.setStyleSheet("font-style: italic;")
            btn_italic.clicked.connect(lambda c: item.update_font(italic=c))
            row.addWidget(btn_italic)
            
            # Underline
            btn_underline = QPushButton("U")
            btn_underline.setCheckable(True)
            btn_underline.setFixedWidth(30)
            try:
                btn_underline.setChecked(f.underline())
            except:
                btn_underline.setChecked(False)
            btn_underline.setStyleSheet("text-decoration: underline;")
            btn_underline.clicked.connect(lambda c: item.update_font(underline=c))
            row.addWidget(btn_underline)
            
            self.props_layout.addLayout(row)
        except (RuntimeError, AttributeError):
            pass # Item was deleted during update

    def change_border_color(self, item, btn=None):
        color = QColorDialog.getColor(item.border_color)
        if color.isValid():
            self.set_border_color(item, color)
            if btn:
                # Update btn style if it was passed (though generic palette handles this now mostly)
                btn.setStyleSheet(f"background-color: {color.name()}; color: {'black' if color.lightness() > 128 else 'white'}")
        
    def change_text_color(self, item):
        color = QColorDialog.getColor(item.defaultTextColor())
        if color.isValid():
            item.setDefaultTextColor(color)

    def save_description(self):
        """Save description text to current step"""
        if self.current_idx is not None and 0 <= self.current_idx < len(self.steps):
            self.steps[self.current_idx].description = self.txt_description.toPlainText()

    def update_thumbnails(self):
        self.thumb_list.clear()
        for i, s in enumerate(self.steps):
            # Create thumbnail with number
            if s.raw_img is not None:
                try:
                    h, w = s.raw_img.shape[:2]
                    thumb_h = 64
                    if h > 0:
                        thumb_w = int(w * thumb_h / h)
                        small = cv2.resize(s.raw_img, (thumb_w, thumb_h))
                        
                        # Add "#1" Badge background
                        cv2.rectangle(small, (0, 0), (28, 22), (20, 20, 20), -1)
                        # Add Number
                        cv2.putText(small, str(i+1), (5, 16), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (255, 255, 255), 1, cv2.LINE_AA)
                        
                        rgb = cv2.cvtColor(small, cv2.COLOR_BGR2RGB)
                        qimg = QImage(rgb.data, thumb_w, thumb_h, thumb_w*3, QImage.Format.Format_RGB888)
                        icon = QIcon(QPixmap.fromImage(qimg))
                        
                        item = QListWidgetItem(icon, f"Schritt {i+1}")
                        self.thumb_list.addItem(item)
                    else:
                        self.thumb_list.addItem(f"Schritt {i+1}")
                except Exception as e:
                    print(f"Thumb error: {e}")
                    self.thumb_list.addItem(f"Schritt {i+1}")
            else:
                self.thumb_list.addItem(f"Schritt {i+1}")

    def load_step(self, idx):
        if idx < 0 or idx >= len(self.steps): return
        self.current_idx = idx
        
        self.scene.clear()
        step = self.steps[idx]
        
        # 1. Background Image - Handle Global Crop
        img = step.raw_img
        
        # Calculate offsets based on crop
        offset_x, offset_y = 0, 0
        if self.global_crop:
            x1, y1, x2, y2 = self.global_crop
            h, w = img.shape[:2]
            # Clip crop rect to image bounds
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(w, x2), min(h, y2)
            
            if x2 > x1 and y2 > y1:
                img = img[y1:y2, x1:x2] # Crop the image logic
                offset_x, offset_y = x1, y1
                
        self.current_offset_x = offset_x
        self.current_offset_y = offset_y
        
        h, w = img.shape[:2]
        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        qimg = QImage(rgb.data, w, h, w*3, QImage.Format.Format_RGB888)
        self.current_pixmap = QPixmap.fromImage(qimg)
        
        bg_item = QGraphicsPixmapItem(self.current_pixmap)
        bg_item.setZValue(-100)
        self.scene.addItem(bg_item)
        
        # Helper helpers
        def shift_rect(c):
            # Transform global coordinates to cropped local coordinates
            return QRectF(c[0]-offset_x, c[1]-offset_y, c[2]-c[0], c[3]-c[1])
            
        def shift_point(x, y):
            return QPointF(x-offset_x, y-offset_y)
        
        # 2. Existing Layers
        
        # Global Layers (persist across steps)
        for gl in self.global_layers:
            # We need to render them relative to the crop
            # gl.data coordinates are always relative to the FULL ORIGINAL IMAGE
            
            if gl.type == 'blur':
                c = gl.data['coords']
                strength = gl.data.get('strength', 40)
                # Only add if visible in crop? QGraphicsView handles clipping anyway.
                item = BlurItem(shift_rect(c), True, self.current_pixmap, uid=gl.uid)
                item.blur_strength = strength
                self.scene.addItem(item)
            elif gl.type == 'zoom':
                d = gl.data
                rect = QRectF(d['x']-offset_x, d['y']-offset_y, d['size'], d['size'])
                target = shift_point(d['target_x'], d['target_y'])
                item = ZoomItem(rect, target, self.current_pixmap, is_global=True, uid=gl.uid)
                if 'color' in d: item.border_color = QColor(*d['color'])
                self.scene.addItem(item)
            elif gl.type == 'arrow':
                d = gl.data
                item = ArrowItem(shift_point(d['sx'], d['sy']), shift_point(d['ex'], d['ey']),
                               QColor(*d['color']), d.get('width', 4), True, gl.uid)
                self.scene.addItem(item)
            elif gl.type == 'icon':
                d = gl.data
                item = IconItem(shift_point(d['x'], d['y']), d['type'], d.get('w', d.get('size', 60)), QColor(*d.get('color', (255,0,0))), True, gl.uid)
                if 'w' in d and 'h' in d:
                    item.setRect(QRectF(0,0,d['w'],d['h']))
                self.scene.addItem(item)
            elif gl.type == 'infobox':
                d = gl.data
                rect = QRectF(d['x']-offset_x, d['y']-offset_y, d['w'], d['h'])
                target = shift_point(d['target_x'], d['target_y'])
                item = InfoBoxItem(rect, target, d['text'], is_global=True, uid=gl.uid)
                if 'color' in d: item.border_color = QColor(*d['color'])
                if 'bg_color' in d: item.bg_color = QColor(*d['bg_color'])
                if 'text_color' in d: item.text_color = QColor(*d['text_color'])
                if 'border_width' in d: item.border_width = d['border_width']
                if 'corner_radius' in d: item.corner_radius = d['corner_radius']
                if 'h_align' in d: item.h_align = d['h_align']
                if 'v_align' in d: item.v_align = d['v_align']
                if 'font' in d:
                    f = d['font']
                    item.update_font(f.get('family'), f.get('size'), f.get('bold'), f.get('italic'), f.get('underline'))
                self.scene.addItem(item)
            elif gl.type == 'spotlight':
                d = gl.data
                item = SpotlightItem(d['x']-offset_x, d['y']-offset_y, d['w'], d['h'], is_global=True, uid=gl.uid)
                item.dim_opacity = d.get('opacity', 0.6)
                item.spotlight_shape = d.get('shape', 'rect')
                item.color = QColor(*d.get('color', (0,0,0)))
                self.scene.addItem(item)
            elif gl.type == 'text':
                d = gl.data
                color = QColor(d['color'][2], d['color'][1], d['color'][0])
                item = EditableTextItem(d['text'], color, is_global=True, uid=gl.uid)
                item.setPos(d['x']-offset_x, d['y']-offset_y)
                if 'font' in d:
                    f = d['font']
                    item.update_font(f.get('family'), f.get('size'), f.get('bold'), f.get('italic'), f.get('underline'))
                self.scene.addItem(item)
                
        # Step Layers
        for l in step.layers:
            if l.type == 'click':
                item = ClickMarkerItem(l.data['x']-offset_x, l.data['y']-offset_y, str(idx+1))
                self.scene.addItem(item)
            elif l.type == 'blur':
                c = l.data['coords']
                strength = l.data.get('strength', 40)
                item = BlurItem(shift_rect(c), False, self.current_pixmap)
                item.blur_strength = strength
                self.scene.addItem(item)
            elif l.type == 'zoom':
                d = l.data
                rect = QRectF(d['x']-offset_x, d['y']-offset_y, d['size'], d['size'])
                target = shift_point(d['target_x'], d['target_y'])
                item = ZoomItem(rect, target, self.current_pixmap, is_global=False)
                if 'color' in d: item.border_color = QColor(*d['color'])
                self.scene.addItem(item)
            elif l.type == 'arrow':
                d = l.data
                item = ArrowItem(shift_point(d['sx'], d['sy']), shift_point(d['ex'], d['ey']),
                               QColor(*d['color']), d.get('width', 4), False)
                self.scene.addItem(item)
            elif l.type == 'icon':
                d = l.data
                item = IconItem(shift_point(d['x'], d['y']), d['type'], d.get('w', d.get('size', 60)), QColor(*d.get('color', (255,0,0))), False)
                if 'w' in d and 'h' in d:
                    item.setRect(QRectF(0,0,d['w'],d['h']))
                self.scene.addItem(item)
            elif l.type == 'infobox':
                d = l.data
                rect = QRectF(d['x']-offset_x, d['y']-offset_y, d['w'], d['h'])
                target = shift_point(d['target_x'], d['target_y'])
                item = InfoBoxItem(rect, target, d['text'], is_global=False)
                if 'color' in d: item.border_color = QColor(*d['color'])
                if 'bg_color' in d: item.bg_color = QColor(*d['bg_color'])
                if 'text_color' in d: item.text_color = QColor(*d['text_color'])
                if 'border_width' in d: item.border_width = d['border_width']
                if 'corner_radius' in d: item.corner_radius = d['corner_radius']
                if 'h_align' in d: item.h_align = d['h_align']
                if 'v_align' in d: item.v_align = d['v_align']
                if 'font' in d:
                    f = d['font']
                    item.update_font(f.get('family'), f.get('size'), f.get('bold'), f.get('italic'), f.get('underline'))
                self.scene.addItem(item)
            elif l.type == 'spotlight':
                d = l.data
                item = SpotlightItem(d['x']-offset_x, d['y']-offset_y, d['w'], d['h'], is_global=False)
                item.dim_opacity = d.get('opacity', 0.6)
                item.spotlight_shape = d.get('shape', 'rect')
                item.color = QColor(*d.get('color', (0,0,0)))
                self.scene.addItem(item)
            elif l.type == 'text':
                d = l.data
                color = QColor(d['color'][2], d['color'][1], d['color'][0])
                item = EditableTextItem(d['text'], color, is_global=False)
                item.setPos(d['x']-offset_x, d['y']-offset_y)
                if 'font' in d:
                    f = d['font']
                    item.update_font(f.get('family'), f.get('size'), f.get('bold'), f.get('italic'), f.get('underline'))
                self.scene.addItem(item)
        
        # 3. Hardcoded Watermark (Non-Deletable) - Added once per step
        watermark = QGraphicsTextItem("Created with ClickStep Guide")
        watermark.setDefaultTextColor(QColor(255, 255, 255, 180)) # White, slighly transparent
        wm_font = QFont("Segoe UI", 10, QFont.Weight.Bold)
        watermark.setFont(wm_font)
        watermark.setZValue(9999) # Always on top
        # No flags = not selectable, not movable, not focusable
        watermark.setFlags(QGraphicsItem.GraphicsItemFlag(0))
        watermark.setAcceptedMouseButtons(Qt.MouseButton.NoButton) # Click through
        
        # Position top right
        pw = self.current_pixmap.width()
        # Estimate text width (approx 160px for this string at size 10)
        watermark.setPos(pw - 210, 20)
        self.scene.addItem(watermark)
        
        # Update Description Field
        if hasattr(self, 'txt_description'):
            self.txt_description.blockSignals(True)
            self.txt_description.setPlainText(step.description if step.description else "")
            self.txt_description.blockSignals(False)
            
        self.refresh_layer_list()
        self.update_properties()
        
        self.scene.setSceneRect(0, 0, w, h)
        
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(10, lambda: self.view.fitInView(self.scene.sceneRect(), Qt.AspectRatioMode.KeepAspectRatio))
        
        self.refresh_layer_list()

# ==================== RECORDER (unchanged) ====================

class Step:
    def __init__(self, raw_img, x, y, label):
        self.raw_img = raw_img
        self.x, self.y = x, y
        self.description = label
        self.layers = [Layer('click', {'x': x, 'y': y}, label)]

class RecordingSignal(QObject):
    click_detected = pyqtSignal(int, int, str, object)

class RecordingThread(QThread):
    def __init__(self):
        super().__init__()
        self.signals = RecordingSignal()
        self.is_running = False
        self.event_queue = queue.Queue()

    def run(self):
        self.is_running = True
        self.mouse_listener = mouse.Listener(on_click=self.on_click)
        self.mouse_listener.start()
        
        while self.is_running:
            try:
                # Wait for click event with short timeout to check is_running
                # Non-blocking allows us to exit clean
                x, y = self.event_queue.get(timeout=0.05)
                
                # Perform the delay and capture here in the thread, NOT in the hook
                time.sleep(0.12)
                
                raw = ImageGrab.grab()
                img = cv2.cvtColor(np.array(raw), cv2.COLOR_RGB2BGR)
                self.signals.click_detected.emit(x, y, "Click", img)
                
            except queue.Empty:
                continue
                
        self.mouse_listener.stop()

    def on_click(self, x, y, button, pressed):
        if not pressed: return
        # FAST: Just put into queue and return immediately to unblock system
        self.event_queue.put((x, y))

class ProRecorder(QMainWindow):
    hotkey_signal = pyqtSignal(str)
    
    def __init__(self):
        super().__init__()
        
        # Windows Taskbar Icon Fix
        if os.name == 'nt':
            try:
                ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("clickstep.guide.pro.1")
            except: pass
            
        self.setWindowTitle("ClickStep Guide")
        self.setFixedSize(1280, 720) # Fixed size to prevent background distortion (16:9)
        self.setWindowIcon(QIcon(resource_path("assets/icon.ico")))
        
        self.settings = AppSettings()
        self.apply_app_theme()
        
        self.steps = []
        self.global_layers = []
        self.global_crop = None
        self.is_recording = False
        self.overlay = RecordingOverlay() # Create overlay
        
        self.recording_thread = RecordingThread()
        self.recording_thread.signals.click_detected.connect(self.handle_click)
        
        # Connect hotkey signal to ensure thread-safe UI calls
        self.hotkey_signal.connect(self.handle_hotkey)
        
        self.setup_ui()
        self.setup_hotkeys()
        self.update_project_list()

    def closeEvent(self, event):
        """Cleanup thread before closing"""
        if hasattr(self, 'recording_thread'):
            self.recording_thread.is_running = False
            self.recording_thread.wait(500)
        event.accept()

    def get_project_dir(self):
        """Returns the project directory in Local AppData for Store compliance"""
        base = os.environ.get('LOCALAPPDATA', os.path.expanduser("~"))
        path = os.path.join(base, "ClickStepGuide", "projects")
        os.makedirs(path, exist_ok=True)
        return path

    def apply_app_theme(self):
        """Apply modern, premium theme with clean gradients and glassmorphism"""
        bg_path = resource_path("assets/background.png").replace("\\", "/")
        has_bg = os.path.exists(resource_path("assets/background.png"))
        
        # Modern Premium Theme (works for both light/dark, optimized for clean look)
        style = """
            QMainWindow { 
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #f8fafc, stop:0.5 #e0e7ff, stop:1 #f1f5f9);
            }
            
            QDialog { background-color: rgba(255, 255, 255, 0.95); }
            
            QWidget { 
                color: #1e293b; 
                font-family: 'Segoe UI', 'Inter', sans-serif; 
                font-size: 13px; 
            }
            
            /* Sidebar - Glassmorphism */
            #Sidebar { 
                background: rgba(255, 255, 255, 0.7);
                border-right: 1px solid rgba(226, 232, 240, 0.8);
                backdrop-filter: blur(10px);
            }
            
            #SidebarHeader { 
                color: #64748b; 
                font-weight: 700; 
                font-size: 11px; 
                letter-spacing: 2px; 
                margin-bottom: 8px;
                padding: 8px 0;
            }
            
            /* Project List */
            #ProjectList { 
                background: transparent; 
                border: none; 
                outline: none; 
            }
            
            #ProjectList::item { 
                padding: 14px 12px; 
                border-radius: 8px; 
                margin-bottom: 4px; 
                color: #475569;
                font-weight: 500;
                background: transparent;
            }
            
            #ProjectList::item:selected { 
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #2563eb, stop:1 #3b82f6);
                color: white; 
                font-weight: 600;
                border-left: 3px solid #60a5fa;
            }
            
            #ProjectList::item:hover:!selected { 
                background: rgba(226, 232, 240, 0.5);
            }
            
            /* Buttons - Modern Style */
            QPushButton { 
                background: rgba(255, 255, 255, 0.9);
                border: 1px solid rgba(226, 232, 240, 0.8);
                padding: 11px 18px; 
                border-radius: 8px; 
                font-weight: 600; 
                color: #334155;
            }
            
            QPushButton:hover { 
                background: rgba(255, 255, 255, 1);
                border-color: #cbd5e1;
                transform: translateY(-1px);
            }
            
            QPushButton:pressed { 
                background: #f1f5f9;
            }
            
            /* Primary Record Button - Hero Style */
            #RecordBtn { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #2563eb, stop:1 #1d4ed8);
                color: white; 
                font-size: 16px; 
                padding: 22px 50px; 
                border: none; 
                border-radius: 14px; 
                font-weight: 700; 
                letter-spacing: 0.5px;
                min-width: 350px;
            }
            
            #RecordBtn:hover { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #3b82f6, stop:1 #2563eb);
            }
            
            /* Sidebar Buttons */
            #PrimarySidebarBtn { 
                background: rgba(255, 255, 255, 0.9);
                border-color: #e2e8f0;
                font-weight: 600;
            }
            
            #PrimarySidebarBtn:hover {
                background: white;
                border-color: #cbd5e1;
            }
            
            #DeleteBtn { 
                background: rgba(254, 226, 226, 0.8);
                color: #dc2626; 
                border: 1px solid rgba(220, 38, 38, 0.2);
                font-weight: 600;
            }
            
            #DeleteBtn:hover { 
                background: rgba(254, 226, 226, 1);
                border-color: #dc2626;
            }
            
            #SettingsBtn {
                background: rgba(255, 255, 255, 0.9);
                border: 1px solid #e2e8f0;
                font-weight: 600;
                color: #475569;
            }
            
            #SettingsBtn:hover {
                background: white;
                border-color: #cbd5e1;
            }
            
            /* Labels */
            #ShortcutLabel { 
                color: #64748b; 
                font-size: 13px; 
                margin-top: 15px;
                font-weight: 500;
            }
            
            #FooterLabel { 
                color: #94a3b8; 
                font-size: 11px; 
                padding-top: 20px; 
                margin-top: 30px; 
                font-weight: 500;
            }
            
            /* Input Fields */
            QLineEdit, QTextEdit { 
                background: rgba(255, 255, 255, 0.9);
                border: 1px solid #e2e8f0;
                border-radius: 8px; 
                padding: 10px; 
                color: #1e293b;
            }
            
            QLineEdit:focus, QTextEdit:focus { 
                border-color: #3b82f6;
                background: white;
            }
            
            /* Group Boxes */
            QGroupBox { 
                font-weight: 600; 
                border: 1px solid #e2e8f0;
                margin-top: 18px; 
                padding-top: 22px; 
                border-radius: 10px;
                background: rgba(255, 255, 255, 0.5);
            }
            
            QGroupBox::title { 
                subcontrol-origin: margin; 
                left: 12px; 
                padding: 0 8px; 
                color: #64748b;
            }
        """
        
        # Override with background image if exists
        if has_bg:
            style = style.replace(
                "QMainWindow { \n                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, \n                    stop:0 #f8fafc, stop:0.5 #e0e7ff, stop:1 #f1f5f9);\n            }",
                f"QMainWindow {{ border-image: url({bg_path}) 0 0 0 0 stretch stretch; }}"
            )
            # Ensure text is readable on custom background
            style += """
                #ShortcutLabel, #FooterLabel { 
                    color: #1e293b; 
                    font-weight: 700;
                    text-shadow: 0 1px 2px rgba(255,255,255,0.8);
                }
                QLabel { 
                    color: #1e293b;
                    font-weight: 600;
                }
            """
        
        self.setStyleSheet(style)

    def setup_hotkeys(self):
        """Setup global keyboard shortcuts from settings - uses signals for thread safety"""
        if hasattr(self, 'hotkey_thread'):
            try:
                self.hotkey_thread.stop()
            except: pass
            
        h = {
            self.settings.shortcut_record: lambda: self.hotkey_signal.emit("record"),
            self.settings.shortcut_editor: lambda: self.hotkey_signal.emit("editor")
        }
        try:
            self.hotkey_thread = keyboard.GlobalHotKeys(h)
            self.hotkey_thread.start()
        except Exception as e:
            print(f"Hotkey Setup Error: {e}")

    def handle_hotkey(self, action):
        """Handler for hotkeys, executes in the UI thread"""
        if action == "record":
            self.btn_record.click()
        elif action == "editor":
            self.open_editor()

    def setup_ui(self):
        self.center_widget = QWidget()
        self.setCentralWidget(self.center_widget)
        
        # Main Layout: Horizontal (Left Sidebar | Right Content)
        main_layout = QHBoxLayout(self.center_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # --- LEFT SIDEBAR (Static Project List) ---
        self.sidebar = QWidget()
        self.sidebar.setObjectName("Sidebar")
        self.sidebar.setFixedWidth(280)
        
        side_layout = QVBoxLayout(self.sidebar)
        side_layout.setContentsMargins(15, 20, 15, 20)
        side_layout.setSpacing(10)
        
        lbl_proj = QLabel("PROJEKTE")
        lbl_proj.setObjectName("SidebarHeader")
        side_layout.addWidget(lbl_proj)
        
        self.proj_list = QListWidget()
        self.proj_list.setObjectName("ProjectList")
        side_layout.addWidget(self.proj_list)
        
        # Sidebar Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(5)
        
        btn_open = QPushButton("📂 Öffnen")
        btn_open.setObjectName("PrimarySidebarBtn")
        btn_open.setFixedHeight(40)
        btn_open.clicked.connect(self.load_project)
        
        btn_del = QPushButton("🗑️")
        btn_del.setObjectName("DeleteBtn")
        btn_del.setFixedSize(40, 40)
        btn_del.clicked.connect(self.delete_project)
        
        btn_layout.addWidget(btn_open)
        btn_layout.addWidget(btn_del)
        side_layout.addLayout(btn_layout)
        
        # Spacer
        side_layout.addStretch()
        
        # Settings Button
        self.btn_settings = QPushButton("⚙️ Einstellungen")
        self.btn_settings.setObjectName("SettingsBtn")
        self.btn_settings.setFixedHeight(45)
        self.btn_settings.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_settings.clicked.connect(self.open_settings)
        side_layout.addWidget(self.btn_settings)
        
        main_layout.addWidget(self.sidebar)
        
        # --- RIGHT CONTENT AREA ---
        content_area = QWidget()
        self.layout = QVBoxLayout(content_area)
        self.layout.setContentsMargins(50, 50, 50, 50)
        self.layout.setSpacing(30)
        
        if os.path.exists(resource_path("assets/logo.png")):
             # Logo removed as requested, keeping code block for structure but logic is disabled
             pass 
        
        self.layout.addStretch(3) # Spacer (Top - weight 3)
        
        self.btn_record = QPushButton("NEUE AUFNAHME STARTEN")
        self.btn_record.setObjectName("RecordBtn")
        self.btn_record.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_record.clicked.connect(self.toggle_recording)
        self.layout.addWidget(self.btn_record, 0, Qt.AlignmentFlag.AlignCenter)
        
        self.lbl_shortcuts = QLabel(f"Shortcuts: {self.settings.shortcut_record.upper()} (Start/Stop) | {self.settings.shortcut_editor.upper()} (Editor)")
        self.lbl_shortcuts.setObjectName("ShortcutLabel")
        self.lbl_shortcuts.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.layout.addWidget(self.lbl_shortcuts)
        
        self.layout.addStretch(1) # Bottom spacer (weight 1)
        
        # Attribution Footer
        footer = QLabel("© 2026 ClickStep Guide | High-Performance Documentation Engine")
        footer.setObjectName("FooterLabel")
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.layout.addWidget(footer)
        
        main_layout.addWidget(content_area)

    def toggle_recording(self):
        if not self.is_recording:
            self.is_recording = True
            self.steps = []
            self.btn_record.setText("AUFNAHME STOPPEN")
            self.showMinimized()
            
            # Reset and show overlay
            self.overlay.update_steps(0)
            screen = QApplication.primaryScreen().geometry()
            self.overlay.setGeometry(screen.width() - 260, 30, 240, 60)
            self.overlay.show() 
            self.recording_thread.start()
        else:
            self.stop_recording()

    def stop_recording(self):
        self.is_recording = False
        self.recording_thread.is_running = False
        self.overlay.hide() # Hide overlay
        self.btn_record.setText("NEUE AUFNAHME STARTEN")
        self.showNormal()
        
        # Give thread a moment to finish
        self.recording_thread.wait(500) # Wait max 500ms
        
        if self.steps:
            self.open_editor()

    def handle_click(self, x, y, label, img):
        self.steps.append(Step(img, x, y, label))
        self.overlay.update_steps(len(self.steps))

    def open_editor(self, project_name=None):
        self.hide() # Hide main recorder
        self.editor = ProEditor(self.steps, self.global_layers, self.global_crop, self.final_export, project_name=project_name, parent_window=self, settings=self.settings)
        self.editor.show()

    def open_settings(self):
        """Open settings dialog and apply changes"""
        dlg = SettingsDialog(self.settings, self)
        if dlg.exec():
            new_data = dlg.get_settings()
            
            # Update settings object
            self.settings.theme = new_data["theme"]
            self.settings.shortcut_record = new_data["shortcut_record"]
            self.settings.shortcut_editor = new_data["shortcut_editor"]
            
            # Persistence
            self.settings.save()
            
            # Apply immediate changes
            self.apply_app_theme()
            self.setup_hotkeys()
            self.lbl_shortcuts.setText(f"Shortcuts: {self.settings.shortcut_record.upper()} (Start/Stop) | {self.settings.shortcut_editor.upper()} (Editor)")
            
            QMessageBox.information(self, "Erfolg", "Einstellungen wurden gespeichert und angewendet!")

    def final_export(self, steps, global_layers, global_crop):
        if not steps: # This is a refresh signal for project list
            self.update_project_list()
            return
            
        path, _ = QFileDialog.getSaveFileName(self, "Export", "", "Word (*.docx)")
        if not path: return
        
        prog = QProgressBar()
        prog.setMaximum(len(steps))
        self.layout.addWidget(prog)
        
        try:
            doc = Document()
            doc.add_heading('ClickStep Guide - Anleitung', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, step_obj in enumerate(steps):
                # 1. Prepare Base Image (Crop)
                raw = step_obj.raw_img
                if global_crop:
                    x1, y1, x2, y2 = global_crop
                    raw = raw[y1:y2, x1:x2]
                
                canvas = raw.copy()
                ox, oy = (global_crop[0], global_crop[1]) if global_crop else (0, 0)
                
                # 2. Render all annotations (using the same logic as our editor's paint methods)
                # First Global Blur
                for gl in global_layers:
                    if gl.type == 'blur':
                        c = gl.data['coords']
                        blur_strength = gl.data.get('strength', 40)
                        self.render_blur_cv2(canvas, (c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy), True, blur_strength)
                
                # Then Step Layers
                for l in step_obj.layers:
                    if l.type == 'blur':
                        c = l.data['coords']
                        blur_strength = l.data.get('strength', 40)
                        self.render_blur_cv2(canvas, (c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy), False, blur_strength)
                    elif l.type == 'click':
                        self.render_click_cv2(canvas, l.data['x']-ox, l.data['y']-oy, i+1)
                    elif l.type == 'zoom':
                        self.render_zoom_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'infobox':
                        self.render_infobox_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'arrow':
                        self.render_arrow_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'icon':
                        self.render_icon_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'text':
                        self.render_text_cv2(canvas, l.data, ox, oy)
                
                # 3. Render Permanent Watermark
                self.render_watermark_cv2(canvas)
                
                # 4. Add to Word
                tmp_file = f"export_tmp_{i}.png"
                cv2.imwrite(tmp_file, canvas)
                doc.add_heading(f"Schritt {i+1}", level=1)
                if step_obj.description: doc.add_paragraph(step_obj.description)
                doc.add_picture(tmp_file, width=Inches(6))
                os.remove(tmp_file)
                prog.setValue(i+1)

            # 4. Add Copyright & Info Footer at the end
            doc.add_page_break()
            doc.add_heading('Über dieses Dokument', level=1)
            p = doc.add_paragraph()
            p.add_run('Diese Anleitung wurde professionell erstellt mit ').italic = True
            p.add_run('ClickStep Guide').bold = True
            p.add_run('.')
            doc.add_paragraph(f"Erstellungsdatum: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
            doc.add_paragraph("© 2026 AutoGuide Automationsysteme")

            doc.save(path)
            QMessageBox.information(self, "Erfolg", f"Export nach {path} abgeschlossen!")
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Export fehlgeschlagen: {str(e)}")
        finally:
            self.layout.removeWidget(prog)
            prog.deleteLater()

    # OpenCV Rendering Helpers for Export
    def render_blur_cv2(self, img, coords, is_global, strength=40):
        x1, y1, x2, y2 = coords
        h, w = img.shape[:2]
        x1, y1, x2, y2 = max(0, x1), max(0, y1), min(w, x2), min(h, y2)
        if x2 > x1 and y2 > y1:
            roi = img[y1:y2, x1:x2]
            k = strength | 1
            img[y1:y2, x1:x2] = cv2.GaussianBlur(roi, (k, k), 0)
            # No border in final export for cleaner look
            # cv2.rectangle(img, (x1, y1), (x2, y2), (0, 255, 0) if is_global else (255, 255, 255), 2)

    def render_click_cv2(self, img, x, y, num):
        settings = ClickMarkerSettings()
        
        # Use settings for rendering
        radius = int(settings.size * 0.6)
        glow_radius = int(settings.size * 0.9)
        
        # Check if transparent
        is_transparent = settings.color.alpha() == 0
        
        # Glow (optional, skip if transparent)
        if settings.show_glow and not is_transparent:
            col_bgr = (settings.color.blue(), settings.color.green(), settings.color.red())
            col_glow = tuple(int(c * 0.7) for c in col_bgr)  # Darker for glow
            cv2.circle(img, (x, y), glow_radius, col_glow, -1, lineType=cv2.LINE_AA)
        
        # Main circle fill (skip if transparent)
        if not is_transparent:
            col_bgr = (settings.color.blue(), settings.color.green(), settings.color.red())
            cv2.circle(img, (x, y), radius, col_bgr, -1, lineType=cv2.LINE_AA)
        
        # Border (always visible)
        cv2.circle(img, (x, y), radius, (255, 255, 255), settings.border_width, lineType=cv2.LINE_AA)
        
        # Number
        font_scale = settings.number_size / 16.0
        text_offset_x = int(settings.size * 0.25)
        text_offset_y = int(settings.size * 0.2)
        
        # Text Color from settings
        text_col = getattr(settings, 'text_color', QColor(255, 255, 255))
        # Convert QColor to BGR for OpenCV
        text_bgr = (text_col.blue(), text_col.green(), text_col.red())
        
        # Only draw outline if text is white (for contrast)
        if text_col.name() == "#ffffff":
            cv2.putText(img, str(num), (x-text_offset_x, y+text_offset_y), 
                       cv2.FONT_HERSHEY_SIMPLEX, font_scale, (0, 0, 0), 4, cv2.LINE_AA)
                       
        cv2.putText(img, str(num), (x-text_offset_x, y+text_offset_y), 
                   cv2.FONT_HERSHEY_SIMPLEX, font_scale, text_bgr, 2, cv2.LINE_AA)

    def render_zoom_cv2(self, img, data, ox, oy):
        zx, zy, sz = data['x']-ox, data['y']-oy, data['size']
        tx, ty = data['target_x']-ox, data['target_y']-oy
        col_rgb = data.get('color', (255, 255, 255))
        col_bgr = (col_rgb[2], col_rgb[1], col_rgb[0])
        
        src_sz = sz // 2
        
        # Crop area
        x1, y1 = max(0, tx - src_sz//2), max(0, ty - src_sz//2)
        x2, y2 = min(img.shape[1], tx + src_sz//2), min(img.shape[0], ty + src_sz//2)
        if x2 > x1 and y2 > y1:
            roi = img[y1:y2, x1:x2]
            zoomed = cv2.resize(roi, (sz, sz))
            try:
                img[zy:zy+sz, zx:zx+sz] = zoomed
            except: pass
            cv2.rectangle(img, (zx, zy), (zx+sz, zy+sz), col_bgr, 3)
            cv2.arrowedLine(img, (zx+sz//2, zy+sz//2), (tx, ty), col_bgr, 3)

    def render_infobox_cv2(self, img, data, ox, oy):
        x, y, w, h = data['x']-ox, data['y']-oy, data['w'], data['h']
        tx, ty = data['target_x']-ox, data['target_y']-oy
        col_rgb = data.get('color', (255, 255, 255))
        col_bgr = (col_rgb[2], col_rgb[1], col_rgb[0])
        
        # 1. Background Box
        sub = img[y:y+h, x:x+w]
        white_rect = np.full(sub.shape, 40, dtype=np.uint8) # Dark grey
        res = cv2.addWeighted(sub, 0.2, white_rect, 0.8, 1.0)
        img[y:y+h, x:x+w] = res
        cv2.rectangle(img, (x, y), (x+w, y+h), col_bgr, 2)
        
        # 2. Arrow (from border)
        cx, cy = x + w//2, y + h//2
        dx, dy = tx - cx, ty - cy
        
        # Liang-Barsky for OpenCV (Basic center to edge logic)
        hw, hh = w/2, h/2
        if dx != 0 or dy != 0:
            scale_x = hw / abs(dx) if dx != 0 else 9999
            scale_y = hh / abs(dy) if dy != 0 else 9999
            scale = min(scale_x, scale_y)
            
            sx = int(cx + dx * scale)
            sy = int(cy + dy * scale)
            
            cv2.line(img, (sx, sy), (tx, ty), col_bgr, 2, cv2.LINE_AA)
            cv2.circle(img, (tx, ty), 4, col_bgr, -1)
        
        # 3. Text (Simple wrapping)
        text = data['text']
        font = cv2.FONT_HERSHEY_SIMPLEX
        
        # Map point size to CV2 scale
        f_data = data.get('font', {})
        pt_size = f_data.get('size', 12)
        scale = pt_size / 24.0 # 12pt -> 0.5 scale
        
        t_col = data.get('text_color', (255, 255, 255))
        color = (t_col[2], t_col[1], t_col[0]) # BGR
        thickness = 1 if pt_size < 18 else 2
        
        line_h = int(pt_size * 1.5)
        dy_text = y + line_h
        for line in text.split('\n'):
            cv2.putText(img, line, (x+10, dy_text), font, scale, color, thickness, cv2.LINE_AA)
            dy_text += line_h

    def render_arrow_cv2(self, img, data, ox, oy):
        sx, sy, ex, ey = data['sx']-ox, data['sy']-oy, data['ex']-ox, data['ey']-oy
        col_rgb = data.get('color', (255, 0, 0))
        col_bgr = (col_rgb[2], col_rgb[1], col_rgb[0])
        w = data.get('width', 4)
        
        cv2.arrowedLine(img, (sx, sy), (ex, ey), col_bgr, w, tipLength=0.2)
        
    def render_icon_cv2(self, img, data, ox, oy):
        x, y = data['x']-ox, data['y']-oy
        t_col = data.get('color', (255, 0, 0))
        col_bgr = (t_col[2], t_col[1], t_col[0])
        size = data.get('w', data.get('size', 60))
        
        # Simple Viz: Circle + Text using selected color
        cv2.circle(img, (x+size//2, y+size//2), size//2, col_bgr, 2)
        
        # Mapping meaningful text for CV2 which doesn't do multicolor emoji
        txt_map = {
            'check': 'OK', 
            'cross': 'X', 
            'warn': '!', 
            'info': 'i', 
            'star': '*', 
            'idea': '?',
            'arrow_up': '^',
            'arrow_down': 'v',
            'heart': '<3'
        }
        t = txt_map.get(data.get('type'), '?')
        
        font_scale = size / 40.0
        cv2.putText(img, t, (x+size//4, y+int(size*0.75)), cv2.FONT_HERSHEY_SIMPLEX, font_scale, col_bgr, 3)

    def render_text_cv2(self, img, data, ox, oy):
        x, y = data['x']-ox, data['y']-oy
        color = data.get('color', (255, 255, 255)) # BGR
        cv2.putText(img, data['text'], (x, y), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (0,0,0), 6, cv2.LINE_AA)
        cv2.putText(img, data['text'], (x, y), cv2.FONT_HERSHEY_SIMPLEX, 1.2, color, 2, cv2.LINE_AA)

    def render_watermark_cv2(self, img):
        """Draw a permanent watermark in the top right corner"""
        h, w = img.shape[:2]
        text = "Created with ClickStep Guide"
        font = cv2.FONT_HERSHEY_SIMPLEX
        scale = 0.6
        thickness = 1
        color = (255, 255, 255) # White (BGR)
        
        size = cv2.getTextSize(text, font, scale, thickness)[0]
        # Top Right Position: x = width - text_width - margin, y = text_height + margin
        tx, ty = w - size[0] - 20, size[1] + 20
        
        # Shadow for visibility on bright backgrounds
        cv2.putText(img, text, (tx+1, ty+1), font, scale, (0, 0, 0), thickness+1, cv2.LINE_AA)
        cv2.putText(img, text, (tx, ty), font, scale, color, thickness, cv2.LINE_AA)

    def update_project_list(self):
        """Update project list in sidebar"""
        self.proj_list.clear()
        path = self.get_project_dir()
        if not os.path.exists(path): 
            os.makedirs(path)
        for d in os.listdir(path):
            if os.path.isdir(os.path.join(path, d)):
                item = QListWidgetItem(d)
                self.proj_list.addItem(item)

    def load_project(self):
        """Load selected project"""
        item = self.proj_list.currentItem()
        if not item: return
        
        project_path = os.path.join(self.get_project_dir(), item.text(), "project.json")
        if not os.path.exists(project_path):
            QMessageBox.warning(self, "Fehler", "Projekt-Datei nicht gefunden!")
            return
        
        try:
            with open(project_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Load steps
            self.steps = []
            img_path = os.path.join(self.get_project_dir(), item.text(), "images")
            for step_data in data.get("steps", []):
                img_file = os.path.join(img_path, step_data["image"])
                if os.path.exists(img_file):
                    img = cv2.imread(img_file)
                    step = Step(img, 0, 0, step_data.get("description", ""))
                    step.layers = [] # Reset default click
                    for l_data in step_data.get("layers", []):
                        step.layers.append(Layer(l_data['type'], l_data['data'], l_data.get('label', 'Layer')))
                    
                    # Update step x, y from the first click layer found
                    click_l = next((l for l in step.layers if l.type == 'click'), None)
                    if click_l:
                        step.x, step.y = click_l.data['x'], click_l.data['y']
                        
                    self.steps.append(step)
            
            self.global_crop = data.get("global_crop")
            self.global_layers = []
            for gl_data in data.get("global_layers", []):
                self.global_layers.append(Layer(gl_data['type'], gl_data['data'], gl_data.get('label', 'Global Layer'), True))
            
            if self.steps:
                self.open_editor(project_name=item.text())
            else:
                QMessageBox.warning(self, "Warnung", "Keine Schritte im Projekt gefunden!")
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Projekt konnte nicht geladen werden: {str(e)}")

    def delete_project(self):
        """Delete selected project"""
        item = self.proj_list.currentItem()
        if not item: return
        
        reply = QMessageBox.question(self, 'Löschen', 
                                    f"Projekt '{item.text()}' wirklich löschen?", 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                shutil.rmtree(os.path.join(self.get_project_dir(), item.text()))
                self.update_project_list()
                QMessageBox.information(self, "Erfolg", "Projekt gelöscht!")
            except Exception as e:
                QMessageBox.critical(self, "Fehler", f"Löschen fehlgeschlagen: {str(e)}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ProRecorder()
    window.show()
    sys.exit(app.exec())
