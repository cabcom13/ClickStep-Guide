import tkinter as tk
from tkinter import filedialog, ttk, messagebox, colorchooser, simpledialog
from pynput import mouse, keyboard
from PIL import ImageGrab, Image, ImageTk
import cv2
import numpy as np
import threading
import datetime
import os
import math
import time
import json
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32api
import win32con

# --- UI STYLE ---
BG_DARK = "#181818"
BG_SIDE = "#202020"
ACCENT = "#007acc"
BGR_ACCENT = (204, 122, 0)
TEXT_COLOR = "#e0e0e0"
BORDER_COLOR = "#333333"

class Layer:
    def __init__(self, type, data, label, is_global=False):
        self.type = type # 'click', 'blur', 'shortcut', 'zoom', 'text'
        self.data = data
        self.label = label
        self.visible = True
        self.is_global = is_global

class Step:
    def __init__(self, raw_img, x, y, button, click_label, shortcut=None):
        self.raw_img = raw_img
        self.description = ""
        self.layers = []
        self.layers.append(Layer('click', {'x': x, 'y': y, 'color': (0,0,255)}, click_label))
        if shortcut:
            self.layers.append(Layer('shortcut', {'text': shortcut}, "Shortcut"))

    @property
    def x(self): return next((l.data['x'] for l in self.layers if l.type == 'click'), 0)
    @x.setter
    def x(self, v): 
        for l in self.layers: 
            if l.type == 'click': l.data['x'] = v
    @property
    def y(self): return next((l.data['y'] for l in self.layers if l.type == 'click'), 0)
    @y.setter
    def y(self, v):
        for l in self.layers:
            if l.type == 'click': l.data['y'] = v

class StepEditor:
    def __init__(self, parent_root, steps_data, global_layers, global_crop, save_callback):
        self.root = tk.Toplevel(parent_root)
        self.root.title("AutoGuide Professional Editor")
        self.root.state('zoomed')
        self.root.configure(bg=BG_DARK)
        
        self.steps = steps_data
        self.global_layers = global_layers
        self.global_crop = global_crop # (x1, y1, x2, y2) or None
        self.save_callback = save_callback
        self.current_idx = 0
        self.selected_layer_idx = -1
        self.selected_is_global = False
        
        self.ratio = 1.0
        self.dragging_layer = False
        self.active_handle = None 
        self.draw_mode = None # 'blur', 'global_blur', 'zoom', 'text', 'crop'
        self.draw_start = None
        self.active_tool_btn = None

        self.setup_ui()
        self.root.bind('<Escape>', self.on_escape)
        self.root.bind('<Delete>', lambda e: self.delete_selected_layer())
        self.root.bind('<BackSpace>', lambda e: self.delete_selected_layer())
        if self.steps: self.root.after(100, lambda: self.select_step(0))

    def setup_ui(self):
        # Toolbar
        top = tk.Frame(self.root, bg=BG_SIDE, height=70, bd=0, highlightthickness=1, highlightbackground=BORDER_COLOR)
        top.pack(side=tk.TOP, fill=tk.X)
        tk.Label(top, text="PRO EDITOR", fg=ACCENT, bg=BG_SIDE, font=("Segoe UI", 14, "bold")).pack(side=tk.LEFT, padx=15, pady=10)
        
        # Tool Groups
        tool_frame = tk.Frame(top, bg=BG_SIDE)
        tool_frame.pack(side=tk.LEFT, padx=20)
        
        btn_style = {"bg": "#2a2a2a", "fg": "white", "relief": tk.FLAT, "padx": 12, "pady": 8, "font": ("Segoe UI", 9)}
        
        # Drawing Tools
        tk.Label(tool_frame, text="DRAWING", bg=BG_SIDE, fg="#888", font=("Segoe UI", 7)).grid(row=0, column=0, columnspan=3, sticky="w")
        self.btn_zoom = tk.Button(tool_frame, text="🔍 Zoom", **btn_style, command=lambda: self.set_tool('zoom'))
        self.btn_zoom.grid(row=1, column=0, padx=2, pady=2)
        self.btn_blur = tk.Button(tool_frame, text="🛡️ Blur", **btn_style, command=lambda: self.set_tool('blur'))
        self.btn_blur.grid(row=1, column=1, padx=2, pady=2)
        self.btn_crop = tk.Button(tool_frame, text="✂️ Crop", **btn_style, command=lambda: self.set_tool('crop'))
        self.btn_crop.grid(row=1, column=2, padx=2, pady=2)
        
        # Global Tools
        tk.Label(tool_frame, text="GLOBAL", bg=BG_SIDE, fg="#888", font=("Segoe UI", 7)).grid(row=0, column=3, sticky="w", padx=(15,0))
        self.btn_global_blur = tk.Button(tool_frame, text="🌍 Global Blur", bg="#2c3e50", fg="white", relief=tk.FLAT, padx=12, pady=8, font=("Segoe UI", 9, "bold"), command=lambda: self.set_tool('global_blur'))
        self.btn_global_blur.grid(row=1, column=3, padx=(15,2), pady=2)
        
        # Text Tools
        tk.Label(tool_frame, text="TEXT", bg=BG_SIDE, fg="#888", font=("Segoe UI", 7)).grid(row=0, column=4, sticky="w", padx=(15,0))
        self.btn_text = tk.Button(tool_frame, text="📝 Text", **btn_style, command=lambda: self.set_tool('text'))
        self.btn_text.grid(row=1, column=4, padx=(15,2), pady=2)
        
        # Actions
        tk.Button(top, text="🗑️ Schritt löschen", bg="#8e0000", fg="white", relief=tk.FLAT, padx=15, pady=8, command=self.delete_current).pack(side=tk.RIGHT, padx=5)
        tk.Button(top, text="📁 Projekt Speichern", bg="#444", fg="white", relief=tk.FLAT, padx=15, pady=8, command=self.save_project).pack(side=tk.RIGHT, padx=5)
        tk.Button(top, text="💾 EXPORT", bg=ACCENT, fg="white", relief=tk.FLAT, padx=25, pady=10, font=("Segoe UI", 10, "bold"), command=self.save_all).pack(side=tk.RIGHT, padx=20)

        self.panes = tk.PanedWindow(self.root, orient=tk.HORIZONTAL, bg=BG_DARK, borderwidth=0, sashwidth=6)
        self.panes.pack(fill=tk.BOTH, expand=True)

        # Left: Thumbnails
        self.side_left = tk.Frame(self.panes, bg=BG_SIDE, width=220)
        self.panes.add(self.side_left, stretch="never")
        self.canvas_thumbs = tk.Canvas(self.side_left, bg=BG_SIDE, highlightthickness=0)
        self.thumb_box = tk.Frame(self.canvas_thumbs, bg=BG_SIDE)
        self.canvas_thumbs.create_window((0,0), window=self.thumb_box, anchor="nw", width=220)
        self.canvas_thumbs.pack(fill=tk.BOTH, expand=True)

        # Center
        self.center = tk.Frame(self.panes, bg=BG_DARK)
        self.panes.add(self.center, stretch="always")
        self.canvas = tk.Canvas(self.center, bg=BG_DARK, highlightthickness=0)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        self.canvas.bind("<Button-1>", self.on_mousedown)
        self.canvas.bind("<B1-Motion>", self.on_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_mouseup)
        self.canvas.bind("<Configure>", lambda e: self.draw_step())

        # Right: Properties & Layers
        self.side_right = tk.Frame(self.panes, bg=BG_SIDE, width=300)
        self.panes.add(self.side_right, stretch="never")
        
        # Properties Section
        tk.Label(self.side_right, text="PROPERTIES", bg=BG_SIDE, fg=ACCENT, font=("Segoe UI", 10, "bold")).pack(pady=(20, 5))
        self.props_frame = tk.Frame(self.side_right, bg=BG_SIDE)
        self.props_frame.pack(fill=tk.X, padx=20, pady=5)
        
        # Layers Section
        tk.Label(self.side_right, text="LAYERS", bg=BG_SIDE, fg=ACCENT, font=("Segoe UI", 10, "bold")).pack(pady=(20, 5))
        self.layer_list = tk.Listbox(self.side_right, bg="#2a2a2a", fg="white", borderwidth=0, 
                                     highlightthickness=1, highlightbackground=BORDER_COLOR, selectbackground=ACCENT, font=("Segoe UI", 9), height=12)
        self.layer_list.pack(fill=tk.X, padx=20, pady=5)
        self.layer_list.bind("<<ListboxSelect>>", self.on_layer_select)
        
        tk.Button(self.side_right, text="🗑️ Ebene löschen (Del)", bg="#444", fg="#ff6666", 
                  relief=tk.FLAT, font=("Arial", 8), command=self.delete_selected_layer).pack(pady=5)

        tk.Label(self.side_right, text="BESCHREIBUNG", bg=BG_SIDE, fg="#888", font=("Segoe UI", 9)).pack(anchor=tk.W, padx=20, pady=(20, 0))
        self.txt_desc = tk.Text(self.side_right, height=6, bg="#2a2a2a", fg="white", borderwidth=0, padx=10, pady=10, font=("Segoe UI", 10))
        self.txt_desc.pack(fill=tk.X, padx=20, pady=5)
        self.txt_desc.bind("<KeyRelease>", self.update_desc)

        self.update_thumbnails()

    def update_properties_panel(self):
        """Update properties panel based on selected layer"""
        for w in self.props_frame.winfo_children(): w.destroy()
        
        if self.selected_layer_idx < 0:
            tk.Label(self.props_frame, text="Keine Ebene ausgewählt", bg=BG_SIDE, fg="#666", font=("Segoe UI", 9, "italic")).pack()
            return
        
        if self.selected_is_global:
            layer = self.global_layers[self.selected_layer_idx]
        else:
            layer = self.steps[self.current_idx].layers[self.selected_layer_idx]
        
        # Show properties based on layer type
        if layer.type == 'text':
            tk.Label(self.props_frame, text="Text:", bg=BG_SIDE, fg="white", font=("Segoe UI", 9)).pack(anchor="w")
            txt_entry = tk.Entry(self.props_frame, bg="#2a2a2a", fg="white", font=("Segoe UI", 10))
            txt_entry.insert(0, layer.data['text'])
            txt_entry.pack(fill=tk.X, pady=5)
            txt_entry.bind("<KeyRelease>", lambda e: self.update_text_layer(layer, txt_entry.get()))
            
            tk.Label(self.props_frame, text="Größe:", bg=BG_SIDE, fg="white", font=("Segoe UI", 9)).pack(anchor="w", pady=(10,0))
            size_scale = tk.Scale(self.props_frame, from_=0.5, to=3.0, resolution=0.1, orient=tk.HORIZONTAL, bg=BG_SIDE, fg="white", highlightthickness=0)
            size_scale.set(layer.data.get('size', 1.0))
            size_scale.pack(fill=tk.X)
            size_scale.bind("<ButtonRelease-1>", lambda e: self.update_text_size(layer, size_scale.get()))
            
            tk.Button(self.props_frame, text="🎨 Farbe ändern", bg="#444", fg="white", relief=tk.FLAT, command=lambda: self.change_text_color(layer)).pack(pady=10)

    def update_text_layer(self, layer, text):
        layer.data['text'] = text
        self.draw_step()
    
    def update_text_size(self, layer, size):
        layer.data['size'] = size
        self.draw_step()
    
    def change_text_color(self, layer):
        color = colorchooser.askcolor(title="Textfarbe wählen", initialcolor="#ffffff", parent=self.root)
        if color[1]:
            # Convert hex to BGR
            hex_color = color[1].lstrip('#')
            r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            layer.data['color'] = (b, g, r) 
            self.draw_step()

    def set_tool(self, tool):
        """Set active tool and update UI"""
        # Reset all buttons
        for btn in [self.btn_zoom, self.btn_blur, self.btn_crop, self.btn_global_blur, self.btn_text]:
            btn.config(bg="#2a2a2a", relief=tk.FLAT)
        
        # Highlight active button
        if tool == 'zoom': 
            self.btn_zoom.config(bg=ACCENT, relief=tk.SUNKEN)
            self.draw_mode = 'zoom'
            self.canvas.config(cursor="crosshair")
        elif tool == 'blur': 
            self.btn_blur.config(bg=ACCENT, relief=tk.SUNKEN)
            self.draw_mode = 'blur'
            self.canvas.config(cursor="crosshair")
        elif tool == 'crop':
            self.btn_crop.config(bg=ACCENT, relief=tk.SUNKEN)
            self.draw_mode = 'crop'
            self.canvas.config(cursor="sizing")
        elif tool == 'global_blur': 
            self.btn_global_blur.config(bg="#1e5a7a", relief=tk.SUNKEN)
            self.draw_mode = 'global_blur'
            self.canvas.config(cursor="crosshair")
        elif tool == 'text':
            self.btn_text.config(bg=ACCENT, relief=tk.SUNKEN)
            self.draw_mode = 'text'
            self.canvas.config(cursor="ibeam") # Fixed Windows cursor name

    def on_escape(self, e):
        """Deselect and cancel modes"""
        self.selected_layer_idx = -1
        self.selected_is_global = False
        self.draw_mode = None
        self.draw_start = None
        self.canvas.config(cursor="")
        
        # Reset toolbar
        for btn in [self.btn_zoom, self.btn_blur, self.btn_global_blur, self.btn_text]:
            btn.config(bg="#2a2a2a", relief=tk.FLAT)
        
        self.update_properties_panel()
        self.draw_step()

    def update_thumbnails(self):
        for w in self.thumb_box.winfo_children(): w.destroy()
        for i, s in enumerate(self.steps):
            f = tk.Frame(self.thumb_box, bg=ACCENT if i == self.current_idx else BG_SIDE, pady=5)
            f.pack(fill=tk.X, padx=10, pady=2)
            
            thumb_img = s.raw_img.copy()
            for gl in self.global_layers:
                if gl.type == 'blur':
                    x1,y1,x2,y2 = gl.data['coords']
                    area = thumb_img[y1:y2, x1:x2]
                    if area.size > 0: thumb_img[y1:y2, x1:x2] = cv2.GaussianBlur(area, (51,51), 20)
            
            small = cv2.resize(thumb_img, (180, 101))
            img_tk = ImageTk.PhotoImage(image=Image.fromarray(cv2.cvtColor(small, cv2.COLOR_BGR2RGB)))
            btn = tk.Label(f, image=img_tk, bg=BG_SIDE); btn.image = img_tk
            btn.bind("<Button-1>", lambda e, idx=i: self.select_step(idx))
            btn.pack()
            tk.Label(f, text=f"SCHRITT {i+1}", fg="white", bg=f.cget("bg"), font=("Segoe UI", 8)).pack()

    def select_step(self, idx):
        self.current_idx = idx
        step = self.steps[idx]
        self.txt_desc.delete("1.0", tk.END); self.txt_desc.insert("1.0", step.description)
        self.refresh_layer_list()
        self.draw_step(); self.update_thumbnails()

    def refresh_layer_list(self):
        self.layer_list.delete(0, tk.END)
        for layer in self.global_layers:
            icon = "🛡️" if layer.type == 'blur' else "📝"
            self.layer_list.insert(tk.END, f"🌍 [GLOBAL] {icon} {layer.label}")
        for layer in self.steps[self.current_idx].layers:
            icons = {'click': '🎯', 'blur': '🛡️', 'zoom': '🔍', 'text': '📝', 'shortcut': '⌨️'}
            icon = icons.get(layer.type, '●')
            self.layer_list.insert(tk.END, f"{icon} {layer.label}")

    def on_layer_select(self, e):
        sel = self.layer_list.curselection()
        if sel:
            idx = sel[0]
            if idx < len(self.global_layers):
                self.selected_layer_idx = idx; self.selected_is_global = True
            else:
                self.selected_layer_idx = idx - len(self.global_layers)
                self.selected_is_global = False
            self.update_properties_panel()
            self.draw_step()

    def delete_selected_layer(self):
        if self.selected_layer_idx >= 0:
            if self.selected_is_global:
                self.global_layers.pop(self.selected_layer_idx)
            else:
                layer = self.steps[self.current_idx].layers[self.selected_layer_idx]
                if layer.type == 'click':
                    messagebox.showwarning("Fehler", "Haupt-Marker kann nicht entfernt werden.")
                    return
                self.steps[self.current_idx].layers.pop(self.selected_layer_idx)
            self.selected_layer_idx = -1
            self.refresh_layer_list(); self.update_properties_panel(); self.draw_step()

    def update_desc(self, e): self.steps[self.current_idx].description = self.txt_desc.get("1.0", "end-1c")

    def draw_step(self):
        if not self.steps: return
        step = self.steps[self.current_idx]
        
        # Base image: Apply Global Crop if exists
        raw = step.raw_img
        if self.global_crop:
            x1, y1, x2, y2 = self.global_crop
            raw = raw[y1:y2, x1:x2]
        
        rendered = raw.copy()
        
        # Adjust layer rendering to handle crop offset
        offset_x = self.global_crop[0] if self.global_crop else 0
        offset_y = self.global_crop[1] if self.global_crop else 0
        # Global Layers
        for i, layer in enumerate(self.global_layers):
            hl = (self.selected_is_global and i == self.selected_layer_idx)
            coords = layer.data['coords']
            # Offset coords for drawing if global_crop exists
            adjusted_coords = (coords[0]-offset_x, coords[1]-offset_y, coords[2]-offset_x, coords[3]-offset_y)
            if layer.type == 'blur': self.render_blur(rendered, adjusted_coords, hl, is_global=True)
            elif layer.type == 'text':
                # Simplified offset for text: render_text handles its own layer
                self.render_text(rendered, layer, hl, offset_x, offset_y)
            
        # Local Layers
        for i, layer in enumerate(step.layers):
            if layer.type in ['zoom', 'text']: continue
            hl = (not self.selected_is_global and i == self.selected_layer_idx)
            if layer.type == 'click':
                self.draw_pro_marker(rendered, layer.data['x']-offset_x, layer.data['y']-offset_y, layer.data['color'], str(self.current_idx+1), layer.label, hl)
            elif layer.type == 'blur':
                coords = layer.data['coords']
                adjusted_coords = (coords[0]-offset_x, coords[1]-offset_y, coords[2]-offset_x, coords[3]-offset_y)
                self.render_blur(rendered, adjusted_coords, hl)
            elif layer.type == 'shortcut':
                click = next((l for l in step.layers if l.type == 'click'), None)
                if click: self.draw_pro_text(rendered, layer.data['text'], (click.data['x']-offset_x+20, click.data['y']-offset_y+80), (100, 255, 100), scale=0.7)

        # Zoom & Text Layers (always on top)
        for i, layer in enumerate(step.layers):
            if layer.type == 'zoom':
                hl = (not self.selected_is_global and i == self.selected_layer_idx)
                self.render_zoom(rendered, layer, hl, offset_x, offset_y)
            elif layer.type == 'text':
                hl = (not self.selected_is_global and i == self.selected_layer_idx)
                self.render_text(rendered, layer, hl, offset_x, offset_y)

        # Display
        cw, ch = self.canvas.winfo_width(), self.canvas.winfo_height()
        if cw < 50: return
        h, w = rendered.shape[:2]
        self.ratio = min((cw-40)/w, (ch-40)/h)
        nw, nh = int(w*self.ratio), int(h*self.ratio)
        rgb = cv2.cvtColor(cv2.resize(rendered, (nw, nh)), cv2.COLOR_BGR2RGB)
        img_tk = ImageTk.PhotoImage(image=Image.fromarray(rgb))
        self.canvas.delete("all")
        self.canvas.create_image(cw//2, ch//2, image=img_tk)
        self.canvas.image = img_tk

    def render_blur(self, img, coords, hl, is_global=False):
        x1, y1, x2, y2 = coords
        h, w = img.shape[:2]
        x1, y1 = max(0, x1), max(0, y1)
        x2, y2 = min(w, x2), min(h, y2)
        
        if x2 > x1 and y2 > y1:
            roi = img[y1:y2, x1:x2]
            if roi.size > 0:
                img[y1:y2, x1:x2] = cv2.GaussianBlur(roi, (75, 75), 40)
                color = (255, 100, 100) if hl else (255, 255, 255) if is_global else (150, 150, 150)
                cv2.rectangle(img, (x1, y1), (x2, y2), color, 2 if hl else 1)
                if is_global: cv2.putText(img, "GLOBAL", (x1, y1-5), cv2.FONT_HERSHEY_SIMPLEX, 0.4, color, 1)

    def render_text(self, img, layer, hl, ox=0, oy=0):
        d = layer.data
        text, x, y = d['text'], d['x']-ox, d['y']-oy
        color = d.get('color', (255, 255, 255))
        size = d.get('size', 1.0)
        
        # High-quality text with outline
        cv2.putText(img, text, (x, y), cv2.FONT_HERSHEY_SIMPLEX, size, (0, 0, 0), int(size*5+2), cv2.LINE_AA)
        cv2.putText(img, text, (x, y), cv2.FONT_HERSHEY_SIMPLEX, size, color, int(size*2), cv2.LINE_AA)
        
        if hl:
            text_size = cv2.getTextSize(text, cv2.FONT_HERSHEY_SIMPLEX, size, int(size*2))[0]
            cv2.rectangle(img, (x-5, y-text_size[1]-5), (x+text_size[0]+5, y+5), BGR_ACCENT, 2)

    def render_zoom(self, img, layer, hl, ox=0, oy=0):
        d = layer.data; zx, zy, sz = d['x']-ox, d['y']-oy, d['size']; tx, ty = d['target_x']-ox, d['target_y']-oy
        src_sz = sz // 2
        x1, y1 = max(0, tx - src_sz//2), max(0, ty - src_sz//2)
        x2, y2 = min(img.shape[1], tx + src_sz//2), min(img.shape[0], ty + src_sz//2)
        roi = self.steps[self.current_idx].raw_img[y1:y2, x1:x2]
        if roi.size > 0:
            zoomed = cv2.resize(roi, (sz, sz))
            cv2.circle(zoomed, (sz//2, sz//2), 30, (0,0,255), 4, cv2.LINE_AA)
            cv2.circle(zoomed, (sz//2, sz//2), 6, (255,255,255), -1, cv2.LINE_AA)

            h, w = img.shape[:2]
            y_start, y_end = max(0, zy), min(h, zy+sz)
            x_start, x_end = max(0, zx), min(w, zx+sz)
            
            if y_end > y_start and x_end > x_start:
                cv2.rectangle(img, (zx-4, zy-4), (zx+sz+4, zy+sz+4), (0,0,0), -1)
                cv2.rectangle(img, (zx-2, zy-2), (zx+sz+2, zy+sz+2), (255,255,255), -1)
                img[y_start:y_end, x_start:x_end] = zoomed[y_start-zy:y_end-zy, x_start-zx:x_end-zx]
                
                if hl:
                    cv2.rectangle(img, (zx-2, zy-2), (zx+sz+2, zy+sz+2), BGR_ACCENT, 4)
                    h_sz = 12
                    for (hx, hy) in [(zx, zy), (zx+sz, zy), (zx, zy+sz), (zx+sz, zy+sz)]:
                        cv2.rectangle(img, (hx-h_sz, hy-h_sz), (hx+h_sz, hy+h_sz), (0,0,0), -1)
                        cv2.rectangle(img, (hx-h_sz+2, hy-h_sz+2), (hx+h_sz-2, hy+h_sz-2), (255,255,255), -1)
                        cv2.rectangle(img, (hx-h_sz+4, hy-h_sz+4), (hx+h_sz-4, hy+h_sz-4), BGR_ACCENT, -1)
            
            cx, cy = zx + sz//2, zy + sz//2
            angle = math.atan2(ty - cy, tx - cx)
            sx, sy = int(cx + (sz//2 + 5) * math.cos(angle)), int(cy + (sz//2 + 5) * math.sin(angle))
            cv2.arrowedLine(img, (sx, sy), (tx, ty), (255,255,255), 6, tipLength=0.1)
            cv2.arrowedLine(img, (sx, sy), (tx, ty), (0,0,0), 2, tipLength=0.1)

    def draw_pro_marker(self, img, x, y, color, num, label, hl):
        cv2.circle(img, (x, y), 38, (255, 255, 255), 6, cv2.LINE_AA)
        cv2.circle(img, (x, y), 38, (0, 0, 0), 2, cv2.LINE_AA)
        ov = img.copy(); cv2.circle(ov, (x, y), 35, color, -1)
        cv2.addWeighted(ov, 0.4 if hl else 0.2, img, 0.6 if hl else 0.8, 0, img)
        cv2.circle(img, (x, y), 32, color, 4, cv2.LINE_AA)
        cv2.circle(img, (x, y), 8, (255, 255, 255), -1, cv2.LINE_AA)
        self.draw_pro_text(img, num, (x+18, y-18), (255, 255, 255), scale=1.3, thick=4)
        self.draw_pro_text(img, label, (x+18, y+45), (255, 255, 255), scale=0.8, thick=2)

    def draw_pro_text(self, img, text, pos, color, scale=1.0, thick=2):
        cv2.putText(img, text, (pos[0]+2, pos[1]+2), cv2.FONT_HERSHEY_SIMPLEX, scale, (0,0,0), thick+2, cv2.LINE_AA)
        cv2.putText(img, text, pos, cv2.FONT_HERSHEY_SIMPLEX, scale, color, thick, cv2.LINE_AA)

    def on_mousedown(self, e):
        rx, ry = self.get_img_coords(e.x, e.y)
        
        if self.draw_mode == 'text':
            # Text placement
            text = simpledialog.askstring("Text eingeben", "Text:", parent=self.root)
            if text:
                color = colorchooser.askcolor(title="Textfarbe")
                if color[0]:
                    r, g, b = [int(c) for c in color[0]]
                    self.steps[self.current_idx].layers.append(
                        Layer('text', {'text': text, 'x': rx, 'y': ry, 'color': (b,g,r), 'size': 1.2}, "Text"))
                    self.refresh_layer_list(); self.draw_step()
            self.draw_mode = None
            self.canvas.config(cursor="")
            for btn in [self.btn_text]: btn.config(bg="#2a2a2a", relief=tk.FLAT)
            return
        
        if self.draw_mode:
            self.draw_start = (rx, ry); return
        
        self.active_handle = None
        if self.selected_layer_idx >= 0 and not self.selected_is_global:
            layer = self.steps[self.current_idx].layers[self.selected_layer_idx]
            if layer.type == 'zoom':
                d = layer.data; zx, zy, sz = d['x'], d['y'], d['size']
                h_sz = 20
                if math.sqrt((rx - zx)**2 + (ry - zy)**2) < h_sz: self.active_handle = 'nw'
                elif math.sqrt((rx - (zx+sz))**2 + (ry - zy)**2) < h_sz: self.active_handle = 'ne'
                elif math.sqrt((rx - zx)**2 + (ry - (zy+sz))**2) < h_sz: self.active_handle = 'sw'
                elif math.sqrt((rx - (zx+sz))**2 + (ry - (zy+sz))**2) < h_sz: self.active_handle = 'se'
                
                if self.active_handle: self.dragging_layer = True; return
        
        step = self.steps[self.current_idx]
        for i, layer in enumerate(reversed(step.layers)):
            idx = len(step.layers) - 1 - i
            if layer.type == 'zoom':
                d = layer.data
                if d['x'] <= rx <= d['x']+d['size'] and d['y'] <= ry <= d['y']+d['size']:
                    self.selected_layer_idx = idx; self.selected_is_global = False; self.dragging_layer = True
                    self.refresh_layer_list(); self.layer_list.select_set(len(self.global_layers) + idx)
                    self.update_properties_panel(); return
            elif layer.type == 'text':
                text_size = cv2.getTextSize(layer.data['text'], cv2.FONT_HERSHEY_SIMPLEX, layer.data.get('size', 1.0), 2)[0]
                if layer.data['x'] <= rx <= layer.data['x']+text_size[0] and layer.data['y']-text_size[1] <= ry <= layer.data['y']:
                    self.selected_layer_idx = idx; self.selected_is_global = False; self.dragging_layer = True
                    self.refresh_layer_list(); self.layer_list.select_set(len(self.global_layers) + idx)
                    self.update_properties_panel(); return
            elif layer.type == 'click':
                if math.sqrt((rx - layer.data['x'])**2 + (ry - layer.data['y'])**2) < 60:
                    self.selected_layer_idx = idx; self.selected_is_global = False; self.dragging_layer = True
                    self.refresh_layer_list(); self.layer_list.select_set(len(self.global_layers) + idx)
                    self.update_properties_panel(); return
        self.dragging_layer = False

    def on_drag(self, e):
        rx, ry = self.get_img_coords(e.x, e.y)
        
        # Live preview while drawing
        if self.draw_mode and self.draw_start:
            self.draw_step()
            cw, ch = self.canvas.winfo_width(), self.canvas.winfo_height()
            
            # Use raw size or cropped size for preview calculation
            if self.global_crop:
                h, w = self.global_crop[3]-self.global_crop[1], self.global_crop[2]-self.global_crop[0]
            else:
                h, w = self.steps[self.current_idx].raw_img.shape[:2]
            
            off_x, off_y = (cw - w*self.ratio)/2, (ch - h*self.ratio)/2
            
            x1_c = off_x + self.draw_start[0] * self.ratio
            y1_c = off_y + self.draw_start[1] * self.ratio
            x2_c = off_x + rx * self.ratio
            y2_c = off_y + ry * self.ratio
            
            colors = {'global_blur': "#00ff00", 'zoom': "#ffff00", 'crop': "#ffffff", 'blur': "#ff0000"}
            color = colors.get(self.draw_mode, "#ff0000")
            self.canvas.create_rectangle(x1_c, y1_c, x2_c, y2_c, outline=color, width=3, dash=(5,5))
            return
        
        if self.dragging_layer and self.selected_layer_idx >= 0:
            layer = self.steps[self.current_idx].layers[self.selected_layer_idx]
            d = layer.data
            if self.active_handle:
                if self.active_handle == 'nw':
                    diff = max(d['x']-rx, d['y']-ry); d['x'] -= diff; d['y'] -= diff; d['size'] += diff
                elif self.active_handle == 'ne':
                    diff = max(rx-(d['x']+d['size']), d['y']-ry); d['y'] -= diff; d['size'] += diff
                elif self.active_handle == 'sw':
                    diff = max(d['x']-rx, ry-(d['y']+d['size'])); d['x'] -= diff; d['size'] += diff
                elif self.active_handle == 'se':
                    d['size'] = max(50, max(rx - d['x'], ry - d['y']))
                d['size'] = max(50, d['size'])
            elif layer.type == 'click':
                layer.data['x'], layer.data['y'] = rx, ry
                for l in self.steps[self.current_idx].layers:
                    if l.type == 'zoom': l.data['target_x'], l.data['target_y'] = rx, ry
            elif layer.type == 'zoom':
                layer.data['x'], layer.data['y'] = rx - layer.data['size']//2, ry - layer.data['size']//2
            elif layer.type == 'text':
                layer.data['x'], layer.data['y'] = rx, ry
            self.draw_step()

    def on_mouseup(self, e):
        if self.draw_mode and self.draw_start:
            rx, ry = self.get_img_coords(e.x, e.y)
            x1, y1 = min(self.draw_start[0], rx), min(self.draw_start[1], ry)
            x2, y2 = max(self.draw_start[0], rx), max(self.draw_start[1], ry)
            
            if abs(x2-x1) > 10:
                if self.draw_mode == 'blur':
                    self.steps[self.current_idx].layers.append(Layer('blur', {'coords': (x1, y1, x2, y2)}, "Zensur"))
                elif self.draw_mode == 'global_blur':
                    self.global_layers.append(Layer('blur', {'coords': (x1, y1, x2, y2)}, "Globaler Blur", True))
                elif self.draw_mode == 'zoom':
                    target = (self.steps[self.current_idx].x, self.steps[self.current_idx].y)
                    self.steps[self.current_idx].layers.append(Layer('zoom', 
                        {'x': x1, 'y': y1, 'size': max(x2-x1, y2-y1), 'target_x': target[0], 'target_y': target[1]}, "Detail Zoom"))
                elif self.draw_mode == 'crop':
                    # Global crop logic
                    if self.global_crop:
                        # Add to existing crop
                        ox, oy = self.global_crop[0], self.global_crop[1]
                        self.global_crop = (ox+x1, oy+y1, ox+x2, oy+y2)
                    else:
                        self.global_crop = (x1, y1, x2, y2)
                self.refresh_layer_list(); self.draw_step()
            self.draw_mode = None; self.canvas.config(cursor="")
            for btn in [self.btn_zoom, self.btn_blur, self.btn_crop, self.btn_global_blur]:
                btn.config(bg="#2a2a2a" if btn != self.btn_global_blur else "#2c3e50", relief=tk.FLAT)
        self.dragging_layer = False; self.draw_start = None

    def delete_current(self):
        if messagebox.askyesno("Löschen", "Schritt entfernen?"):
            self.steps.pop(self.current_idx); self.current_idx = max(0, self.current_idx-1)
            if self.steps: self.select_step(self.current_idx)
            else: self.canvas.delete("all"); self.update_thumbnails()
    
    def save_project(self):
        name = simpledialog.askstring("Projekt speichern", "Name des Projekts:", parent=self.root)
        if not name: return
        
        base_path = os.path.join(os.getcwd(), "projects", name)
        img_path = os.path.join(base_path, "images")
        os.makedirs(img_path, exist_ok=True)
        
        project_data = {
            "global_crop": self.global_crop,
            "global_layers": [],
            "steps": []
        }
        
        for gl in self.global_layers:
            project_data["global_layers"].append({"type": gl.type, "data": gl.data, "label": gl.label})
            
        for i, s in enumerate(self.steps):
            filename = f"step_{i}.png"
            cv2.imwrite(os.path.join(img_path, filename), s.raw_img)
            
            step_layers = []
            for l in s.layers:
                step_layers.append({"type": l.type, "data": l.data, "label": l.label})
            
            project_data["steps"].append({
                "image": filename,
                "description": s.description,
                "layers": step_layers
            })
            
        with open(os.path.join(base_path, "project.json"), "w", encoding="utf-8") as f:
            json.dump(project_data, f, indent=4)
        
        messagebox.showinfo("Gespeichert", f"Projekt '{name}' wurde erfolgreich gespeichert!")
        if hasattr(self.save_callback, '__self__') and hasattr(self.save_callback.__self__, 'update_project_list'):
            self.save_callback.__self__.update_project_list()

    def save_all(self): self.save_callback(self.steps, self.global_layers); self.root.destroy()
    
    def get_img_coords(self, ex, ey):
        cw, ch = self.canvas.winfo_width(), self.canvas.winfo_height()
        if self.global_crop:
            h, w = self.global_crop[3]-self.global_crop[1], self.global_crop[2]-self.global_crop[0]
        else:
            h, w = self.steps[self.current_idx].raw_img.shape[:2]
        off_x, off_y = (cw - w*self.ratio)/2, (ch - h*self.ratio)/2
        return int((ex - off_x)/self.ratio), int((ey - off_y)/self.ratio)

class StatusOverlay:
    def __init__(self, root):
        self.window = tk.Toplevel(root); self.window.overrideredirect(True)
        self.window.attributes("-topmost", True); self.window.geometry("200x50+20+20")
        self.window.configure(bg=BG_SIDE)
        self.label = tk.Label(self.window, text="VORSCHAU", fg="white", bg=BG_SIDE, font=("Segoe UI", 11, "bold"))
        self.label.pack(expand=True, fill=tk.BOTH)
        self.window.bind("<Button-1>", self.st); self.window.bind("<B1-Motion>", self.mv)
        self.x=0; self.y=0
    def st(self, e): self.x=e.x; self.y=e.y
    def mv(self, e): self.window.geometry(f"+{self.window.winfo_x()+(e.x-self.x)}+{self.window.winfo_y()+(e.y-self.y)}")
    def update_status(self, s):
        colors = {"REC": ("#e53935", "🔴 REC (Ctrl+Alt+E)"), "PAUSE": ("#ffa000", "⏸ PAUSE"), "STOP": (BG_SIDE, "⏹ STOP")}
        bg, text = colors.get(s, (BG_SIDE, s))
        self.window.configure(bg=bg); self.label.configure(text=text, bg=bg)
    def hide(self): self.window.withdraw(); self.window.update()
    def show(self): self.window.deiconify(); self.window.update()

class SmartRecorder:
    def __init__(self, root):
        self.root = root; self.root.title("AutoGuide Pro"); self.root.geometry("500x350"); self.root.configure(bg=BG_DARK)
        self.is_recording = False; self.is_paused = False; self.steps = []
        self.global_layers = []
        self.global_crop = None # (x1, y1, x2, y2)
        self.overlay = StatusOverlay(root); self.last_shortcut = None; self.last_shortcut_time = 0
        self.setup_ui(); self.start_listeners()

    def setup_ui(self):
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
        main_frame = tk.Frame(self.root, bg=BG_DARK)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Left Side: Projects
        self.side_proj = tk.Frame(main_frame, bg=BG_SIDE, width=200)
        self.side_proj.pack(side=tk.LEFT, fill=tk.Y)
        tk.Label(self.side_proj, text="PROJEKTE", fg=ACCENT, bg=BG_SIDE, font=("Segoe UI", 10, "bold")).pack(pady=10)
        
        self.proj_list = tk.Listbox(self.side_proj, bg="#2a2a2a", fg="white", borderwidth=0, highlightthickness=0, selectbackground=ACCENT)
        self.proj_list.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        btn_frame = tk.Frame(self.side_proj, bg=BG_SIDE)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        tk.Button(btn_frame, text="📂 Öffnen", bg="#444", fg="white", relief=tk.FLAT, command=self.load_project).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=2)
        tk.Button(btn_frame, text="🗑️", bg="#444", fg="#ff4444", relief=tk.FLAT, command=self.delete_project).pack(side=tk.LEFT, padx=2)
        
        # Right Side: Recording
        f = tk.Frame(main_frame, bg=BG_DARK)
        f.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        
        tk.Label(f, text="AUTOGUIDE PRO", font=("Segoe UI", 26, "bold"), fg=ACCENT, bg=BG_DARK).pack(pady=(50, 0))
        self.btn = tk.Button(f, text="NEUE AUFNAHME", bg=ACCENT, fg="white", relief=tk.FLAT, font=("Segoe UI", 12, "bold"), padx=50, pady=15, command=self.toggle)
        self.btn.pack(pady=25)
        tk.Label(f, text="Shortcuts: Ctrl+Alt+ S(tart), P(ause), E(ditor)", fg="#666", bg=BG_DARK).pack()
        
        self.update_project_list()

    def update_project_list(self):
        self.proj_list.delete(0, tk.END)
        path = os.path.join(os.getcwd(), "projects")
        if not os.path.exists(path): os.makedirs(path)
        for d in os.listdir(path):
            if os.path.isdir(os.path.join(path, d)):
                self.proj_list.insert(tk.END, d)

    def load_project(self):
        sel = self.proj_list.curselection()
        if not sel: return
        name = self.proj_list.get(sel[0])
        path = os.path.join(os.getcwd(), "projects", name)
        
        with open(os.path.join(path, "project.json"), "r", encoding="utf-8") as f:
            data = json.load(f)
            
        self.global_crop = data.get("global_crop")
        self.global_layers = [Layer(l["type"], l["data"], l["label"], True) for l in data.get("global_layers", [])]
        self.steps = []
        
        for s_data in data["steps"]:
            img_file = os.path.join(path, "images", s_data["image"])
            img = cv2.imread(img_file)
            # Find the click layer to reconstruct the Step object
            click = next((l for l in s_data["layers"] if l["type"] == "click"), None)
            if click:
                step = Step(img, click["data"]["x"], click["data"]["y"], None, click["label"])
                step.description = s_data["description"]
                step.layers = [Layer(l["type"], l["data"], l["label"]) for l in s_data["layers"]]
                self.steps.append(step)
        
        self.overlay.hide()
        StepEditor(self.root, self.steps, self.global_layers, self.global_crop, self.final_export)

    def delete_project(self):
        sel = self.proj_list.curselection()
        if not sel: return
        name = self.proj_list.get(sel[0])
        if messagebox.askyesno("Löschen", f"Projekt '{name}' wirklich löschen?"):
            shutil.rmtree(os.path.join(os.getcwd(), "projects", name))
            self.update_project_list()

    def start_listeners(self):
        h = {'<ctrl>+<alt>+s': lambda: self.root.after(0, self.start_rec),
             '<ctrl>+<alt>+p': lambda: self.root.after(0, self.pause_rec),
             '<ctrl>+<alt>+e': lambda: self.root.after(0, self.stop_rec)}
        keyboard.GlobalHotKeys(h).start()
        keyboard.Listener(on_press=self.on_kp).start()

    def on_kp(self, key):
        try:
            kn = self.get_key_name(key)
            if not kn or any(m in kn for m in ['CTRL', 'ALT', 'SHIFT']): return 
            m = []
            if win32api.GetKeyState(win32con.VK_CONTROL) < 0: m.append("STRG")
            if win32api.GetKeyState(win32con.VK_MENU) < 0: m.append("ALT")
            if win32api.GetKeyState(win32con.VK_SHIFT) < 0: m.append("SHIFT")
            if m: self.last_shortcut = f"[{'+'.join(m)}] + [{kn}]"; self.last_shortcut_time = time.time()
            elif kn in ['ENTER', 'TAB', 'ESC']: self.last_shortcut = f"[{kn}]"; self.last_shortcut_time = time.time()
        except: pass

    def get_key_name(self, k):
        if hasattr(k, 'char') and k.char: return k.char.upper()
        return {'enter':'ENTER', 'tab':'TAB', 'esc':'ESC'}.get(getattr(k, 'name', None))

    def toggle(self): self.stop_rec() if self.is_recording else self.start_rec()
    def start_rec(self):
        if not self.is_recording:
            self.is_recording = True; self.steps = []
            self.overlay.update_status("REC"); self.root.iconify()
            self.mouse_l = mouse.Listener(on_click=self.on_click); self.mouse_l.start()
            self.btn.config(text="STOP (Ctrl+Alt+E)", bg="#8e0000")

    def pause_rec(self):
        if self.is_recording: self.is_paused = not self.is_paused; self.overlay.update_status("PAUSE" if self.is_paused else "REC")
    
    def stop_rec(self):
        if self.is_recording:
            self.is_recording = False
            if hasattr(self, 'mouse_l'): self.mouse_l.stop()
            self.root.deiconify()
            self.overlay.hide() # Hide overlay in editor
            self.btn.config(text="START (Ctrl+Alt+S)", bg=ACCENT)
            StepEditor(self.root, self.steps, self.global_layers, self.global_crop, self.final_export)

    def on_click(self, x, y, button, pressed):
        if not self.is_recording or self.is_paused or pressed: return
        self.overlay.hide(); time.sleep(0.12)
        raw = ImageGrab.grab(); self.overlay.show()
        img = cv2.cvtColor(np.array(raw), cv2.COLOR_RGB2BGR)
        sc = self.last_shortcut if (time.time() - self.last_shortcut_time) < 2.5 else None
        lbl = "Rechtsklick" if button == mouse.Button.right else "Linksklick"
        self.steps.append(Step(img, x, y, button, lbl, sc))

    def final_export(self, steps, global_layers):
        if not steps: return
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not path: return
        doc = Document()
        doc.add_heading('Anleitung', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        ox, oy = (self.global_crop[0], self.global_crop[1]) if self.global_crop else (0, 0)
        
        for i, s in enumerate(steps):
            raw = s.raw_img
            if self.global_crop:
                x1, y1, x2, y2 = self.global_crop
                raw = raw[y1:y2, x1:x2]
            img = raw.copy()
            
            # Render Globals
            for l in global_layers:
                if l.type == 'blur':
                    c = l.data['coords']
                    x1, y1, x2, y2 = c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy
                    ih, iw = img.shape[:2]
                    x1, y1 = max(0, x1), max(0, y1)
                    x2, y2 = min(iw, x2), min(ih, y2)
                    if x2 > x1 and y2 > y1:
                        img[y1:y2, x1:x2] = cv2.GaussianBlur(img[y1:y2, x1:x2], (75,75), 40)
                elif l.type == 'text':
                    d = l.data
                    tx, ty = d['x']-ox, d['y']-oy
                    cv2.putText(img, d['text'], (tx, ty), cv2.FONT_HERSHEY_SIMPLEX, d.get('size', 1.0), (0,0,0), int(d.get('size', 1.0)*5+2), cv2.LINE_AA)
                    cv2.putText(img, d['text'], (tx, ty), cv2.FONT_HERSHEY_SIMPLEX, d.get('size', 1.0), d.get('color', (255,255,255)), int(d.get('size', 1.0)*2), cv2.LINE_AA)
            
            # Render Locals
            for layer in s.layers:
                if not layer.visible: continue
                if layer.type == 'blur':
                    c = layer.data['coords']
                    x1, y1, x2, y2 = c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy
                    ih, iw = img.shape[:2]
                    x1, y1 = max(0, x1), max(0, y1)
                    x2, y2 = min(iw, x2), min(ih, y2)
                    if x2 > x1 and y2 > y1:
                        img[y1:y2, x1:x2] = cv2.GaussianBlur(img[y1:y2, x1:x2], (75,75), 40)
                elif layer.type == 'click':
                    x, y = layer.data['x']-ox, layer.data['y']-oy
                    cv2.circle(img, (x, y), 38, (255, 255, 255), 6, cv2.LINE_AA)
                    cv2.circle(img, (x, y), 32, (0,0,255), 4, cv2.LINE_AA)
                    cv2.putText(img, str(i+1), (x+18, y-18), cv2.FONT_HERSHEY_SIMPLEX, 1.3, (0,0,0), 6, cv2.LINE_AA)
                    cv2.putText(img, str(i+1), (x+18, y-18), cv2.FONT_HERSHEY_SIMPLEX, 1.3, (255,255,255), 3, cv2.LINE_AA)
                elif layer.type == 'text':
                    d = layer.data
                    tx, ty = d['x']-ox, d['y']-oy
                    cv2.putText(img, d['text'], (tx, ty), cv2.FONT_HERSHEY_SIMPLEX, d.get('size', 1.0), (0,0,0), int(d.get('size', 1.0)*5+2), cv2.LINE_AA)
                    cv2.putText(img, d['text'], (tx, ty), cv2.FONT_HERSHEY_SIMPLEX, d.get('size', 1.0), d.get('color', (255,255,255)), int(d.get('size', 1.0)*2), cv2.LINE_AA)
                elif layer.type == 'zoom':
                    d = layer.data; sz = d['size']
                    tx_o, ty_o = d['target_x']-ox, d['target_y']-oy
                    roi = s.raw_img[max(0, d['target_y']-sz//4):min(s.raw_img.shape[0], d['target_y']+sz//4), 
                                    max(0, d['target_x']-sz//4):min(s.raw_img.shape[1], d['target_x']+sz//4)]
                    if roi.size>0:
                        zoomed_roi = cv2.resize(roi, (sz,sz))
                        cv2.circle(zoomed_roi, (sz//2, sz//2), 30, (0,0,255), 4, cv2.LINE_AA)
                        cv2.circle(zoomed_roi, (sz//2, sz//2), 6, (255,255,255), -1, cv2.LINE_AA)
                        
                        h, w = img.shape[:2]
                        zy, zx = d['y']-oy, d['x']-ox
                        y_start, y_end = max(0, zy), min(h, zy+sz)
                        x_start, x_end = max(0, zx), min(w, zx+sz)
                        if y_end > y_start and x_end > x_start:
                            img[y_start:y_end, x_start:x_end] = zoomed_roi[y_start-zy:y_end-zy, x_start-zx:x_end-zx]
                            cv2.rectangle(img, (zx, zy), (zx+sz, zy+sz), (255,255,255), 3)
                        
                        cx, cy = zx+sz//2, zy+sz//2
                        cv2.arrowedLine(img, (cx, cy), (tx_o, ty_o), (255,255,255), 6, tipLength=0.1)

            tmp = f"tmp_{i}.png"; cv2.imwrite(tmp, img)
            doc.add_heading(f"Schritt {i+1}", level=1)
            if s.description: doc.add_paragraph(s.description)
            doc.add_picture(tmp, width=Inches(6.0)); os.remove(tmp)
        doc.save(path); messagebox.showinfo("Erfolg", "Dokument wurde erfolgreich erstellt!")

if __name__ == "__main__":
    root = tk.Tk(); app = SmartRecorder(root); root.mainloop()