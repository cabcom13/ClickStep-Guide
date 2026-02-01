import sys
import os
import time
import json
import shutil
import math
import uuid
import queue
import ctypes
from datetime import datetime

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QPushButton, QLabel, QFileDialog, QListWidget, QListWidgetItem, 
    QDockWidget, QToolBar, QSlider, QInputDialog, QMessageBox, 
    QGraphicsScene, QGraphicsView, QGraphicsItem, QGraphicsRectItem,
    QGraphicsPixmapItem, QGraphicsTextItem, QGraphicsLineItem, QMenu,
    QSpinBox, QColorDialog, QFontComboBox, QComboBox, QDialog, QLineEdit, 
    QDialogButtonBox, QAbstractItemView, QCheckBox, QTextEdit, QFrame,
    QFormLayout, QGroupBox, QRadioButton, QButtonGroup, QProgressBar
)
from PyQt6.QtCore import Qt, QTimer, QPointF, QRectF, QRect, QSize, pyqtSignal, QObject, QLineF, QThread
from PyQt6.QtGui import (
    QPixmap, QPainter, QPen, QColor, QFont, QAction, QIcon, 
    QBrush, QImage, QPainterPath
)

from pynput import mouse, keyboard
from PIL import ImageGrab
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from clickstep.config.settings import AppSettings, ClickMarkerSettings
from clickstep.models import Step, Layer
from clickstep.ui.dialogs import ModernDialog, SettingsDialog
from clickstep.ui.overlay import RecordingOverlay
from clickstep.ui.styles import get_recorder_stylesheet, get_editor_stylesheet
from clickstep.utils.helpers import resource_path
from clickstep.editor.editor import ProEditor

# ==================== RECORDER LOGIC ====================

class RecordingSignal(QObject):
    click_detected = pyqtSignal(int, int, str, object)

class RecordingThread(QThread):
    def __init__(self):
        super().__init__()
        self.signals = RecordingSignal()
        self.is_running = False
        self.event_queue = queue.Queue()

    def run(self):
        self.is_running = True
        self.mouse_listener = mouse.Listener(on_click=self.on_click)
        self.mouse_listener.start()
        
        while self.is_running:
            try:
                # Wait for click event with short timeout to check is_running
                # Non-blocking allows us to exit clean
                x, y = self.event_queue.get(timeout=0.05)
                
                # Perform the delay and capture here in the thread, NOT in the hook
                time.sleep(0.12)
                
                raw = ImageGrab.grab()
                img = cv2.cvtColor(np.array(raw), cv2.COLOR_RGB2BGR)
                self.signals.click_detected.emit(x, y, "Click", img)
                
            except queue.Empty:
                continue
                
        self.mouse_listener.stop()

    def on_click(self, x, y, button, pressed):
        if not pressed: return
        # FAST: Just put into queue and return immediately to unblock system
        self.event_queue.put((x, y))

class ProRecorder(QMainWindow):
    hotkey_signal = pyqtSignal(str)
    
    def __init__(self):
        super().__init__()
        
        # Windows Taskbar Icon Fix
        if os.name == 'nt':
            try:
                ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("clickstep.guide.pro.1")
            except: pass
            
        self.setWindowTitle("ClickStep Guide")
        self.setFixedSize(1280, 720) # Fixed size to prevent background distortion (16:9)
        self.setWindowIcon(QIcon(resource_path("icon.ico")))
        
        self.settings = AppSettings()
        self.apply_app_theme()
        
        self.steps = []
        self.global_layers = []
        self.global_crop = None
        self.is_recording = False
        self.overlay = RecordingOverlay() # Create overlay
        
        self.recording_thread = RecordingThread()
        self.recording_thread.signals.click_detected.connect(self.handle_click)
        
        # Connect hotkey signal to ensure thread-safe UI calls
        self.hotkey_signal.connect(self.handle_hotkey)
        
        self.setup_ui()
        self.setup_hotkeys()
        self.update_project_list()

    def closeEvent(self, event):
        """Cleanup thread before closing"""
        if hasattr(self, 'recording_thread'):
            self.recording_thread.is_running = False
            self.recording_thread.wait(500)
        event.accept()

    def get_project_dir(self):
        """Returns the project directory in Local AppData for Store compliance"""
        base = os.environ.get('LOCALAPPDATA', os.path.expanduser("~"))
        path = os.path.join(base, "ClickStepGuide", "projects")
        os.makedirs(path, exist_ok=True)
        return path

    def apply_app_theme(self):
        """Apply modern, premium theme with clean gradients and glassmorphism"""
        bg_path = resource_path("background.png").replace("\\", "/")
        has_bg = os.path.exists(resource_path("background.png"))
        
        # Modern Premium Theme (works for both light/dark, optimized for clean look)
        style = """
            QMainWindow { 
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #f8fafc, stop:0.5 #e0e7ff, stop:1 #f1f5f9);
            }
            
            QDialog { background-color: rgba(255, 255, 255, 0.95); }
            
            QWidget { 
                color: #1e293b; 
                font-family: 'Segoe UI', 'Inter', sans-serif; 
                font-size: 13px; 
            }
            
            /* Sidebar - Glassmorphism */
            #Sidebar { 
                background: rgba(255, 255, 255, 0.7);
                border-right: 1px solid rgba(226, 232, 240, 0.8);
                backdrop-filter: blur(10px);
            }
            
            #SidebarHeader { 
                color: #64748b; 
                font-weight: 700; 
                font-size: 11px; 
                letter-spacing: 2px; 
                margin-bottom: 8px;
                padding: 8px 0;
            }
            
            /* Project List */
            #ProjectList { 
                background: transparent; 
                border: none; 
                outline: none; 
            }
            
            #ProjectList::item { 
                padding: 14px 12px; 
                border-radius: 8px; 
                margin-bottom: 4px; 
                color: #475569;
                font-weight: 500;
                background: transparent;
            }
            
            #ProjectList::item:selected { 
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #2563eb, stop:1 #3b82f6);
                color: white; 
                font-weight: 600;
                border-left: 3px solid #60a5fa;
            }
            
            #ProjectList::item:hover:!selected { 
                background: rgba(226, 232, 240, 0.5);
            }
            
            /* Buttons - Modern Style */
            QPushButton { 
                background: rgba(255, 255, 255, 0.9);
                border: 1px solid rgba(226, 232, 240, 0.8);
                padding: 11px 18px; 
                border-radius: 8px; 
                font-weight: 600; 
                color: #334155;
            }
            
            QPushButton:hover { 
                background: rgba(255, 255, 255, 1);
                border-color: #cbd5e1;
                transform: translateY(-1px);
            }
            
            QPushButton:pressed { 
                background: #f1f5f9;
            }
            
            /* Primary Record Button - Hero Style */
            #RecordBtn { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #2563eb, stop:1 #1d4ed8);
                color: white; 
                font-size: 16px; 
                padding: 22px 50px; 
                border: none; 
                border-radius: 14px; 
                font-weight: 700; 
                letter-spacing: 0.5px;
                min-width: 350px;
            }
            
            #RecordBtn:hover { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #3b82f6, stop:1 #2563eb);
            }
            
            /* Sidebar Buttons */
            #PrimarySidebarBtn { 
                background: rgba(255, 255, 255, 0.9);
                border-color: #e2e8f0;
                font-weight: 600;
            }
            
            #PrimarySidebarBtn:hover {
                background: white;
                border-color: #cbd5e1;
            }
            
            #DeleteBtn { 
                background: rgba(254, 226, 226, 0.8);
                color: #dc2626; 
                border: 1px solid rgba(220, 38, 38, 0.2);
                font-weight: 600;
            }
            
            #DeleteBtn:hover { 
                background: rgba(254, 226, 226, 1);
                border-color: #dc2626;
            }
            
            #SettingsBtn {
                background: rgba(255, 255, 255, 0.9);
                border: 1px solid #e2e8f0;
                font-weight: 600;
                color: #475569;
            }
            
            #SettingsBtn:hover {
                background: white;
                border-color: #cbd5e1;
            }
            
            /* Labels */
            #ShortcutLabel { 
                color: #64748b; 
                font-size: 13px; 
                margin-top: 15px;
                font-weight: 500;
            }
            
            #FooterLabel { 
                color: #94a3b8; 
                font-size: 11px; 
                padding-top: 20px; 
                margin-top: 30px; 
                font-weight: 500;
            }
            
            /* Input Fields */
            QLineEdit, QTextEdit { 
                background: rgba(255, 255, 255, 0.9);
                border: 1px solid #e2e8f0;
                border-radius: 8px; 
                padding: 10px; 
                color: #1e293b;
            }
            
            QLineEdit:focus, QTextEdit:focus { 
                border-color: #3b82f6;
                background: white;
            }
            
            /* Group Boxes */
            QGroupBox { 
                font-weight: 600; 
                border: 1px solid #e2e8f0;
                margin-top: 18px; 
                padding-top: 22px; 
                border-radius: 10px;
                background: rgba(255, 255, 255, 0.5);
            }
            
            QGroupBox::title { 
                subcontrol-origin: margin; 
                left: 12px; 
                padding: 0 8px; 
                color: #64748b;
            }
        """
        
        # Override with background image if exists
        if has_bg:
            style = style.replace(
                "QMainWindow { \n                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, \n                    stop:0 #f8fafc, stop:0.5 #e0e7ff, stop:1 #f1f5f9);\n            }",
                f"QMainWindow {{ border-image: url({bg_path}) 0 0 0 0 stretch stretch; }}"
            )
            # Ensure text is readable on custom background
            style += """
                #ShortcutLabel, #FooterLabel { 
                    color: #1e293b; 
                    font-weight: 700;
                    text-shadow: 0 1px 2px rgba(255,255,255,0.8);
                }
                QLabel { 
                    color: #1e293b;
                    font-weight: 600;
                }
            """
        
        self.setStyleSheet(style)

    def setup_hotkeys(self):
        """Setup global keyboard shortcuts from settings - uses signals for thread safety"""
        if hasattr(self, 'hotkey_thread'):
            try:
                self.hotkey_thread.stop()
            except: pass
            
        h = {
            self.settings.shortcut_record: lambda: self.hotkey_signal.emit("record"),
            self.settings.shortcut_editor: lambda: self.hotkey_signal.emit("editor")
        }
        try:
            self.hotkey_thread = keyboard.GlobalHotKeys(h)
            self.hotkey_thread.start()
        except Exception as e:
            print(f"Hotkey Setup Error: {e}")

    def handle_hotkey(self, action):
        """Handler for hotkeys, executes in the UI thread"""
        if action == "record":
            self.btn_record.click()
        elif action == "editor":
            self.open_editor()

    def setup_ui(self):
        self.center_widget = QWidget()
        self.setCentralWidget(self.center_widget)
        
        # Main Layout: Horizontal (Left Sidebar | Right Content)
        main_layout = QHBoxLayout(self.center_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # --- LEFT SIDEBAR (Static Project List) ---
        self.sidebar = QWidget()
        self.sidebar.setObjectName("Sidebar")
        self.sidebar.setFixedWidth(280)
        
        side_layout = QVBoxLayout(self.sidebar)
        side_layout.setContentsMargins(15, 20, 15, 20)
        side_layout.setSpacing(10)
        
        lbl_proj = QLabel("PROJEKTE")
        lbl_proj.setObjectName("SidebarHeader")
        side_layout.addWidget(lbl_proj)
        
        self.proj_list = QListWidget()
        self.proj_list.setObjectName("ProjectList")
        side_layout.addWidget(self.proj_list)
        
        # Sidebar Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(5)
        
        btn_open = QPushButton("📂 Öffnen")
        btn_open.setObjectName("PrimarySidebarBtn")
        btn_open.setFixedHeight(40)
        btn_open.clicked.connect(self.load_project)
        
        btn_del = QPushButton("🗑️")
        btn_del.setObjectName("DeleteBtn")
        btn_del.setFixedSize(40, 40)
        btn_del.clicked.connect(self.delete_project)
        
        btn_layout.addWidget(btn_open)
        btn_layout.addWidget(btn_del)
        side_layout.addLayout(btn_layout)
        
        # Spacer
        side_layout.addStretch()
        
        # Settings Button
        self.btn_settings = QPushButton("⚙️ Einstellungen")
        self.btn_settings.setObjectName("SettingsBtn")
        self.btn_settings.setFixedHeight(45)
        self.btn_settings.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_settings.clicked.connect(self.open_settings)
        side_layout.addWidget(self.btn_settings)
        
        main_layout.addWidget(self.sidebar)
        
        # --- RIGHT CONTENT AREA ---
        content_area = QWidget()
        self.layout = QVBoxLayout(content_area)
        self.layout.setContentsMargins(50, 50, 50, 50)
        self.layout.setSpacing(30)
        
        if os.path.exists(resource_path("logo.png")):
             # Logo removed as requested, keeping code block for structure but logic is disabled
             pass 
        
        self.layout.addStretch(3) # Spacer (Top - weight 3)
        
        self.btn_record = QPushButton("NEUE AUFNAHME STARTEN")
        self.btn_record.setObjectName("RecordBtn")
        self.btn_record.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_record.clicked.connect(self.toggle_recording)
        self.layout.addWidget(self.btn_record, 0, Qt.AlignmentFlag.AlignCenter)
        
        self.lbl_shortcuts = QLabel(f"Shortcuts: {self.settings.shortcut_record.upper()} (Start/Stop) | {self.settings.shortcut_editor.upper()} (Editor)")
        self.lbl_shortcuts.setObjectName("ShortcutLabel")
        self.lbl_shortcuts.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.layout.addWidget(self.lbl_shortcuts)
        
        self.layout.addStretch(1) # Bottom spacer (weight 1)
        
        # Attribution Footer
        footer = QLabel("© 2026 ClickStep Guide | High-Performance Documentation Engine")
        footer.setObjectName("FooterLabel")
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.layout.addWidget(footer)
        
        main_layout.addWidget(content_area)

    def toggle_recording(self):
        if not self.is_recording:
            self.is_recording = True
            self.steps = []
            self.btn_record.setText("AUFNAHME STOPPEN")
            self.showMinimized()
            
            # Reset and show overlay
            self.overlay.update_steps(0)
            screen = QApplication.primaryScreen().geometry()
            self.overlay.setGeometry(screen.width() - 260, 30, 240, 60)
            self.overlay.show() 
            self.recording_thread.start()
        else:
            self.stop_recording()

    def stop_recording(self):
        self.is_recording = False
        self.recording_thread.is_running = False
        self.overlay.hide() # Hide overlay
        self.btn_record.setText("NEUE AUFNAHME STARTEN")
        self.showNormal()
        
        # Give thread a moment to finish
        self.recording_thread.wait(500) # Wait max 500ms
        
        if self.steps:
            self.open_editor()

    def handle_click(self, x, y, label, img):
        # Create modular Step and add initial 'click' Layer
        new_step = Step(img, x, y, label)
        new_step.layers.append(Layer('click', {'x': x, 'y': y}, label))
        self.steps.append(new_step)
        self.overlay.update_steps(len(self.steps))

    def open_editor(self):
        self.hide() # Hide main recorder
        self.editor = ProEditor(self.steps, self.global_layers, self.global_crop, self.final_export, parent_window=self, settings=self.settings)
        self.editor.show()

    def open_settings(self):
        """Open settings dialog and apply changes"""
        dlg = SettingsDialog(self.settings, self)
        if dlg.exec():
            new_data = dlg.get_settings()
            
            # Update settings object
            self.settings.theme = new_data["theme"]
            self.settings.shortcut_record = new_data["shortcut_record"]
            self.settings.shortcut_editor = new_data["shortcut_editor"]
            
            # Persistence
            self.settings.save()
            
            # Apply immediate changes
            self.apply_app_theme()
            self.setup_hotkeys()
            self.lbl_shortcuts.setText(f"Shortcuts: {self.settings.shortcut_record.upper()} (Start/Stop) | {self.settings.shortcut_editor.upper()} (Editor)")
            
            QMessageBox.information(self, "Erfolg", "Einstellungen wurden gespeichert und angewendet!")

    def final_export(self, steps, global_layers, global_crop):
        if not steps: # This is a refresh signal for project list
            self.update_project_list()
            return
            
        path, _ = QFileDialog.getSaveFileName(self, "Export", "", "Word (*.docx)")
        if not path: return
        
        prog = QProgressBar()
        prog.setMaximum(len(steps))
        self.layout.addWidget(prog)
        
        try:
            doc = Document()
            doc.add_heading('ClickStep Guide - Anleitung', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, s in enumerate(steps):
                # 1. Prepare Base Image (Crop)
                raw = s.raw_img
                if global_crop:
                    x1, y1, x2, y2 = global_crop
                    raw = raw[y1:y2, x1:x2]
                
                canvas = raw.copy()
                ox, oy = (global_crop[0], global_crop[1]) if global_crop else (0, 0)
                
                # 2. Render all annotations (using the same logic as our editor's paint methods)
                # First Global Blur
                for gl in global_layers:
                    if gl.type == 'blur':
                        c = gl.data['coords']
                        s = gl.data.get('strength', 40)
                        self.render_blur_cv2(canvas, (c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy), True, s)
                
                # Then Step Layers
                for l in s.layers:
                    if l.type == 'blur':
                        c = l.data['coords']
                        s = l.data.get('strength', 40)
                        self.render_blur_cv2(canvas, (c[0]-ox, c[1]-oy, c[2]-ox, c[3]-oy), False, s)
                    elif l.type == 'click':
                        self.render_click_cv2(canvas, l.data['x']-ox, l.data['y']-oy, i+1)
                    elif l.type == 'zoom':
                        self.render_zoom_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'infobox':
                        self.render_infobox_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'arrow':
                        self.render_arrow_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'icon':
                        self.render_icon_cv2(canvas, l.data, ox, oy)
                    elif l.type == 'text':
                        self.render_text_cv2(canvas, l.data, ox, oy)
                
                # 3. Render Permanent Watermark
                self.render_watermark_cv2(canvas)
                
                # 4. Add to Word
                tmp_file = f"export_tmp_{i}.png"
                cv2.imwrite(tmp_file, canvas)
                doc.add_heading(f"Schritt {i+1}", level=1)
                if s.description: doc.add_paragraph(s.description)
                doc.add_picture(tmp_file, width=Inches(6))
                os.remove(tmp_file)
                prog.setValue(i+1)

            # 4. Add Copyright & Info Footer at the end
            doc.add_page_break()
            doc.add_heading('Über dieses Dokument', level=1)
            p = doc.add_paragraph()
            p.add_run('Diese Anleitung wurde professionell erstellt mit ').italic = True
            p.add_run('ClickStep Guide').bold = True
            p.add_run('.')
            doc.add_paragraph(f"Erstellungsdatum: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
            doc.add_paragraph("© 2026 AutoGuide Automationsysteme")

            doc.save(path)
            QMessageBox.information(self, "Erfolg", f"Export nach {path} abgeschlossen!")
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Export fehlgeschlagen: {str(e)}")
        finally:
            self.layout.removeWidget(prog)
            prog.deleteLater()

    # OpenCV Rendering Helpers for Export
    def render_blur_cv2(self, img, coords, is_global, strength=40):
        x1, y1, x2, y2 = coords
        h, w = img.shape[:2]
        x1, y1, x2, y2 = max(0, x1), max(0, y1), min(w, x2), min(h, y2)
        if x2 > x1 and y2 > y1:
            roi = img[y1:y2, x1:x2]
            k = strength | 1
            img[y1:y2, x1:x2] = cv2.GaussianBlur(roi, (k, k), 0)
            # No border in final export for cleaner look
            # cv2.rectangle(img, (x1, y1), (x2, y2), (0, 255, 0) if is_global else (255, 255, 255), 2)

    def render_click_cv2(self, img, x, y, num):
        settings = ClickMarkerSettings()
        
        # Use settings for rendering
        radius = int(settings.size * 0.6)
        glow_radius = int(settings.size * 0.9)
        
        # Check if transparent
        is_transparent = settings.color.alpha() == 0
        
        # Glow (optional, skip if transparent)
        if settings.show_glow and not is_transparent:
            col_bgr = (settings.color.blue(), settings.color.green(), settings.color.red())
            col_glow = tuple(int(c * 0.7) for c in col_bgr)  # Darker for glow
            cv2.circle(img, (x, y), glow_radius, col_glow, -1, lineType=cv2.LINE_AA)
        
        # Main circle fill (skip if transparent)
        if not is_transparent:
            col_bgr = (settings.color.blue(), settings.color.green(), settings.color.red())
            cv2.circle(img, (x, y), radius, col_bgr, -1, lineType=cv2.LINE_AA)
        
        # Border (always visible)
        cv2.circle(img, (x, y), radius, (255, 255, 255), settings.border_width, lineType=cv2.LINE_AA)
        
        # Number
        font_scale = settings.number_size / 16.0
        text_offset_x = int(settings.size * 0.25)
        text_offset_y = int(settings.size * 0.2)
        
        # Text Color from settings
        text_col = getattr(settings, 'text_color', QColor(255, 255, 255))
        # Convert QColor to BGR for OpenCV
        text_bgr = (text_col.blue(), text_col.green(), text_col.red())
        
        # Only draw outline if text is white (for contrast)
        if text_col.name() == "#ffffff":
            cv2.putText(img, str(num), (x-text_offset_x, y+text_offset_y), 
                       cv2.FONT_HERSHEY_SIMPLEX, font_scale, (0, 0, 0), 4, cv2.LINE_AA)
                       
        cv2.putText(img, str(num), (x-text_offset_x, y+text_offset_y), 
                   cv2.FONT_HERSHEY_SIMPLEX, font_scale, text_bgr, 2, cv2.LINE_AA)

    def render_zoom_cv2(self, img, data, ox, oy):
        zx, zy, sz = data['x']-ox, data['y']-oy, data['size']
        tx, ty = data['target_x']-ox, data['target_y']-oy
        col_rgb = data.get('color', (255, 255, 255))
        col_bgr = (col_rgb[2], col_rgb[1], col_rgb[0])
        
        src_sz = sz // 2
        
        # Crop area
        x1, y1 = max(0, tx - src_sz//2), max(0, ty - src_sz//2)
        x2, y2 = min(img.shape[1], tx + src_sz//2), min(img.shape[0], ty + src_sz//2)
        if x2 > x1 and y2 > y1:
            roi = img[y1:y2, x1:x2]
            zoomed = cv2.resize(roi, (sz, sz))
            try:
                img[zy:zy+sz, zx:zx+sz] = zoomed
            except: pass
            cv2.rectangle(img, (zx, zy), (zx+sz, zy+sz), col_bgr, 3)
            cv2.arrowedLine(img, (zx+sz//2, zy+sz//2), (tx, ty), col_bgr, 3)

    def render_infobox_cv2(self, img, data, ox, oy):
        x, y, w, h = data['x']-ox, data['y']-oy, data['w'], data['h']
        tx, ty = data['target_x']-ox, data['target_y']-oy
        col_rgb = data.get('color', (255, 255, 255))
        col_bgr = (col_rgb[2], col_rgb[1], col_rgb[0])
        
        # 1. Background Box
        sub = img[y:y+h, x:x+w]
        white_rect = np.full(sub.shape, 40, dtype=np.uint8) # Dark grey
        res = cv2.addWeighted(sub, 0.2, white_rect, 0.8, 1.0)
        img[y:y+h, x:x+w] = res
        cv2.rectangle(img, (x, y), (x+w, y+h), col_bgr, 2)
        
        # 2. Arrow (from border)
        cx, cy = x + w//2, y + h//2
        dx, dy = tx - cx, ty - cy
        
        # Liang-Barsky for OpenCV (Basic center to edge logic)
        hw, hh = w/2, h/2
        if dx != 0 or dy != 0:
            scale_x = hw / abs(dx) if dx != 0 else 9999
            scale_y = hh / abs(dy) if dy != 0 else 9999
            scale = min(scale_x, scale_y)
            
            sx = int(cx + dx * scale)
            sy = int(cy + dy * scale)
            
            cv2.line(img, (sx, sy), (tx, ty), col_bgr, 2, cv2.LINE_AA)
            cv2.circle(img, (tx, ty), 4, col_bgr, -1)
        
        # 3. Text (Simple wrapping)
        text = data['text']
        font = cv2.FONT_HERSHEY_SIMPLEX
        
        # Map point size to CV2 scale
        f_data = data.get('font', {})
        pt_size = f_data.get('size', 12)
        scale = pt_size / 24.0 # 12pt -> 0.5 scale
        
        t_col = data.get('text_color', (255, 255, 255))
        color = (t_col[2], t_col[1], t_col[0]) # BGR
        thickness = 1 if pt_size < 18 else 2
        
        line_h = int(pt_size * 1.5)
        dy_text = y + line_h
        for line in text.split('\n'):
            cv2.putText(img, line, (x+10, dy_text), font, scale, color, thickness, cv2.LINE_AA)
            dy_text += line_h

    def render_arrow_cv2(self, img, data, ox, oy):
        sx, sy, ex, ey = data['sx']-ox, data['sy']-oy, data['ex']-ox, data['ey']-oy
        col_rgb = data.get('color', (255, 0, 0))
        col_bgr = (col_rgb[2], col_rgb[1], col_rgb[0])
        w = data.get('width', 4)
        
        cv2.arrowedLine(img, (sx, sy), (ex, ey), col_bgr, w, tipLength=0.2)
        
    def render_icon_cv2(self, img, data, ox, oy):
        x, y = data['x']-ox, data['y']-oy
        t_col = data.get('color', (255, 0, 0))
        col_bgr = (t_col[2], t_col[1], t_col[0])
        size = data.get('w', data.get('size', 60))
        
        # Simple Viz: Circle + Text using selected color
        cv2.circle(img, (x+size//2, y+size//2), size//2, col_bgr, 2)
        
        # Mapping meaningful text for CV2 which doesn't do multicolor emoji
        txt_map = {
            'check': 'OK', 
            'cross': 'X', 
            'warn': '!', 
            'info': 'i', 
            'star': '*', 
            'idea': '?',
            'arrow_up': '^',
            'arrow_down': 'v',
            'heart': '<3'
        }
        t = txt_map.get(data.get('type'), '?')
        
        font_scale = size / 40.0
        cv2.putText(img, t, (x+size//4, y+int(size*0.75)), cv2.FONT_HERSHEY_SIMPLEX, font_scale, col_bgr, 3)

    def render_text_cv2(self, img, data, ox, oy):
        x, y = data['x']-ox, data['y']-oy
        color = data.get('color', (255, 255, 255)) # BGR
        cv2.putText(img, data['text'], (x, y), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (0,0,0), 6, cv2.LINE_AA)
        cv2.putText(img, data['text'], (x, y), cv2.FONT_HERSHEY_SIMPLEX, 1.2, color, 2, cv2.LINE_AA)

    def render_watermark_cv2(self, img):
        """Draw a permanent watermark in the top right corner"""
        h, w = img.shape[:2]
        text = "Created with ClickStep Guide"
        font = cv2.FONT_HERSHEY_SIMPLEX
        scale = 0.6
        thickness = 1
        color = (255, 255, 255) # White (BGR)
        
        size = cv2.getTextSize(text, font, scale, thickness)[0]
        # Top Right Position: x = width - text_width - margin, y = text_height + margin
        tx, ty = w - size[0] - 20, size[1] + 20
        
        # Shadow for visibility on bright backgrounds
        cv2.putText(img, text, (tx+1, ty+1), font, scale, (0, 0, 0), thickness+1, cv2.LINE_AA)
        cv2.putText(img, text, (tx, ty), font, scale, color, thickness, cv2.LINE_AA)

    def update_project_list(self):
        """Update project list in sidebar"""
        self.proj_list.clear()
        path = self.get_project_dir()
        if not os.path.exists(path): 
            os.makedirs(path)
        for d in os.listdir(path):
            if os.path.isdir(os.path.join(path, d)):
                item = QListWidgetItem(d)
                self.proj_list.addItem(item)

    def load_project(self):
        """Load selected project"""
        item = self.proj_list.currentItem()
        if not item: return
        
        project_path = os.path.join(self.get_project_dir(), item.text(), "project.json")
        if not os.path.exists(project_path):
            QMessageBox.warning(self, "Fehler", "Projekt-Datei nicht gefunden!")
            return
        
        try:
            with open(project_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Load steps
            self.steps = []
            img_path = os.path.join(self.get_project_dir(), item.text(), "images")
            for step_data in data.get("steps", []):
                img_file = os.path.join(img_path, step_data["image"])
                if os.path.exists(img_file):
                    img = cv2.imread(img_file)
                    step = Step(img, 0, 0, step_data.get("description", ""))
                    step.layers = [] # Reset default click
                    for l_data in step_data.get("layers", []):
                        step.layers.append(Layer(l_data['type'], l_data['data'], l_data.get('label', 'Layer')))
                    
                    # Update step x, y from the first click layer found
                    click_l = next((l for l in step.layers if l.type == 'click'), None)
                    if click_l:
                        step.x, step.y = click_l.data['x'], click_l.data['y']
                        
                    self.steps.append(step)
            
            self.global_crop = data.get("global_crop")
            self.global_layers = []
            for gl_data in data.get("global_layers", []):
                self.global_layers.append(Layer(gl_data['type'], gl_data['data'], gl_data.get('label', 'Global Layer'), True))
            
            if self.steps:
                self.open_editor()
            else:
                QMessageBox.warning(self, "Warnung", "Keine Schritte im Projekt gefunden!")
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Projekt konnte nicht geladen werden: {str(e)}")

    def delete_project(self):
        """Delete selected project"""
        item = self.proj_list.currentItem()
        if not item: return
        
        reply = QMessageBox.question(self, 'Löschen', 
                                    f"Projekt '{item.text()}' wirklich löschen?", 
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                shutil.rmtree(os.path.join(self.get_project_dir(), item.text()))
                self.update_project_list()
                QMessageBox.information(self, "Erfolg", "Projekt gelöscht!")
            except Exception as e:
                QMessageBox.critical(self, "Fehler", f"Löschen fehlgeschlagen: {str(e)}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ProRecorder()
    window.show()
    sys.exit(app.exec())
